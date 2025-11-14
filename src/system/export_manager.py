"""
Export Manager for AAWT
Handles exporting projects to multiple formats: TXT, MD, DOCX, PDF, EPUB, JSON.
"""

import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class ExportManager:
    """Manages project exports to various formats."""
    
    def __init__(self, settings_manager):
        """
        Initialize export manager.
        
        Args:
            settings_manager: Settings manager instance
        """
        self.settings = settings_manager
        logger.info("Export Manager initialized")
    
    def export_project(self, project_data: Dict[str, Any], format: str,
                      output_path: Optional[str] = None, include_metadata: bool = True,
                      include_statistics: bool = True) -> Dict[str, Any]:
        """
        Export project to specified format.
        
        Args:
            project_data: Project data dictionary
            format: Export format (txt, md, docx, pdf, epub, json)
            output_path: Optional custom output path
            include_metadata: Include project metadata
            include_statistics: Include text statistics
        
        Returns:
            Result dictionary with success status and file path
        """
        format = format.lower()
        
        # Check if format is enabled
        if not self.settings.get(f'export.formats_enabled.{format}', True):
            return {'success': False, 'error': f'Export format {format} is disabled'}
        
        # Generate output path if not provided
        if not output_path:
            export_dir = self.settings.get('export.default_export_dir', 'exports')
            Path(export_dir).mkdir(parents=True, exist_ok=True)
            
            project_name = project_data.get('name', 'untitled')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{project_name}_{timestamp}.{format}"
            output_path = Path(export_dir) / filename
        else:
            # Convert string path to Path object
            output_path = Path(output_path)
        
        # Route to appropriate export method
        export_methods = {
            'txt': self._export_txt,
            'md': self._export_markdown,
            'docx': self._export_docx,
            'pdf': self._export_pdf,
            'epub': self._export_epub,
            'json': self._export_json
        }
        
        if format not in export_methods:
            return {'success': False, 'error': f'Unsupported format: {format}'}
        
        try:
            result = export_methods[format](
                project_data, output_path, include_metadata, include_statistics
            )
            
            if result.get('success'):
                logger.info(f"Exported project to {output_path}")
            
            return result
            
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def _prepare_content(self, project_data: Dict, include_metadata: bool, include_statistics: bool) -> str:
        """Prepare content with optional metadata and statistics."""
        lines = []
        
        if include_metadata:
            lines.append("=" * 60)
            lines.append(f"Project: {project_data.get('name', 'Untitled')}")
            lines.append(f"Created: {project_data.get('created_date', 'Unknown')}")
            lines.append(f"Last Modified: {project_data.get('last_modified', 'Unknown')}")
            
            metadata = project_data.get('metadata', {})
            if isinstance(metadata, str):
                try:
                    metadata = json.loads(metadata)
                except:
                    metadata = {}
            
            if metadata:
                if 'genre' in metadata:
                    lines.append(f"Genre: {metadata['genre']}")
                if 'target_audience' in metadata:
                    lines.append(f"Target Audience: {metadata['target_audience']}")
            
            lines.append("=" * 60)
            lines.append("")
        
        if include_statistics:
            lines.append("Statistics:")
            lines.append(f"  Word Count: {project_data.get('word_count', 0)}")
            lines.append(f"  Target Words: {project_data.get('target_words', 0)}")
            
            if project_data.get('word_count', 0) > 0 and project_data.get('target_words', 0) > 0:
                progress = (project_data['word_count'] / project_data['target_words']) * 100
                lines.append(f"  Progress: {progress:.1f}%")
            
            lines.append("")
            lines.append("-" * 60)
            lines.append("")
        
        # Add main content
        content = project_data.get('content', '')
        lines.append(content)
        
        return '\n'.join(lines)
    
    def _export_txt(self, project_data: Dict, output_path: Path, 
                   include_metadata: bool, include_statistics: bool) -> Dict:
        """Export to plain text format."""
        try:
            content = self._prepare_content(project_data, include_metadata, include_statistics)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'success': True,
                'format': 'txt',
                'path': str(output_path),
                'size': output_path.stat().st_size
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _export_markdown(self, project_data: Dict, output_path: Path,
                        include_metadata: bool, include_statistics: bool) -> Dict:
        """Export to Markdown format."""
        try:
            lines = []
            
            # Title
            lines.append(f"# {project_data.get('name', 'Untitled')}")
            lines.append("")
            
            if include_metadata:
                lines.append("## Metadata")
                lines.append("")
                lines.append(f"- **Created:** {project_data.get('created_date', 'Unknown')}")
                lines.append(f"- **Last Modified:** {project_data.get('last_modified', 'Unknown')}")
                
                metadata = project_data.get('metadata', {})
                if isinstance(metadata, str):
                    try:
                        metadata = json.loads(metadata)
                    except:
                        metadata = {}
                
                if metadata:
                    if 'genre' in metadata:
                        lines.append(f"- **Genre:** {metadata['genre']}")
                    if 'target_audience' in metadata:
                        lines.append(f"- **Target Audience:** {metadata['target_audience']}")
                
                lines.append("")
            
            if include_statistics:
                lines.append("## Statistics")
                lines.append("")
                lines.append(f"- **Word Count:** {project_data.get('word_count', 0)}")
                lines.append(f"- **Target Words:** {project_data.get('target_words', 0)}")
                
                if project_data.get('word_count', 0) > 0 and project_data.get('target_words', 0) > 0:
                    progress = (project_data['word_count'] / project_data['target_words']) * 100
                    lines.append(f"- **Progress:** {progress:.1f}%")
                
                lines.append("")
            
            lines.append("## Content")
            lines.append("")
            lines.append(project_data.get('content', ''))
            
            content = '\n'.join(lines)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'success': True,
                'format': 'md',
                'path': str(output_path),
                'size': output_path.stat().st_size
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _export_docx(self, project_data: Dict, output_path: Path,
                    include_metadata: bool, include_statistics: bool) -> Dict:
        """Export to DOCX format."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # Title
            title = doc.add_heading(project_data.get('name', 'Untitled'), 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if include_metadata:
                doc.add_heading('Metadata', 2)
                
                p = doc.add_paragraph()
                p.add_run('Created: ').bold = True
                p.add_run(project_data.get('created_date', 'Unknown'))
                
                p = doc.add_paragraph()
                p.add_run('Last Modified: ').bold = True
                p.add_run(project_data.get('last_modified', 'Unknown'))
                
                metadata = project_data.get('metadata', {})
                if isinstance(metadata, str):
                    try:
                        metadata = json.loads(metadata)
                    except:
                        metadata = {}
                
                if metadata and 'genre' in metadata:
                    p = doc.add_paragraph()
                    p.add_run('Genre: ').bold = True
                    p.add_run(metadata['genre'])
            
            if include_statistics:
                doc.add_heading('Statistics', 2)
                
                p = doc.add_paragraph()
                p.add_run('Word Count: ').bold = True
                p.add_run(str(project_data.get('word_count', 0)))
                
                p = doc.add_paragraph()
                p.add_run('Target Words: ').bold = True
                p.add_run(str(project_data.get('target_words', 0)))
            
            # Page break before content
            doc.add_page_break()
            
            # Content
            content = project_data.get('content', '')
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            doc.save(str(output_path))
            
            return {
                'success': True,
                'format': 'docx',
                'path': str(output_path),
                'size': output_path.stat().st_size
            }
        except ImportError:
            return {'success': False, 'error': 'python-docx not installed. Run: pip install python-docx'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _export_pdf(self, project_data: Dict, output_path: Path,
                   include_metadata: bool, include_statistics: bool) -> Dict:
        """Export to PDF format."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            
            doc = SimpleDocTemplate(str(output_path), pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor='black',
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor='black',
                spaceAfter=12,
                spaceBefore=12
            )
            
            # Title
            story.append(Paragraph(project_data.get('name', 'Untitled'), title_style))
            story.append(Spacer(1, 0.2*inch))
            
            if include_metadata:
                story.append(Paragraph('Metadata', heading_style))
                story.append(Paragraph(f"<b>Created:</b> {project_data.get('created_date', 'Unknown')}", styles['Normal']))
                story.append(Paragraph(f"<b>Last Modified:</b> {project_data.get('last_modified', 'Unknown')}", styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            if include_statistics:
                story.append(Paragraph('Statistics', heading_style))
                story.append(Paragraph(f"<b>Word Count:</b> {project_data.get('word_count', 0)}", styles['Normal']))
                story.append(Paragraph(f"<b>Target Words:</b> {project_data.get('target_words', 0)}", styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # Page break before content
            story.append(PageBreak())
            
            # Content
            content = project_data.get('content', '')
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    # Escape HTML characters
                    safe_text = paragraph.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe_text, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            
            return {
                'success': True,
                'format': 'pdf',
                'path': str(output_path),
                'size': output_path.stat().st_size
            }
        except ImportError:
            return {'success': False, 'error': 'reportlab not installed. Run: pip install reportlab'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _export_epub(self, project_data: Dict, output_path: Path,
                    include_metadata: bool, include_statistics: bool) -> Dict:
        """Export to EPUB format."""
        try:
            from ebooklib import epub
            
            book = epub.EpubBook()
            
            # Metadata
            book.set_identifier(f"aawt_{project_data.get('name', 'untitled')}")
            book.set_title(project_data.get('name', 'Untitled'))
            book.set_language('en')
            
            metadata = project_data.get('metadata', {})
            if isinstance(metadata, str):
                try:
                    metadata = json.loads(metadata)
                except:
                    metadata = {}
            
            if metadata and 'author' in metadata:
                book.add_author(metadata['author'])
            else:
                book.add_author('Unknown Author')
            
            # Create chapters
            chapters = []
            
            # Info chapter if metadata/stats included
            if include_metadata or include_statistics:
                info_content = ['<h1>Project Information</h1>']
                
                if include_metadata:
                    info_content.append('<h2>Metadata</h2>')
                    info_content.append(f"<p><strong>Created:</strong> {project_data.get('created_date', 'Unknown')}</p>")
                    info_content.append(f"<p><strong>Last Modified:</strong> {project_data.get('last_modified', 'Unknown')}</p>")
                
                if include_statistics:
                    info_content.append('<h2>Statistics</h2>')
                    info_content.append(f"<p><strong>Word Count:</strong> {project_data.get('word_count', 0)}</p>")
                    info_content.append(f"<p><strong>Target Words:</strong> {project_data.get('target_words', 0)}</p>")
                
                c1 = epub.EpubHtml(title='Information', file_name='info.xhtml', lang='en')
                c1.content = '\n'.join(info_content)
                book.add_item(c1)
                chapters.append(c1)
            
            # Main content chapter
            content = project_data.get('content', '')
            content_html = ['<h1>Content</h1>']
            
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    content_html.append(f'<p>{paragraph.strip()}</p>')
            
            c2 = epub.EpubHtml(title='Content', file_name='content.xhtml', lang='en')
            c2.content = '\n'.join(content_html)
            book.add_item(c2)
            chapters.append(c2)
            
            # Table of contents
            book.toc = tuple(chapters)
            
            # Navigation files
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Spine
            book.spine = ['nav'] + chapters
            
            # Write file
            epub.write_epub(str(output_path), book, {})
            
            return {
                'success': True,
                'format': 'epub',
                'path': str(output_path),
                'size': output_path.stat().st_size
            }
        except ImportError:
            return {'success': False, 'error': 'ebooklib not installed. Run: pip install ebooklib'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _export_json(self, project_data: Dict, output_path: Path,
                    include_metadata: bool, include_statistics: bool) -> Dict:
        """Export to JSON format."""
        try:
            export_data = {
                'export_date': datetime.now().isoformat(),
                'export_version': '1.0',
                'project': project_data
            }
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            return {
                'success': True,
                'format': 'json',
                'path': str(output_path),
                'size': output_path.stat().st_size
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def get_export_formats(self) -> list:
        """Get list of available export formats."""
        formats = []
        format_names = {
            'txt': 'Plain Text (.txt)',
            'md': 'Markdown (.md)',
            'docx': 'Microsoft Word (.docx)',
            'pdf': 'PDF Document (.pdf)',
            'epub': 'EPUB E-book (.epub)',
            'json': 'JSON Data (.json)'
        }
        
        for format_key, format_name in format_names.items():
            if self.settings.get(f'export.formats_enabled.{format_key}', True):
                formats.append({
                    'key': format_key,
                    'name': format_name,
                    'enabled': True
                })
        
        return formats

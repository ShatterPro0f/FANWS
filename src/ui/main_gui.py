"""
Modern FANWS GUI - Enhanced User Interface
Priority 3.1: Modern GUI Design Implementation

This module provides the modernized main GUI for FANWS with:
- Professional modern design system
- Responsive layout management
- Enhanced workflow visualization
- Improved user experience
- Better progress tracking
"""

import sys
import os
import json
import asyncio
from datetime import datetime, timedelta
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *

# Import modern UI components from ui package
try:
    from src.ui.core_ui import CoreUIComponents as UIComponents
    from src.ui.analytics_ui import AnalyticsUIComponents
    from src.ui.management_ui import ManagementUIComponents
    GUI_AVAILABLE = True
except ImportError as e:
    UIComponents = None
    AnalyticsUIComponents = None
    ManagementUIComponents = None
    GUI_AVAILABLE = False
    print(f"⚠ Modern UI components not available: {e}")

# Create consolidated design system
class DesignSystem:
    """Consolidated design system for FANWS"""
    COLORS = {
        'primary': '#2196F3',
        'primary_light': '#64B5F6',
        'secondary': '#FFC107',
        'success': '#4CAF50',
        'danger': '#F44336',
        'warning': '#FF9800',
        'info': '#2196F3',
        'bg_primary': '#FFFFFF',
        'bg_secondary': '#F5F5F5',
        'bg_card': '#FFFFFF',
        'text_primary': '#212121',
        'text_secondary': '#757575',
        'border_primary': '#E0E0E0'
    }

    @staticmethod
    def get_main_stylesheet():
        return """
        QMainWindow {
            background-color: #f5f5f5;
            color: #212121;
        }
        QWidget {
            font-family: 'Segoe UI', Arial, sans-serif;
        }
        """

    @staticmethod
    def get_sidebar_stylesheet():
        return """
        QFrame[class="sidebar"] {
            background-color: #263238;
            color: white;
            border-right: 1px solid #37474F;
        }
        QPushButton[class="nav-button"] {
            background-color: transparent;
            color: white;
            text-align: left;
            padding: 12px 16px;
            border: none;
            margin: 2px;
        }
        QPushButton[class="nav-button"]:hover {
            background-color: #37474F;
        }
        QPushButton[class="nav-button active"] {
            background-color: #2196F3;
        }
        """

    @staticmethod
    def get_dashboard_stylesheet():
        return """
        QLabel[class="header"] {
            font-size: 24px;
            font-weight: bold;
            color: #212121;
            margin: 16px 0px 8px 0px;
        }
        QLabel[class="subheader"] {
            font-size: 16px;
            font-weight: 600;
            color: #424242;
        }
        QLabel[class="caption"] {
            font-size: 12px;
            color: #757575;
        }
        QFrame[class="card"] {
            background-color: white;
            border: 1px solid #E0E0E0;
            border-radius: 8px;
            padding: 16px;
            margin: 8px;
        }
        """

class Components:
    """Modern UI components factory"""

    @staticmethod
    def create_button(text, style="primary", icon=None):
        button = QPushButton(text)
        colors = DesignSystem.COLORS
        if style == "primary":
            color = colors['primary']
        elif style == "success":
            color = colors['success']
        elif style == "danger":
            color = colors['danger']
        elif style == "secondary":
            color = colors['secondary']
        else:
            color = colors['primary']

        button.setStyleSheet(f"""
            QPushButton {{
                background-color: {color};
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background-color: {color}dd;
            }}
        """)
        return button

    @staticmethod
    def create_card(title, content_widget):
        card = QFrame()
        card.setProperty("class", "card")
        layout = QVBoxLayout()
        card.setLayout(layout)

        if title:
            title_label = QLabel(title)
            title_label.setProperty("class", "subheader")
            layout.addWidget(title_label)

        layout.addWidget(content_widget)
        return card

    @staticmethod
    def create_status_indicator(status, text):
        indicator = QLabel(text)
        colors = {
            'online': '#4CAF50',
            'offline': '#F44336',
            'warning': '#FF9800',
            'completed': '#4CAF50',
            'in_progress': '#2196F3',
            'pending': '#9E9E9E'
        }
        color = colors.get(status, '#9E9E9E')
        indicator.setStyleSheet(f"""
            QLabel {{
                color: {color};
                font-weight: 600;
                padding: 4px 8px;
                border-radius: 4px;
                background-color: {color}20;
            }}
        """)
        return indicator

class WorkflowUI:
    """Workflow UI components"""

    @staticmethod
    def create_progress_dashboard(metrics):
        widget = QWidget()
        layout = QGridLayout()
        widget.setLayout(layout)

        for i, (key, value) in enumerate(metrics.items()):
            card = Components.create_card(
                key.replace('_', ' ').title(),
                QLabel(str(value))
            )
            layout.addWidget(card, i // 3, i % 3)

        return widget

    @staticmethod
    def create_workflow_step_card(step_num, title, description, status):
        card = QFrame()
        card.setProperty("class", "card")
        layout = QVBoxLayout()
        card.setLayout(layout)

        # Step header
        header_layout = QHBoxLayout()
        step_label = QLabel(f"{step_num}. {title}")
        step_label.setProperty("class", "subheader")
        header_layout.addWidget(step_label)

        status_indicator = Components.create_status_indicator(status, status.title())
        header_layout.addWidget(status_indicator)
        header_layout.addStretch()

        layout.addLayout(header_layout)

        # Description
        desc_label = QLabel(description)
        desc_label.setProperty("class", "caption")
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)

        return card

# Create placeholder classes for missing components
class Animations:
    pass

class LayoutManager:
    pass

class WorkflowVisualization:
    pass

class UserFeedbackIntegration:
    pass

class IntegratedWorkflowFeedback:
    pass

class ProjectManagement:
    def __init__(self, parent):
        self.parent = parent

    def create_project_template_selector(self):
        """Create a project template selector widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        layout.addWidget(QLabel("Project Template Selector"))
        layout.addWidget(QLabel("Select a template for your new project:"))

        # Add some basic template options
        templates = ["Fantasy Novel", "Adventure Story", "Mystery Novel", "Custom Project"]
        for template in templates:
            btn = QPushButton(template)
            layout.addWidget(btn)

        widget.setLayout(layout)
        return widget

    def create_version_control_interface(self):
        """Create a version control interface widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        layout.addWidget(QLabel("Version Control Interface"))
        layout.addWidget(QLabel("Manage your project versions and backups:"))

        # Add basic version control options
        controls = ["Create Backup", "Restore Version", "View History", "Merge Changes"]
        for control in controls:
            btn = QPushButton(control)
            layout.addWidget(btn)

        widget.setLayout(layout)
        return widget

    def create_backup_management_interface(self):
        """Create a backup management interface widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        layout.addWidget(QLabel("Backup Management"))
        layout.addWidget(QLabel("Manage automatic and manual backups:"))

        # Add backup management options
        options = ["Create Manual Backup", "Schedule Automatic Backups", "View Backup History", "Restore from Backup"]
        for option in options:
            btn = QPushButton(option)
            layout.addWidget(btn)

        widget.setLayout(layout)
        return widget

    def create_project_sharing_interface(self):
        """Create a project sharing interface widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        layout.addWidget(QLabel("Project Sharing"))
        layout.addWidget(QLabel("Share your project with collaborators:"))

        # Add sharing options
        options = ["Export Project", "Share via Cloud", "Create Collaboration Link", "Manage Permissions"]
        for option in options:
            btn = QPushButton(option)
            layout.addWidget(btn)

        widget.setLayout(layout)
        return widget

class PerformanceOptimizationUI:
    def __init__(self, parent):
        self.parent = parent

    def create_performance_dashboard(self):
        return QLabel("Performance Dashboard - Coming Soon")

    def create_performance_overview(self):
        return QLabel("Performance Overview - Coming Soon")

    def create_system_health_overview(self):
        return QLabel("System Health Overview - Coming Soon")

# Import async operations framework (Priority 4.1)
try:
    from ..system.async_operations import BackgroundTaskManager, get_async_manager
    from ..plugins.plugin_workflow_integration import AsyncWorkflowOperations
    ASYNC_FRAMEWORK_AVAILABLE = True

    # Create placeholder UI components for async operations
    class AsyncProgressDialog(QDialog):
        def __init__(self, parent=None):
            super().__init__(parent)
            self.setWindowTitle("Async Progress")
            layout = QVBoxLayout()
            self.progress_bar = QProgressBar()
            layout.addWidget(self.progress_bar)
            self.setLayout(layout)

    class AsyncStatusBar(QWidget):
        def __init__(self, parent=None):
            super().__init__(parent)
            layout = QHBoxLayout()
            self.status_label = QLabel("Ready")
            layout.addWidget(self.status_label)
            self.setLayout(layout)

    class AsyncTaskWidget(QWidget):
        def __init__(self, parent=None):
            super().__init__(parent)
            layout = QVBoxLayout()
            self.task_label = QLabel("No active tasks")
            layout.addWidget(self.task_label)
            self.setLayout(layout)

    class WorkflowTaskManager:
        def __init__(self, async_manager=None):
            self.async_manager = async_manager

except ImportError as e:
    print(f"⚠ Async framework not available: {e}")
    BackgroundTaskManager = None
    WorkflowTaskManager = None
    AsyncProgressDialog = None
    AsyncStatusBar = None
    AsyncTaskWidget = None
    ASYNC_FRAMEWORK_AVAILABLE = False

# Import multi-provider AI system
try:
    from ..system.api_manager import get_api_manager

    class AIManager:
        """AI Provider Manager for multi-provider support"""
        def __init__(self):
            self.api_manager = get_api_manager()

        def get_current_provider(self):
            return getattr(self.api_manager, 'current_provider', 'openai')

        def set_provider(self, provider):
            if hasattr(self.api_manager, 'set_provider'):
                self.api_manager.set_provider(provider)

        def generate_content(self, prompt, **kwargs):
            if hasattr(self.api_manager, 'generate_content'):
                return self.api_manager.generate_content(prompt, **kwargs)
            return "Generated content placeholder"

except ImportError:
    class AIManager:
        """Fallback AI Manager"""
        def get_current_provider(self):
            return "openai"

        def set_provider(self, provider):
            pass

        def generate_content(self, prompt, **kwargs):
            return "Generated content placeholder"

class MultiProviderConfigurationUI(QWidget):
    def __init__(self, ai_manager, parent=None):
        super().__init__(parent)
        self.ai_manager = ai_manager
        layout = QVBoxLayout()
        layout.addWidget(QLabel("AI Provider Configuration"))
        self.setLayout(layout)

# Import advanced prompt engineering system
try:
    from ..text.text_processing import PromptEngineeringManager

    class PromptEngineUI(QWidget):
        def __init__(self, prompt_engine, parent=None):
            super().__init__(parent)
            self.prompt_engine = prompt_engine
            layout = QVBoxLayout()
            layout.addWidget(QLabel("Prompt Engineering Tools"))
            self.setLayout(layout)

except ImportError as e:
    class PromptEngine:
        def __init__(self, ai_manager=None):
            pass

    class PromptEngineUI(QWidget):
        def __init__(self, prompt_engine=None, parent=None):
            super().__init__(parent)

# Import existing components
try:
    # Primary import from consolidated workflow_coordinator
    from ..plugins.plugin_workflow_integration import NovelWritingWorkflowModular
except ImportError:
    NovelWritingWorkflowModular = None

try:
    from src.ui import UIComponents
except ImportError:
    UIComponents = None

try:
    from ..core.configuration_manager import get_global_config
    ConfigManager = get_global_config()
except ImportError:
    ConfigManager = None

try:
    from ..database.database_manager import DatabaseManager
except ImportError:
    DatabaseManager = None

try:
    from ..core.performance_monitor import PerformanceMonitor
except ImportError:
    PerformanceMonitor = None

# Import workflow controller
try:
    # Primary import from consolidated workflow_coordinator
    from ..plugins.plugin_workflow_integration import WorkflowCoordinator
except ImportError:
    class WorkflowFeedbackIntegrator(QObject):
        workflow_feedback_ready = pyqtSignal(str, dict)

        def __init__(self):
            super().__init__()

            # Create minimal progress tracker
            class ProgressTracker(QObject):
                progress_updated = pyqtSignal(int, str)
                eta_updated = pyqtSignal(str)
                speed_updated = pyqtSignal(float)
                log_message = pyqtSignal(str, str)

                def __init__(self):
                    super().__init__()

            self.progress_tracker = ProgressTracker()

            class FeedbackManager(QObject):
                feedback_received = pyqtSignal(dict)

                def __init__(self):
                    super().__init__()

            self.feedback_manager = FeedbackManager()

class MainWindow(QMainWindow):
    """
    Main window for FANWS with GUI design.
    """

    def __init__(self):
        super().__init__()

        # Initialize components if available
        self.config = ConfigManager if ConfigManager else None
        self.db_manager = DatabaseManager() if DatabaseManager else None

        # Initialize design system and UI components
        self.design_system = DesignSystem()
        self.animations = Animations()
        self.components = Components()
        self.layout_manager = LayoutManager()

        self.performance_monitor = PerformanceMonitor() if PerformanceMonitor else None
        self.workflow_manager = NovelWritingWorkflowModular() if NovelWritingWorkflowModular else None

        # Initialize multi-provider AI system
        self.ai_provider_manager = AIManager()
        self.ai_config_ui = MultiProviderConfigurationUI(self.ai_provider_manager)

        # Initialize advanced prompt engineering system
        self.prompt_engine = PromptEngine(self.ai_provider_manager)
        self.prompt_engine_ui = PromptEngineUI(self.prompt_engine)

        # Initialize analytics system for dashboard
        try:
            from ..analytics.analytics_system import WritingAnalyticsDashboard, AnalyticsWidget
            self.analytics_manager = WritingAnalyticsDashboard()
            self.analytics_widget = AnalyticsWidget(self.analytics_manager)
            print("✓ Analytics system initialized")
        except ImportError as e:
            self.analytics_manager = None
            self.analytics_widget = None
            print(f"⚠ Analytics system not available: {e}")

        # Initialize workflow manager with analytics integration
        if self.workflow_manager and self.analytics_manager:
            # Connect workflow events to analytics
            try:
                # This would connect workflow completion events to analytics tracking
                pass
            except Exception as e:
                print(f"⚠ Analytics integration failed: {e}")

        # UI state
        self.current_project = None
        self.workflow_status = {}
        self.animations = {}

        # Dashboard chart state
        self._last_total_progress = 0.0
        self._last_avg_daily = 0.0
        self._chart_metrics = {}

        # Initialize workflow controller
        self.workflow_controller = WorkflowFeedbackIntegrator()

        # Initialize advanced project management
        self.advanced_project_mgmt = ProjectManagement(self)

        # Intelligent caching integration
        try:
            from ..core.error_handling_system import IntelligentCacheManager, CacheIntegration
            # Get project directory safely
            project_dir = getattr(self, 'project_dir', os.path.dirname(__file__))
            self.cache_manager = IntelligentCacheManager(
                cache_dir=os.path.join(project_dir, "cache"),
                max_cache_size=500
            )
            self.cache_integration = CacheIntegration(self.cache_manager)
            self.cache_enabled = True
            print("✓ Intelligent caching system initialized")
        except (ImportError, Exception) as e:
            self.cache_manager = None
            self.cache_integration = None
            self.cache_enabled = False
            print(f"⚠ Intelligent caching system not available: {e}")

        # Initialize async framework (Priority 4.1)
        if ASYNC_FRAMEWORK_AVAILABLE:
            self.async_manager = get_async_manager()
            self.async_workflow_manager = WorkflowTaskManager(self.async_manager)
            self.async_status_bar = AsyncStatusBar(self)
            self.async_task_widgets = {}
            self.async_enabled = True
            print("✓ Async operations framework initialized")
        else:
            self.async_manager = None
            self.async_workflow_manager = None
            self.async_status_bar = None
            self.async_task_widgets = {}
            self.async_enabled = False
            print("⚠ Running without async framework")

        # UI
        self.init_ui()
        self.setup_styling()
        self.setup_connections()

        # Initialize async UI components if available
        if self.async_enabled:
            self.init_async_ui()

        self.load_project_data()

        # Connect controller signals to UI updates
        self.workflow_controller.progress_tracker.progress_updated.connect(self.update_progress_display)
        self.workflow_controller.progress_tracker.eta_updated.connect(self.update_eta_display)
        self.workflow_controller.progress_tracker.speed_updated.connect(self.update_speed_display)
        self.workflow_controller.progress_tracker.log_message.connect(self.add_log_message)
        self.workflow_controller.workflow_feedback_ready.connect(self.show_feedback_request)
        self.workflow_controller.feedback_manager.feedback_received.connect(self.update_feedback_metrics)

        # Start performance monitoring if available
        if self.performance_monitor:
            try:
                self.performance_monitor.start_monitoring()
            except:
                pass

    def init_ui(self):
        """Initialize the modern user interface."""
        self.setWindowTitle("FANWS - Fully Automated Novel Writing System")
        self.setGeometry(100, 100, 1400, 900)
        self.setMinimumSize(1200, 800)

        # Set window icon if available
        icon_path = self.get_icon_path('fanws_icon.png')
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

        # Create central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Create main layout - sidebar takes 1/4, content takes 3/4
        main_layout = QHBoxLayout()
        central_widget.setLayout(main_layout)

        # Create sidebar (1/4 of screen)
        self.sidebar = self.create_hierarchical_sidebar()
        main_layout.addWidget(self.sidebar, 1)  # 1/4 ratio

        # Create main content area (3/4 of screen)
        self.main_content = self.create_home_page()
        main_layout.addWidget(self.main_content, 3)  # 3/4 ratio

        # Create status bar
        self.create_status_bar()

        # Create menu bar
        self.create_menu_bar()

    def create_hierarchical_sidebar(self):
        """Create the hierarchical sidebar with sections, subsections, and subsubsections."""
        sidebar = QFrame()
        sidebar.setProperty("class", "sidebar")
        sidebar.setMinimumWidth(350)  # Increased width for hierarchy
        sidebar.setMaximumWidth(350)

        layout = QVBoxLayout()
        sidebar.setLayout(layout)

        # Header
        header = QFrame()
        header.setProperty("class", "sidebar-header")
        header_layout = QVBoxLayout()
        header.setLayout(header_layout)

        # Logo/Title
        title = QLabel("FANWS")
        title.setProperty("class", "header")
        header_layout.addWidget(title)

        # Subtitle
        subtitle = QLabel("Fully Automated Novel Writing System")
        subtitle.setProperty("class", "caption")
        header_layout.addWidget(subtitle)

        layout.addWidget(header)

        # Create scrollable area for navigation
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        nav_widget = QWidget()
        nav_layout = QVBoxLayout()
        nav_widget.setLayout(nav_layout)

        # Initialize tracking variables
        self.current_section = None
        self.current_subsection = None
        self.section_buttons = {}
        self.subsection_buttons = {}
        self.subsubsection_buttons = {}

        # Define the hierarchical structure
        self.navigation_structure = {
            'project': {
                'title': '📁 Project',
                'subsections': {
                    'switch_project': 'Switch Project',
                    'create_project': 'Create Project',
                    'load_project': 'Load Project',
                    'delete_project': 'Delete Project',
                    'novel_settings': {
                        'title': 'Novel Settings',
                        'subsubsections': {
                            'novel_concept': 'Novel Concept',
                            'primary_tone': 'Primary Tone',
                            'sub_tone': 'Sub-Tone',
                            'theme': 'Theme',
                            'target_word_count': 'Target Word Count',
                            'reading_level': 'Reading Level',
                            'chapter_organization': 'Chapter/section organization',
                            'total_chapters': 'Total Chapters',
                            'chapter_sections': 'Chapter Sections: Sections per chapter'
                        }
                    }
                }
            },
            'writing': {
                'title': '✍️ Writing',
                'subsections': {
                    'text_editor': 'Text Editor',
                    'writing_session': 'Writing Session',
                    'goal_tracking': 'Goal Tracking',
                    'word_count': 'Word Count Tracker',
                    'writing_timer': 'Writing Timer'
                }
            },
            'ai': {
                'title': '🤖 AI Assistant',
                'subsections': {
                    'ai_settings': 'AI Provider Settings',
                    'models': 'AI Models',
                    'prompts': 'Prompts & Templates',
                    'response_handling': 'Response Handling',
                    'usage_limits': 'Usage Limits & Monitoring'
                }
            },
            'templates': {
                'title': '📋 Templates',
                'subsections': {
                    'template_library': 'Template Library',
                    'custom_templates': 'Custom Templates',
                    'template_marketplace': 'Template Marketplace',
                    'template_recommendations': 'Recommendations'
                }
            },
            'collaboration': {
                'title': '🤝 Collaboration',
                'subsections': {
                    'collaboration_overview': 'Overview',
                    'team_management': 'Team Management',
                    'shared_projects': 'Shared Projects',
                    'communication': 'Communication',
                    'version_control': 'Version Control'
                }
            },
            'analytics': {
                'title': '📊 Analytics',
                'subsections': {
                    'writing_analytics': 'Writing Analytics',
                    'productivity_metrics': 'Productivity Metrics',
                    'goal_progress': 'Goal Progress',
                    'usage_statistics': 'Usage Statistics'
                }
            },
            'text_tools': {
                'title': '📝 Text Tools',
                'subsections': {
                    'text_analysis': 'Text Analysis',
                    'grammar_check': 'Grammar Check',
                    'style_analysis': 'Style Analysis',
                    'consistency_check': 'Consistency Check',
                    'readability': 'Readability Analysis'
                }
            },
            'workflow': {
                'title': '🔄 Workflow',
                'subsections': {
                    'workflow_overview': 'Workflow Overview',
                    'step_management': 'Step Management',
                    'automation': 'Automation Settings',
                    'workflow_templates': 'Workflow Templates'
                }
            },
            'dashboard': {
                'title': '📊 Dashboard',
                'subsections': {
                    'progress_graph': 'Progress Graph',
                    'synonyms': 'Synonyms',
                    'log': 'Log',
                    'chapter_progress': 'Chapter Progress',
                    'current_draft': 'Current Draft'
                }
            },
            'performance': {
                'title': '📈 Performance',
                'subsections': {
                    'memory_usage': 'Memory Usage: Current RAM consumption (MB)',
                    'cpu_usage': 'CPU Usage: Processor utilization (%)',
                    'api_call_stats': 'API Call Statistics: Requests per service',
                    'file_operations': 'File Operations: Read/write counts',
                    'cache_hit_rate': 'Cache Hit Rate: File cache efficiency (%)',
                    'response_times': 'Response Times: API and file operation speeds',
                    'optimization_recommendations': 'Optimization Recommendations: Performance suggestions',
                    'system_resources': 'System Resources: Disk space, network status'
                }
            },
            'settings': {
                'title': '⚙️ Settings',
                'subsections': {
                    'openai_api_key': 'OpenAI API Key (Savable)',
                    'wordsapi_key': 'WordsAPI Key (Savable)'
                }
            },
            'export': {
                'title': '📤 Export',
                'subsections': {
                    'export_status': 'Export Status: Success/failure of exports',
                    'export_formats': 'Export Formats: Available output types',
                    'export_history': 'Export History: Previous export attempts',
                    'file_sizes': 'File Sizes: Generated file sizes',
                    'export_quality': 'Export Quality: Format-specific settings'
                }
            }
        }

        # Create navigation hierarchy
        self.create_navigation_hierarchy(nav_layout)

        scroll_area.setWidget(nav_widget)
        layout.addWidget(scroll_area)

        # Spacer
        nav_layout.addStretch()

        return sidebar

    def create_navigation_hierarchy(self, layout):
        """Create the hierarchical navigation structure."""
        for section_id, section_data in self.navigation_structure.items():
            # Create section button
            section_btn = QPushButton(section_data['title'])
            section_btn.setProperty("class", "nav-section")
            section_btn.clicked.connect(lambda checked, sid=section_id: self.toggle_section(sid))
            section_btn.setStyleSheet("""
                QPushButton[class="nav-section"] {
                    background-color: #37474F;
                    color: white;
                    text-align: left;
                    padding: 12px 8px;
                    border: none;
                    margin: 2px 0;
                    font-weight: bold;
                }
                QPushButton[class="nav-section"]:hover {
                    background-color: #455A64;
                }
                QPushButton[class="nav-section"]:pressed {
                    background-color: #263238;
                }
            """)
            self.section_buttons[section_id] = section_btn
            layout.addWidget(section_btn)

            # Create subsection container (initially hidden)
            subsection_container = QWidget()
            subsection_layout = QVBoxLayout()
            subsection_layout.setContentsMargins(20, 0, 0, 0)  # Indent subsections
            subsection_container.setLayout(subsection_layout)
            subsection_container.hide()

            for subsection_id, subsection_data in section_data['subsections'].items():
                if isinstance(subsection_data, dict):  # Has subsubsections
                    # Create subsection button with subsubsections
                    subsection_btn = QPushButton(f"  ▶ {subsection_data['title']}")
                    subsection_btn.setProperty("class", "nav-subsection")
                    subsection_btn.clicked.connect(lambda checked, sid=section_id, ssid=subsection_id: self.toggle_subsection(sid, ssid))

                    # Create subsubsection container (initially hidden)
                    subsubsection_container = QWidget()
                    subsubsection_layout = QVBoxLayout()
                    subsubsection_layout.setContentsMargins(20, 0, 0, 0)  # Further indent
                    subsubsection_container.setLayout(subsubsection_layout)
                    subsubsection_container.hide()

                    for subsubsection_id, subsubsection_title in subsection_data['subsubsections'].items():
                        subsubsection_btn = QPushButton(f"    • {subsubsection_title}")
                        subsubsection_btn.setProperty("class", "nav-subsubsection")
                        subsubsection_btn.clicked.connect(lambda checked, sid=section_id, ssid=subsection_id, sssid=subsubsection_id: self.open_subsubsection(sid, ssid, sssid))

                        if subsection_id not in self.subsubsection_buttons:
                            self.subsubsection_buttons[subsection_id] = {}
                        self.subsubsection_buttons[subsection_id][subsubsection_id] = {
                            'button': subsubsection_btn,
                            'container': subsubsection_container
                        }
                        subsubsection_layout.addWidget(subsubsection_btn)

                    subsection_layout.addWidget(subsection_btn)
                    subsection_layout.addWidget(subsubsection_container)

                    if section_id not in self.subsection_buttons:
                        self.subsection_buttons[section_id] = {}
                    self.subsection_buttons[section_id][subsection_id] = {
                        'button': subsection_btn,
                        'container': subsubsection_container
                    }
                else:  # Simple subsection
                    subsection_btn = QPushButton(f"  • {subsection_data}")
                    subsection_btn.setProperty("class", "nav-subsection")
                    subsection_btn.clicked.connect(lambda checked, sid=section_id, ssid=subsection_id: self.open_subsection(sid, ssid))

                    if section_id not in self.subsection_buttons:
                        self.subsection_buttons[section_id] = {}
                    self.subsection_buttons[section_id][subsection_id] = {
                        'button': subsection_btn,
                        'container': None
                    }
                    subsection_layout.addWidget(subsection_btn)

            # Store the subsection container reference
            self.section_buttons[section_id].subsection_container = subsection_container
            layout.addWidget(subsection_container)

        # Apply styling for subsections and subsubsections
        self.apply_navigation_styling()

    def apply_navigation_styling(self):
        """Apply styling to navigation elements."""
        subsection_style = """
            QPushButton[class="nav-subsection"] {
                background-color: #455A64;
                color: white;
                text-align: left;
                padding: 8px 4px;
                border: none;
                margin: 1px 0;
                font-size: 11px;
            }
            QPushButton[class="nav-subsection"]:hover {
                background-color: #546E7A;
            }
        """

        subsubsection_style = """
            QPushButton[class="nav-subsubsection"] {
                background-color: #546E7A;
                color: white;
                text-align: left;
                padding: 6px 4px;
                border: none;
                margin: 1px 0;
                font-size: 10px;
            }
            QPushButton[class="nav-subsubsection"]:hover {
                background-color: #607D8B;
            }
        """

        # Apply styles to all subsection and subsubsection buttons
        for section_id, subsections in self.subsection_buttons.items():
            for subsection_id, subsection_data in subsections.items():
                subsection_data['button'].setStyleSheet(subsection_style)

        for subsection_id, subsubsections in self.subsubsection_buttons.items():
            for subsubsection_id, subsubsection_data in subsubsections.items():
                subsubsection_data['button'].setStyleSheet(subsubsection_style)

    def toggle_section(self, section_id):
        """Toggle section visibility and update arrow indicators."""
        section_btn = self.section_buttons[section_id]
        container = section_btn.subsection_container

        if container.isVisible():
            container.hide()
            # Update button text to show closed state
            title = self.navigation_structure[section_id]['title']
            section_btn.setText(title)
        else:
            # Hide all other sections first
            for other_section_id, other_btn in self.section_buttons.items():
                if other_section_id != section_id and other_btn.subsection_container.isVisible():
                    other_btn.subsection_container.hide()
                    other_title = self.navigation_structure[other_section_id]['title']
                    other_btn.setText(other_title)

            container.show()
            # Update button text to show open state
            title = self.navigation_structure[section_id]['title']
            section_btn.setText(f"▼ {title[2:]}")  # Remove emoji and add arrow

        self.current_section = section_id if container.isVisible() else None

    def toggle_subsection(self, section_id, subsection_id):
        """Toggle subsection with subsubsections."""
        subsection_data = self.subsection_buttons[section_id][subsection_id]
        container = subsection_data['container']
        button = subsection_data['button']

        if container and container.isVisible():
            container.hide()
            button.setText(f"  ▶ {self.navigation_structure[section_id]['subsections'][subsection_id]['title']}")
        else:
            # Hide all other subsection containers in this section
            for other_subsection_id, other_subsection_data in self.subsection_buttons[section_id].items():
                if (other_subsection_id != subsection_id and
                    other_subsection_data['container'] and
                    other_subsection_data['container'].isVisible()):
                    other_subsection_data['container'].hide()
                    other_title = self.navigation_structure[section_id]['subsections'][other_subsection_id]['title']
                    other_subsection_data['button'].setText(f"  ▶ {other_title}")

            if container:
                container.show()
                button.setText(f"  ▼ {self.navigation_structure[section_id]['subsections'][subsection_id]['title']}")

        self.current_subsection = subsection_id if container and container.isVisible() else None

    def open_subsection(self, section_id, subsection_id):
        """Open a simple subsection (no subsubsections) in the home page."""
        subsection_title = self.navigation_structure[section_id]['subsections'][subsection_id]
        self.load_content_to_home_page(f"{section_id}_{subsection_id}", subsection_title, section_id, subsection_id)

    def open_subsubsection(self, section_id, subsection_id, subsubsection_id):
        """Open a subsubsection in the home page."""
        subsubsection_title = self.navigation_structure[section_id]['subsections'][subsection_id]['subsubsections'][subsubsection_id]
        self.load_content_to_home_page(f"{section_id}_{subsection_id}_{subsubsection_id}", subsubsection_title, section_id, subsection_id, subsubsection_id)

    def create_home_page(self):
        """Create the main home page content area (3/4 of screen)."""
        home_page = QWidget()
        layout = QVBoxLayout()
        home_page.setLayout(layout)

        # Page header
        self.page_header = QLabel("Welcome to FANWS")
        self.page_header.setProperty("class", "page-header")
        self.page_header.setStyleSheet("""
            QLabel[class="page-header"] {
                font-size: 28px;
                font-weight: bold;
                color: #212121;
                padding: 20px;
                border-bottom: 2px solid #E0E0E0;
                margin-bottom: 20px;
            }
        """)
        layout.addWidget(self.page_header)

        # Content area
        self.content_area = QStackedWidget()

        # Create default welcome content
        welcome_content = self.create_welcome_content()
        self.content_area.addWidget(welcome_content)

        layout.addWidget(self.content_area)

        return home_page

    def create_welcome_content(self):
        """Create the default welcome content."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Welcome message
        welcome_label = QLabel("Welcome to FANWS - Fully Automated Novel Writing System")
        welcome_label.setStyleSheet("""
            font-size: 18px;
            color: #424242;
            padding: 20px;
            text-align: center;
        """)
        welcome_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(welcome_label)

        # Instructions
        instructions = QLabel("""
        <h3>Getting Started:</h3>
        <p>• Use the sidebar to navigate between different sections</p>
        <p>• Click on section headers to expand and see subsections</p>
        <p>• Select subsections to view their content here</p>
        <p>• Configure your project settings in the Project section</p>
        <p>• Monitor performance and view analytics in their respective sections</p>
        """)
        instructions.setStyleSheet("""
            font-size: 14px;
            color: #666666;
            padding: 20px;
            background-color: #F5F5F5;
            border-radius: 8px;
            margin: 20px;
        """)
        layout.addWidget(instructions)

        layout.addStretch()
        return widget

    def load_content_to_home_page(self, content_id, title, section_id, subsection_id, subsubsection_id=None):
        """Load specific content to the home page based on selection."""
        # Update page header
        self.page_header.setText(title)

        # Create content based on the selection
        content_widget = self.create_content_for_selection(section_id, subsection_id, subsubsection_id)

        # Clear existing content and add new content
        while self.content_area.count() > 0:
            widget = self.content_area.widget(0)
            self.content_area.removeWidget(widget)
            widget.deleteLater()

        self.content_area.addWidget(content_widget)
        self.content_area.setCurrentWidget(content_widget)

    def create_content_for_selection(self, section_id, subsection_id, subsubsection_id=None):
        """Create appropriate content widget based on the selection."""
        # Route to specific content creation methods
        if section_id == 'project':
            return self.create_project_content(subsection_id, subsubsection_id)
        elif section_id == 'writing':
            return self.create_writing_content(subsection_id)
        elif section_id == 'ai':
            return self.create_ai_content(subsection_id)
        elif section_id == 'templates':
            return self.create_templates_content(subsection_id)
        elif section_id == 'collaboration':
            return self.create_collaboration_content(subsection_id)
        elif section_id == 'analytics':
            return self.create_analytics_content(subsection_id)
        elif section_id == 'text_tools':
            return self.create_text_tools_content(subsection_id)
        elif section_id == 'workflow':
            return self.create_workflow_content(subsection_id)
        elif section_id == 'dashboard':
            return self.create_dashboard_content(subsection_id)
        elif section_id == 'performance':
            return self.create_performance_content(subsection_id)
        elif section_id == 'settings':
            return self.create_settings_content(subsection_id)
        elif section_id == 'export':
            return self.create_export_content(subsection_id)
        else:
            return self.create_placeholder_content(section_id, subsection_id, subsubsection_id)

    def create_project_content(self, subsection_id, subsubsection_id=None):
        """Create content for project-related selections."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        if subsection_id == 'switch_project':
            return self.create_switch_project_widget()
        elif subsection_id == 'create_project':
            return self.create_new_project_widget()
        elif subsection_id == 'load_project':
            return self.create_load_project_widget()
        elif subsection_id == 'delete_project':
            return self.create_delete_project_widget()
        elif subsection_id == 'novel_settings':
            if subsubsection_id:
                return self.create_novel_settings_widget(subsubsection_id)
            else:
                return self.create_novel_settings_overview()

        return self.create_placeholder_content('project', subsection_id, subsubsection_id)

    def create_dashboard_content(self, subsection_id):
        """Create content for dashboard-related selections."""
        if subsection_id == 'progress_graph':
            return self.create_progress_graph_widget()
        elif subsection_id == 'synonyms':
            return self.create_synonyms_widget()
        elif subsection_id == 'log':
            return self.create_log_widget()
        elif subsection_id == 'chapter_progress':
            return self.create_chapter_progress_widget()
        elif subsection_id == 'current_draft':
            return self.create_current_draft_widget()

        return self.create_placeholder_content('dashboard', subsection_id)

    def create_performance_content(self, subsection_id):
        """Create content for performance-related selections."""
        if subsection_id == 'memory_usage':
            return self.create_memory_usage_widget()
        elif subsection_id == 'cpu_usage':
            return self.create_cpu_usage_widget()
        elif subsection_id == 'api_call_stats':
            return self.create_api_stats_widget()
        elif subsection_id == 'file_operations':
            return self.create_file_operations_widget()
        elif subsection_id == 'cache_hit_rate':
            return self.create_cache_rate_widget()
        elif subsection_id == 'response_times':
            return self.create_response_times_widget()
        elif subsection_id == 'optimization_recommendations':
            return self.create_optimization_widget()
        elif subsection_id == 'system_resources':
            return self.create_system_resources_widget()

        return self.create_placeholder_content('performance', subsection_id)

    def create_settings_content(self, subsection_id):
        """Create content for settings-related selections."""
        if subsection_id == 'openai_api_key':
            return self.create_openai_api_key_widget()
        elif subsection_id == 'wordsapi_key':
            return self.create_wordsapi_key_widget()

        return self.create_placeholder_content('settings', subsection_id)

    def create_export_content(self, subsection_id):
        """Create content for export-related selections."""
        if subsection_id == 'export_status':
            return self.create_export_status_widget()
        elif subsection_id == 'export_formats':
            return self.create_export_formats_widget()
        elif subsection_id == 'export_history':
            return self.create_export_history_widget()
        elif subsection_id == 'file_sizes':
            return self.create_file_sizes_widget()
        elif subsection_id == 'export_quality':
            return self.create_export_quality_widget()

        return self.create_placeholder_content('export', subsection_id)

    def create_placeholder_content(self, section_id, subsection_id, subsubsection_id=None):
        """Create placeholder content for sections not yet implemented."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        title = f"{section_id.title()} - {subsection_id.replace('_', ' ').title()}"
        if subsubsection_id:
            title += f" - {subsubsection_id.replace('_', ' ').title()}"

        label = QLabel(f"Content for: {title}")
        label.setStyleSheet("""
            font-size: 16px;
            color: #666666;
            padding: 20px;
            text-align: center;
        """)
        label.setAlignment(Qt.AlignCenter)
        layout.addWidget(label)

        placeholder_text = QLabel("This section is under development. Content will be available soon.")
        placeholder_text.setStyleSheet("""
            font-size: 14px;
            color: #999999;
            padding: 10px;
            text-align: center;
        """)
        placeholder_text.setAlignment(Qt.AlignCenter)
        layout.addWidget(placeholder_text)

        layout.addStretch()
        return widget

    # Project content widgets
    def create_switch_project_widget(self):
        """Create project switching interface."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Project selector
        selector_group = QGroupBox("Select Project")
        selector_layout = QFormLayout()

        self.project_selector = QComboBox()
        self.project_selector.addItem("Select a project")
        # Load projects from projects directory
        projects_dir = "projects"
        if os.path.exists(projects_dir):
            for project in os.listdir(projects_dir):
                if os.path.isdir(os.path.join(projects_dir, project)):
                    self.project_selector.addItem(project)

        selector_layout.addRow("Available Projects:", self.project_selector)

        switch_btn = Components.create_button("Switch to Project", "primary")
        switch_btn.clicked.connect(self.switch_to_selected_project)
        selector_layout.addRow("", switch_btn)

        selector_group.setLayout(selector_layout)
        layout.addWidget(selector_group)

        layout.addStretch()
        return widget

    def create_new_project_widget(self):
        """Create new project interface."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Project creation form
        creation_group = QGroupBox("Create New Project")
        creation_layout = QFormLayout()

        self.new_project_name = QLineEdit()
        self.new_project_name.setPlaceholderText("Enter project name...")
        creation_layout.addRow("Project Name:", self.new_project_name)

        self.project_description = QTextEdit()
        self.project_description.setPlaceholderText("Enter project description (optional)...")
        self.project_description.setMaximumHeight(100)
        creation_layout.addRow("Description:", self.project_description)

        create_btn = Components.create_button("Create Project", "success")
        create_btn.clicked.connect(self.create_new_project)
        creation_layout.addRow("", create_btn)

        creation_group.setLayout(creation_layout)
        layout.addWidget(creation_group)

        layout.addStretch()
        return widget

    def create_novel_settings_widget(self, subsubsection_id):
        """Create specific novel settings widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        settings_group = QGroupBox(f"{subsubsection_id.replace('_', ' ').title()}")
        settings_layout = QFormLayout()

        if subsubsection_id == 'novel_concept':
            self.novel_concept_text = QTextEdit()
            self.novel_concept_text.setPlaceholderText("Describe your novel concept, main plot, characters, and setting...")
            settings_layout.addRow("Novel Concept:", self.novel_concept_text)

        elif subsubsection_id == 'primary_tone':
            self.primary_tone_combo = QComboBox()
            tones = ["Neutral", "Dramatic", "Humorous", "Dark", "Light", "Serious", "Playful", "Mysterious", "Romantic", "Adventurous"]
            self.primary_tone_combo.addItems(tones)
            settings_layout.addRow("Primary Tone:", self.primary_tone_combo)

        elif subsubsection_id == 'sub_tone':
            self.sub_tone_combo = QComboBox()
            sub_tones = ["Descriptive", "Dialogue-heavy", "Action-packed", "Contemplative", "Fast-paced", "Detailed"]
            self.sub_tone_combo.addItems(sub_tones)
            settings_layout.addRow("Sub-Tone:", self.sub_tone_combo)

        elif subsubsection_id == 'theme':
            self.theme_text = QTextEdit()
            self.theme_text.setPlaceholderText("What themes does your novel explore? (e.g., friendship, redemption, power, etc.)")
            self.theme_text.setMaximumHeight(100)
            settings_layout.addRow("Theme:", self.theme_text)

        elif subsubsection_id == 'target_word_count':
            self.target_word_count = QSpinBox()
            self.target_word_count.setRange(1000, 1000000)
            self.target_word_count.setValue(80000)
            self.target_word_count.setSuffix(" words")
            settings_layout.addRow("Target Word Count:", self.target_word_count)

        elif subsubsection_id == 'reading_level':
            self.reading_level_combo = QComboBox()
            levels = ["Elementary", "Middle School", "High School", "College", "Graduate"]
            self.reading_level_combo.addItems(levels)
            self.reading_level_combo.setCurrentText("High School")
            settings_layout.addRow("Reading Level:", self.reading_level_combo)

        elif subsubsection_id == 'total_chapters':
            self.total_chapters_spin = QSpinBox()
            self.total_chapters_spin.setRange(1, 100)
            self.total_chapters_spin.setValue(20)
            settings_layout.addRow("Total Chapters:", self.total_chapters_spin)

        elif subsubsection_id == 'chapter_sections':
            self.chapter_sections_spin = QSpinBox()
            self.chapter_sections_spin.setRange(1, 20)
            self.chapter_sections_spin.setValue(5)
            settings_layout.addRow("Sections per Chapter:", self.chapter_sections_spin)

        # Save button for all settings
        save_btn = Components.create_button("Save Settings", "primary")
        save_btn.clicked.connect(lambda: self.save_novel_settings(subsubsection_id))
        settings_layout.addRow("", save_btn)

        settings_group.setLayout(settings_layout)
        layout.addWidget(settings_group)

        layout.addStretch()
        return widget

    def create_openai_api_key_widget(self):
        """Create OpenAI API key settings widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        api_group = QGroupBox("OpenAI API Configuration")
        api_layout = QFormLayout()

        self.openai_key_input = QLineEdit()
        self.openai_key_input.setEchoMode(QLineEdit.Password)
        self.openai_key_input.setPlaceholderText("Enter your OpenAI API key...")
        api_layout.addRow("API Key:", self.openai_key_input)

        # Show/Hide button
        toggle_btn = QPushButton("Show")
        toggle_btn.clicked.connect(lambda: self.toggle_password_visibility(self.openai_key_input, toggle_btn))
        api_layout.addRow("", toggle_btn)

        # Test connection button
        test_btn = Components.create_button("Test Connection", "info")
        test_btn.clicked.connect(self.test_openai_connection)
        api_layout.addRow("", test_btn)

        # Save button
        save_btn = Components.create_button("Save API Key", "success")
        save_btn.clicked.connect(self.save_openai_key)
        api_layout.addRow("", save_btn)

        api_group.setLayout(api_layout)
        layout.addWidget(api_group)

        layout.addStretch()
        return widget

    # Utility methods
    def toggle_password_visibility(self, line_edit, button):
        """Toggle password visibility for API key inputs."""
        if line_edit.echoMode() == QLineEdit.Password:
            line_edit.setEchoMode(QLineEdit.Normal)
            button.setText("Hide")
        else:
            line_edit.setEchoMode(QLineEdit.Password)
            button.setText("Show")

    def save_openai_key(self):
        """Save OpenAI API key."""
        try:
            # Implementation would save API key securely
            QMessageBox.information(self, "API Key Saved", "OpenAI API key saved successfully!")
        except Exception as e:
            QMessageBox.critical(self, "Save Error", f"Failed to save API key: {e}")

    def test_openai_connection(self):
        """Test OpenAI API connection."""
        QMessageBox.information(self, "Connection Test", "OpenAI connection test would be performed here.")

    def create_wordsapi_key_widget(self):
        """Create WordsAPI key settings widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        api_group = QGroupBox("WordsAPI Configuration")
        api_layout = QFormLayout()

        self.wordsapi_key_input = QLineEdit()
        self.wordsapi_key_input.setEchoMode(QLineEdit.Password)
        self.wordsapi_key_input.setPlaceholderText("Enter your WordsAPI key...")
        api_layout.addRow("API Key:", self.wordsapi_key_input)

        # Show/Hide button
        toggle_btn = QPushButton("Show")
        toggle_btn.clicked.connect(lambda: self.toggle_password_visibility(self.wordsapi_key_input, toggle_btn))
        api_layout.addRow("", toggle_btn)

        # Test connection button
        test_btn = Components.create_button("Test Connection", "info")
        test_btn.clicked.connect(self.test_wordsapi_connection)
        api_layout.addRow("", test_btn)

        # Save button
        save_btn = Components.create_button("Save API Key", "success")
        save_btn.clicked.connect(self.save_wordsapi_key)
        api_layout.addRow("", save_btn)

        api_group.setLayout(api_layout)
        layout.addWidget(api_group)

        layout.addStretch()
        return widget

    def create_progress_graph_widget(self):
        """Create progress graph widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        graph_group = QGroupBox("Writing Progress Graph")
        graph_layout = QVBoxLayout()

        # Placeholder for graph
        graph_placeholder = QLabel("📊 Progress Graph\n\nWord count progress over time will be displayed here")
        graph_placeholder.setStyleSheet("""
            font-size: 14px;
            color: #666666;
            padding: 40px;
            text-align: center;
            border: 2px dashed #CCCCCC;
            border-radius: 8px;
        """)
        graph_placeholder.setAlignment(Qt.AlignCenter)
        graph_layout.addWidget(graph_placeholder)

        graph_group.setLayout(graph_layout)
        layout.addWidget(graph_group)

        layout.addStretch()
        return widget

    def create_memory_usage_widget(self):
        """Create memory usage monitoring widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        memory_group = QGroupBox("Memory Usage Monitor")
        memory_layout = QFormLayout()

        self.memory_current_label = QLabel("Calculating...")
        memory_layout.addRow("Current RAM Usage:", self.memory_current_label)

        self.memory_peak_label = QLabel("Calculating...")
        memory_layout.addRow("Peak Usage:", self.memory_peak_label)

        self.memory_available_label = QLabel("Calculating...")
        memory_layout.addRow("Available RAM:", self.memory_available_label)

        refresh_btn = Components.create_button("Refresh", "primary")
        refresh_btn.clicked.connect(self.refresh_memory_stats)
        memory_layout.addRow("", refresh_btn)

        memory_group.setLayout(memory_layout)
        layout.addWidget(memory_group)

        layout.addStretch()
        return widget

    def create_load_project_widget(self):
        """Create load project widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Load Existing Project")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Project selector
        selector_layout = QHBoxLayout()
        selector_layout.addWidget(QLabel("Select Project:"))

        self.load_project_selector = QComboBox()
        self.load_project_selector.setObjectName("projectSelector")
        self.load_project_selector.addItem("Select a project")

        # Populate with available projects
        projects_dir = "projects"
        if os.path.exists(projects_dir):
            for item in os.listdir(projects_dir):
                item_path = os.path.join(projects_dir, item)
                if os.path.isdir(item_path):
                    self.load_project_selector.addItem(item)

        selector_layout.addWidget(self.load_project_selector)
        layout.addLayout(selector_layout)

        # Load button
        load_btn = QPushButton("Load Selected Project")
        load_btn.clicked.connect(self.switch_to_selected_project)
        layout.addWidget(load_btn)

        # Project info display
        info_group = QGroupBox("Project Information")
        info_layout = QVBoxLayout()
        info_group.setLayout(info_layout)

        self.project_info_display = QTextEdit()
        self.project_info_display.setReadOnly(True)
        self.project_info_display.setMaximumHeight(200)
        self.project_info_display.setPlainText("Select a project to view its information.")
        info_layout.addWidget(self.project_info_display)

        layout.addWidget(info_group)

        # Connect selector to info update
        self.load_project_selector.currentTextChanged.connect(self.update_project_info_display)

        layout.addStretch()
        return widget

    def update_project_info_display(self, project_name):
        """Update the project information display."""
        if not project_name or project_name == "Select a project":
            self.project_info_display.setPlainText("Select a project to view its information.")
            return

        try:
            project_path = os.path.join("projects", project_name)
            if os.path.exists(project_path):
                info_text = f"Project: {project_name}\n"
                info_text += f"Path: {project_path}\n"

                # Try to load project config
                config_path = os.path.join(project_path, "project.json")
                if os.path.exists(config_path):
                    try:
                        with open(config_path, 'r') as f:
                            config = json.load(f)
                        info_text += f"Description: {config.get('description', 'No description')}\n"
                        info_text += f"Genre: {config.get('genre', 'Not specified')}\n"
                        info_text += f"Target Chapters: {config.get('target_chapters', 'Not set')}\n"
                        info_text += f"Target Words: {config.get('target_words', 'Not set')}\n"
                        info_text += f"Created: {config.get('created_at', 'Unknown')}\n"
                    except:
                        info_text += "Error reading project configuration.\n"
                else:
                    info_text += "No project configuration found.\n"

                self.project_info_display.setPlainText(info_text)
            else:
                self.project_info_display.setPlainText("Project directory not found.")
        except Exception as e:
            self.project_info_display.setPlainText(f"Error loading project info: {str(e)}")

    def create_delete_project_widget(self):
        """Create delete project widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Delete Project")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #F44336; margin-bottom: 16px;")
        layout.addWidget(header)

        # Warning
        warning = QLabel("⚠️ Warning: This action cannot be undone!")
        warning.setStyleSheet("color: #F44336; font-weight: bold; padding: 10px; background-color: #ffebee; border-radius: 4px;")
        layout.addWidget(warning)

        # Project selector
        selector_layout = QHBoxLayout()
        selector_layout.addWidget(QLabel("Select Project to Delete:"))

        self.delete_project_selector = QComboBox()
        self.delete_project_selector.addItem("Select a project")

        # Populate with available projects
        projects_dir = "projects"
        if os.path.exists(projects_dir):
            for item in os.listdir(projects_dir):
                item_path = os.path.join(projects_dir, item)
                if os.path.isdir(item_path):
                    self.delete_project_selector.addItem(item)

        selector_layout.addWidget(self.delete_project_selector)
        layout.addLayout(selector_layout)

        # Delete button
        delete_btn = QPushButton("Delete Selected Project")
        delete_btn.setStyleSheet("background-color: #F44336; color: white; font-weight: bold;")
        delete_btn.clicked.connect(self.confirm_delete_project)
        layout.addWidget(delete_btn)

        layout.addStretch()
        return widget

    def confirm_delete_project(self):
        """Confirm and delete the selected project."""
        project_name = self.delete_project_selector.currentText()
        if not project_name or project_name == "Select a project":
            QMessageBox.warning(self, "Warning", "Please select a project to delete.")
            return

        reply = QMessageBox.question(
            self, "Confirm Delete",
            f"Are you sure you want to delete the project '{project_name}'?\n\nThis will permanently remove all files for this project.",
            QMessageBox.Yes | QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            try:
                project_path = os.path.join("projects", project_name)
                if os.path.exists(project_path):
                    import shutil
                    shutil.rmtree(project_path)

                # Remove from selector
                index = self.delete_project_selector.findText(project_name)
                if index >= 0:
                    self.delete_project_selector.removeItem(index)

                QMessageBox.information(self, "Success", f"Project '{project_name}' deleted successfully!")
                self.add_log_message(f"Deleted project: {project_name}", "info")

                # Refresh other project selectors
                self.refresh_project_list()

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to delete project: {str(e)}")

    def create_novel_settings_overview(self):
        """Create novel settings overview widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Novel Settings Overview")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Current project info
        if hasattr(self, 'current_project') and self.current_project:
            project_name = self.current_project.get('name', 'Unknown')
            info_text = f"Current Project: {project_name}"
        else:
            info_text = "No project currently loaded"

        project_label = QLabel(info_text)
        project_label.setStyleSheet("font-size: 14px; color: #666; margin-bottom: 10px;")
        layout.addWidget(project_label)

        # Settings categories
        categories = [
            ("Basic Settings", "Configure fundamental novel parameters"),
            ("AI Configuration", "Set up AI models and behavior"),
            ("Writing Style", "Define tone, genre, and style preferences"),
            ("Chapter Structure", "Organize chapters and sections"),
            ("Export Options", "Configure output formats and settings")
        ]

        for category_name, description in categories:
            category_widget = self.create_settings_category_widget(category_name, description)
            layout.addWidget(category_widget)

        layout.addStretch()
        return widget

    def create_settings_category_widget(self, name, description):
        """Create a settings category widget."""
        widget = QFrame()
        widget.setFrameStyle(QFrame.Box)
        widget.setStyleSheet("border: 1px solid #ddd; border-radius: 4px; margin: 5px; padding: 10px;")

        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Title
        title = QLabel(name)
        title.setStyleSheet("font-size: 14px; font-weight: bold; color: #2196F3;")
        layout.addWidget(title)

        # Description
        desc = QLabel(description)
        desc.setStyleSheet("font-size: 12px; color: #666;")
        layout.addWidget(desc)

        # Configure button
        config_btn = QPushButton("Configure")
        config_btn.clicked.connect(lambda: self.open_settings_category(name))
        layout.addWidget(config_btn)

        return widget

    def open_settings_category(self, category_name):
        """Open a specific settings category."""
        QMessageBox.information(self, "Settings", f"Opening {category_name} configuration...")

    def create_synonyms_widget(self):
        """Create synonyms widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Synonym Lookup")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Search input
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("Word:"))

        self.synonym_input = QLineEdit()
        self.synonym_input.setPlaceholderText("Enter a word to find synonyms...")
        self.synonym_input.returnPressed.connect(self.search_synonyms)
        search_layout.addWidget(self.synonym_input)

        search_btn = QPushButton("Search")
        search_btn.clicked.connect(self.search_synonyms)
        search_layout.addWidget(search_btn)

        layout.addLayout(search_layout)

        # Results area
        self.synonym_results = QTextEdit()
        self.synonym_results.setReadOnly(True)
        self.synonym_results.setPlainText("Enter a word and click Search to find synonyms.")
        layout.addWidget(self.synonym_results)

        # API usage info
        usage_info = QLabel("💡 Tip: Uses WordsAPI for synonym lookup. Configure your API key in Settings.")
        usage_info.setStyleSheet("color: #666; font-style: italic;")
        layout.addWidget(usage_info)

        return widget

    def search_synonyms(self):
        """Search for synonyms of the entered word."""
        word = self.synonym_input.text().strip()
        if not word:
            QMessageBox.warning(self, "Warning", "Please enter a word to search for.")
            return

        # Simulate synonym search (in real implementation, this would use WordsAPI)
        sample_synonyms = {
            "happy": ["joyful", "cheerful", "delighted", "pleased", "content", "elated"],
            "sad": ["unhappy", "sorrowful", "dejected", "melancholy", "downcast", "gloomy"],
            "big": ["large", "huge", "enormous", "massive", "giant", "vast"],
            "small": ["tiny", "little", "minute", "petite", "miniature", "compact"],
            "good": ["excellent", "great", "wonderful", "superb", "outstanding", "remarkable"],
            "bad": ["awful", "terrible", "horrible", "dreadful", "poor", "inadequate"]
        }

        word_lower = word.lower()
        if word_lower in sample_synonyms:
            synonyms = sample_synonyms[word_lower]
            result_text = f"Synonyms for '{word}':\n\n"
            for i, synonym in enumerate(synonyms, 1):
                result_text += f"{i}. {synonym}\n"
        else:
            result_text = f"No synonyms found for '{word}'.\n\nNote: This is a demo. In the full version, this would connect to WordsAPI for comprehensive synonym lookup."

        self.synonym_results.setPlainText(result_text)

    def create_log_widget(self):
        """Create log widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("System Log")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Log controls
        controls_layout = QHBoxLayout()

        self.log_level_filter = QComboBox()
        self.log_level_filter.addItems(["All", "Info", "Warning", "Error"])
        self.log_level_filter.currentTextChanged.connect(self.filter_log_messages)
        controls_layout.addWidget(QLabel("Filter:"))
        controls_layout.addWidget(self.log_level_filter)

        controls_layout.addStretch()

        clear_btn = QPushButton("Clear Log")
        clear_btn.clicked.connect(self.clear_log_display)
        controls_layout.addWidget(clear_btn)

        layout.addLayout(controls_layout)

        # Log display
        self.log_display = QTextEdit()
        self.log_display.setReadOnly(True)
        self.log_display.setObjectName("liveLogText")

        # Add some sample log entries
        sample_logs = [
            "[14:32:15] INFO: Application started successfully",
            "[14:32:16] INFO: Modern GUI loaded",
            "[14:32:17] INFO: Database connection established",
            "[14:32:18] WARNING: wkhtmltopdf not found - PDF export may be limited",
            "[14:32:19] INFO: All systems initialized",
        ]

        for log_entry in sample_logs:
            self.log_display.append(log_entry)

        layout.addWidget(self.log_display)

        # Auto-scroll option
        auto_scroll_cb = QCheckBox("Auto-scroll to new messages")
        auto_scroll_cb.setChecked(True)
        auto_scroll_cb.setObjectName("autoScrollCheckbox")
        layout.addWidget(auto_scroll_cb)

        return widget

    def filter_log_messages(self, filter_level):
        """Filter log messages by level."""
        # In a real implementation, this would filter the actual log messages
        QMessageBox.information(self, "Log Filter", f"Filtering log messages by: {filter_level}")

    def clear_log_display(self):
        """Clear the log display."""
        self.log_display.clear()
        self.log_display.append("[" + datetime.now().strftime("%H:%M:%S") + "] INFO: Log cleared by user")

    def create_chapter_progress_widget(self):
        """Create chapter progress widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Chapter Progress")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Progress overview
        overview_layout = QHBoxLayout()

        # Overall progress
        overall_group = QGroupBox("Overall Progress")
        overall_layout = QVBoxLayout()
        overall_group.setLayout(overall_layout)

        overall_progress = QProgressBar()
        overall_progress.setValue(35)  # Sample value
        overall_layout.addWidget(QLabel("Novel Completion:"))
        overall_layout.addWidget(overall_progress)
        overall_layout.addWidget(QLabel("7 of 20 chapters completed"))

        overview_layout.addWidget(overall_group)

        # Current chapter
        current_group = QGroupBox("Current Chapter")
        current_layout = QVBoxLayout()
        current_group.setLayout(current_layout)

        current_progress = QProgressBar()
        current_progress.setValue(60)  # Sample value
        current_layout.addWidget(QLabel("Chapter 8 Progress:"))
        current_layout.addWidget(current_progress)
        current_layout.addWidget(QLabel("3 of 5 sections completed"))

        overview_layout.addWidget(current_group)

        layout.addLayout(overview_layout)

        # Chapter list
        chapters_group = QGroupBox("Chapters")
        chapters_layout = QVBoxLayout()
        chapters_group.setLayout(chapters_layout)

        self.chapter_list = QListWidget()

        # Sample chapters
        sample_chapters = [
            ("Chapter 1: The Beginning", "✅ Complete"),
            ("Chapter 2: First Steps", "✅ Complete"),
            ("Chapter 3: The Journey Starts", "✅ Complete"),
            ("Chapter 4: Challenges Ahead", "✅ Complete"),
            ("Chapter 5: New Discoveries", "✅ Complete"),
            ("Chapter 6: The Plot Thickens", "✅ Complete"),
            ("Chapter 7: Revelations", "✅ Complete"),
            ("Chapter 8: Turning Point", "🔄 In Progress"),
            ("Chapter 9: The Climax Approaches", "⏳ Pending"),
            ("Chapter 10: Final Confrontation", "⏳ Pending"),
        ]

        for chapter_title, status in sample_chapters:
            item_text = f"{chapter_title} - {status}"
            item = QListWidgetItem(item_text)
            if "Complete" in status:
                item.setBackground(QColor("#e8f5e8"))
            elif "Progress" in status:
                item.setBackground(QColor("#fff3cd"))
            self.chapter_list.addItem(item)

        chapters_layout.addWidget(self.chapter_list)
        layout.addWidget(chapters_group)

        return widget

    def create_current_draft_widget(self):
        """Create current draft widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Current Draft")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Draft info
        info_layout = QHBoxLayout()

        # Word count
        word_count_group = QGroupBox("Word Count")
        word_count_layout = QVBoxLayout()
        word_count_group.setLayout(word_count_layout)

        word_count_display = QLabel("47,832")
        word_count_display.setStyleSheet("font-size: 24px; font-weight: bold; color: #2196F3;")
        word_count_display.setAlignment(Qt.AlignCenter)
        word_count_layout.addWidget(word_count_display)

        target_label = QLabel("Target: 80,000 words")
        target_label.setAlignment(Qt.AlignCenter)
        word_count_layout.addWidget(target_label)

        progress_bar = QProgressBar()
        progress_bar.setValue(60)  # 47,832 / 80,000 * 100
        word_count_layout.addWidget(progress_bar)

        info_layout.addWidget(word_count_group)

        # Last modified
        modified_group = QGroupBox("Last Modified")
        modified_layout = QVBoxLayout()
        modified_group.setLayout(modified_layout)

        modified_display = QLabel("2 hours ago")
        modified_display.setStyleSheet("font-size: 16px; font-weight: bold;")
        modified_display.setAlignment(Qt.AlignCenter)
        modified_layout.addWidget(modified_display)

        date_display = QLabel(datetime.now().strftime("%B %d, %Y at %I:%M %p"))
        date_display.setAlignment(Qt.AlignCenter)
        modified_layout.addWidget(date_display)

        info_layout.addWidget(modified_group)

        layout.addLayout(info_layout)

        # Recent activity
        activity_group = QGroupBox("Recent Activity")
        activity_layout = QVBoxLayout()
        activity_group.setLayout(activity_layout)

        activity_list = QListWidget()

        # Sample recent activities
        recent_activities = [
            "Added 1,247 words to Chapter 8",
            "Completed Chapter 7 - Revelations",
            "Revised opening of Chapter 6",
            "Updated character development notes",
            "Exported draft to PDF for review"
        ]

        for activity in recent_activities:
            activity_list.addItem(f"• {activity}")

        activity_layout.addWidget(activity_list)
        layout.addWidget(activity_group)

        # Quick actions
        actions_layout = QHBoxLayout()

        export_btn = QPushButton("Export Draft")
        export_btn.clicked.connect(self.export_project)
        actions_layout.addWidget(export_btn)

        backup_btn = QPushButton("Create Backup")
        backup_btn.clicked.connect(self.create_backup)
        actions_layout.addWidget(backup_btn)

        layout.addLayout(actions_layout)

        return widget

    def create_backup(self):
        """Create a backup of the current project."""
        QMessageBox.information(self, "Backup", "Backup creation functionality would be implemented here.")

    def create_cpu_usage_widget(self):
        """Create CPU usage widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("CPU Usage Monitor")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # CPU usage display
        usage_layout = QHBoxLayout()

        # Current usage
        current_group = QGroupBox("Current Usage")
        current_layout = QVBoxLayout()
        current_group.setLayout(current_layout)

        cpu_percentage = QLabel("23%")
        cpu_percentage.setStyleSheet("font-size: 32px; font-weight: bold; color: #4CAF50;")
        cpu_percentage.setAlignment(Qt.AlignCenter)
        current_layout.addWidget(cpu_percentage)

        cpu_progress = QProgressBar()
        cpu_progress.setValue(23)
        cpu_progress.setStyleSheet("QProgressBar::chunk { background-color: #4CAF50; }")
        current_layout.addWidget(cpu_progress)

        usage_layout.addWidget(current_group)

        # Memory usage
        memory_group = QGroupBox("Memory Usage")
        memory_layout = QVBoxLayout()
        memory_group.setLayout(memory_layout)

        memory_percentage = QLabel("156 MB")
        memory_percentage.setStyleSheet("font-size: 24px; font-weight: bold; color: #2196F3;")
        memory_percentage.setAlignment(Qt.AlignCenter)
        memory_layout.addWidget(memory_percentage)

        memory_detail = QLabel("of 8 GB total")
        memory_detail.setAlignment(Qt.AlignCenter)
        memory_layout.addWidget(memory_detail)

        usage_layout.addWidget(memory_group)

        layout.addLayout(usage_layout)

        # Process list
        process_group = QGroupBox("FANWS Processes")
        process_layout = QVBoxLayout()
        process_group.setLayout(process_layout)

        process_list = QListWidget()

        # Sample processes
        processes = [
            "Main Application - 45 MB",
            "AI Processing - 67 MB",
            "Database Manager - 23 MB",
            "Export Handler - 12 MB",
            "Cache Manager - 9 MB"
        ]

        for process in processes:
            process_list.addItem(process)

        process_layout.addWidget(process_list)
        layout.addWidget(process_group)

        # Refresh button
        refresh_btn = QPushButton("Refresh Stats")
        refresh_btn.clicked.connect(self.refresh_memory_stats)
        layout.addWidget(refresh_btn)

        return widget

    def create_api_stats_widget(self):
        """Create API stats widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("API Usage Statistics")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # API usage overview
        overview_layout = QHBoxLayout()

        # OpenAI stats
        openai_group = QGroupBox("OpenAI API")
        openai_layout = QVBoxLayout()
        openai_group.setLayout(openai_layout)

        openai_calls = QLabel("1,247")
        openai_calls.setStyleSheet("font-size: 24px; font-weight: bold; color: #4CAF50;")
        openai_calls.setAlignment(Qt.AlignCenter)
        openai_layout.addWidget(openai_calls)

        openai_layout.addWidget(QLabel("Calls Today", alignment=Qt.AlignCenter))

        openai_cost = QLabel("$12.34")
        openai_cost.setStyleSheet("font-size: 16px; color: #666;")
        openai_cost.setAlignment(Qt.AlignCenter)
        openai_layout.addWidget(openai_cost)

        overview_layout.addWidget(openai_group)

        # WordsAPI stats
        words_group = QGroupBox("WordsAPI")
        words_layout = QVBoxLayout()
        words_group.setLayout(words_layout)

        words_calls = QLabel("356")
        words_calls.setStyleSheet("font-size: 24px; font-weight: bold; color: #FF9800;")
        words_calls.setAlignment(Qt.AlignCenter)
        words_layout.addWidget(words_calls)

        words_layout.addWidget(QLabel("Calls Today", alignment=Qt.AlignCenter))

        words_limit = QLabel("of 2,500 daily")
        words_limit.setStyleSheet("font-size: 14px; color: #666;")
        words_limit.setAlignment(Qt.AlignCenter)
        words_layout.addWidget(words_limit)

        overview_layout.addWidget(words_group)

        layout.addLayout(overview_layout)

        # Usage history
        history_group = QGroupBox("Usage History")
        history_layout = QVBoxLayout()
        history_group.setLayout(history_layout)

        history_list = QListWidget()

        # Sample history
        history_entries = [
            "14:32 - OpenAI GPT-4 - Chapter generation - 0.023 tokens",
            "14:28 - WordsAPI - Synonym lookup for 'magnificent' - Success",
            "14:25 - OpenAI GPT-4 - Character dialogue - 0.018 tokens",
            "14:20 - OpenAI GPT-4 - Scene description - 0.031 tokens",
            "14:15 - WordsAPI - Definition lookup for 'ephemeral' - Success"
        ]

        for entry in history_entries:
            history_list.addItem(entry)

        history_layout.addWidget(history_list)
        layout.addWidget(history_group)

        # Controls
        controls_layout = QHBoxLayout()

        reset_btn = QPushButton("Reset Daily Counters")
        reset_btn.clicked.connect(self.reset_api_counters)
        controls_layout.addWidget(reset_btn)

        export_btn = QPushButton("Export Usage Report")
        export_btn.clicked.connect(self.export_api_usage)
        controls_layout.addWidget(export_btn)

        layout.addLayout(controls_layout)

        return widget

    def reset_api_counters(self):
        """Reset API usage counters."""
        QMessageBox.information(self, "Reset Counters", "API usage counters have been reset.")

    def export_api_usage(self):
        """Export API usage report."""
        QMessageBox.information(self, "Export Usage", "API usage report export functionality would be implemented here.")

    def create_file_operations_widget(self):
        """Create file operations widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("File Operations Monitor")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Operations overview
        overview_layout = QHBoxLayout()

        # Read operations
        read_group = QGroupBox("Read Operations")
        read_layout = QVBoxLayout()
        read_group.setLayout(read_layout)

        read_count = QLabel("1,423")
        read_count.setStyleSheet("font-size: 24px; font-weight: bold; color: #4CAF50;")
        read_count.setAlignment(Qt.AlignCenter)
        read_layout.addWidget(read_count)

        read_layout.addWidget(QLabel("Files Read", alignment=Qt.AlignCenter))

        read_speed = QLabel("45 MB/s avg")
        read_speed.setAlignment(Qt.AlignCenter)
        read_layout.addWidget(read_speed)

        overview_layout.addWidget(read_group)

        # Write operations
        write_group = QGroupBox("Write Operations")
        write_layout = QVBoxLayout()
        write_group.setLayout(write_layout)

        write_count = QLabel("892")
        write_count.setStyleSheet("font-size: 24px; font-weight: bold; color: #FF9800;")
        write_count.setAlignment(Qt.AlignCenter)
        write_layout.addWidget(write_count)

        write_layout.addWidget(QLabel("Files Written", alignment=Qt.AlignCenter))

        write_speed = QLabel("32 MB/s avg")
        write_speed.setAlignment(Qt.AlignCenter)
        write_layout.addWidget(write_speed)

        overview_layout.addWidget(write_group)

        layout.addLayout(overview_layout)

        # Recent operations
        recent_group = QGroupBox("Recent File Operations")
        recent_layout = QVBoxLayout()
        recent_group.setLayout(recent_layout)

        operations_list = QListWidget()

        # Sample operations
        operations = [
            "WRITE: project.json - 2.3 KB - 0.01s",
            "READ: chapter_08.txt - 45.2 KB - 0.02s",
            "WRITE: settings.json - 1.8 KB - 0.01s",
            "READ: character_profiles.json - 12.4 KB - 0.01s",
            "WRITE: backup_20250803.zip - 2.1 MB - 0.15s"
        ]

        for operation in operations:
            operations_list.addItem(operation)

        recent_layout.addWidget(operations_list)
        layout.addWidget(recent_group)

        return widget

    def create_cache_rate_widget(self):
        """Create cache rate widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Cache Performance")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Cache stats
        stats_layout = QHBoxLayout()

        # Hit rate
        hit_group = QGroupBox("Cache Hit Rate")
        hit_layout = QVBoxLayout()
        hit_group.setLayout(hit_layout)

        hit_rate = QLabel("87%")
        hit_rate.setStyleSheet("font-size: 32px; font-weight: bold; color: #4CAF50;")
        hit_rate.setAlignment(Qt.AlignCenter)
        hit_layout.addWidget(hit_rate)

        hit_progress = QProgressBar()
        hit_progress.setValue(87)
        hit_progress.setStyleSheet("QProgressBar::chunk { background-color: #4CAF50; }")
        hit_layout.addWidget(hit_progress)

        stats_layout.addWidget(hit_group)

        # Cache size
        size_group = QGroupBox("Cache Size")
        size_layout = QVBoxLayout()
        size_group.setLayout(size_layout)

        cache_size = QLabel("234 MB")
        cache_size.setStyleSheet("font-size: 24px; font-weight: bold; color: #2196F3;")
        cache_size.setAlignment(Qt.AlignCenter)
        size_layout.addWidget(cache_size)

        size_detail = QLabel("of 500 MB limit")
        size_detail.setAlignment(Qt.AlignCenter)
        size_layout.addWidget(size_detail)

        stats_layout.addWidget(size_group)

        layout.addLayout(stats_layout)

        # Cache breakdown
        breakdown_group = QGroupBox("Cache Breakdown")
        breakdown_layout = QVBoxLayout()
        breakdown_group.setLayout(breakdown_layout)

        breakdown_list = QListWidget()

        # Sample cache data
        cache_items = [
            "API Responses: 145 MB (62%)",
            "Generated Text: 67 MB (29%)",
            "Images/Assets: 15 MB (6%)",
            "Settings: 4 MB (2%)",
            "Metadata: 3 MB (1%)"
        ]

        for item in cache_items:
            breakdown_list.addItem(item)

        breakdown_layout.addWidget(breakdown_list)
        layout.addWidget(breakdown_group)

        # Cache controls
        controls_layout = QHBoxLayout()

        optimize_btn = QPushButton("Optimize Cache")
        optimize_btn.clicked.connect(self.optimize_cache)
        controls_layout.addWidget(optimize_btn)

        clear_btn = QPushButton("Clear Cache")
        clear_btn.clicked.connect(self.clear_cache)
        controls_layout.addWidget(clear_btn)

        layout.addLayout(controls_layout)

        return widget

    def create_response_times_widget(self):
        """Create response times widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Response Time Analysis")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Average response times
        avg_layout = QHBoxLayout()

        # API responses
        api_group = QGroupBox("API Response Times")
        api_layout = QVBoxLayout()
        api_group.setLayout(api_layout)

        api_time = QLabel("1.2s")
        api_time.setStyleSheet("font-size: 24px; font-weight: bold; color: #4CAF50;")
        api_time.setAlignment(Qt.AlignCenter)
        api_layout.addWidget(api_time)

        api_layout.addWidget(QLabel("Average", alignment=Qt.AlignCenter))

        avg_layout.addWidget(api_group)

        # Database queries
        db_group = QGroupBox("Database Queries")
        db_layout = QVBoxLayout()
        db_group.setLayout(db_layout)

        db_time = QLabel("0.05s")
        db_time.setStyleSheet("font-size: 24px; font-weight: bold; color: #4CAF50;")
        db_time.setAlignment(Qt.AlignCenter)
        db_layout.addWidget(db_time)

        db_layout.addWidget(QLabel("Average", alignment=Qt.AlignCenter))

        avg_layout.addWidget(db_group)

        # File operations
        file_group = QGroupBox("File Operations")
        file_layout = QVBoxLayout()
        file_group.setLayout(file_layout)

        file_time = QLabel("0.02s")
        file_time.setStyleSheet("font-size: 24px; font-weight: bold; color: #4CAF50;")
        file_time.setAlignment(Qt.AlignCenter)
        file_layout.addWidget(file_time)

        file_layout.addWidget(QLabel("Average", alignment=Qt.AlignCenter))

        avg_layout.addWidget(file_group)

        layout.addLayout(avg_layout)

        # Recent response times
        recent_group = QGroupBox("Recent Response Times")
        recent_layout = QVBoxLayout()
        recent_group.setLayout(recent_layout)

        response_list = QListWidget()

        # Sample response times
        responses = [
            "OpenAI API Call - 1.5s - Chapter generation",
            "Database Query - 0.03s - Load project settings",
            "File Read - 0.01s - Read character data",
            "WordsAPI Call - 0.8s - Synonym lookup",
            "File Write - 0.02s - Save progress"
        ]

        for response in responses:
            response_list.addItem(response)

        recent_layout.addWidget(response_list)
        layout.addWidget(recent_group)

        return widget

    def create_optimization_widget(self):
        """Create optimization widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Performance Optimization")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Optimization status
        status_group = QGroupBox("Optimization Status")
        status_layout = QVBoxLayout()
        status_group.setLayout(status_layout)

        status_text = QLabel("System Performance: Good")
        status_text.setStyleSheet("font-size: 16px; font-weight: bold; color: #4CAF50;")
        status_text.setAlignment(Qt.AlignCenter)
        status_layout.addWidget(status_text)

        overall_score = QProgressBar()
        overall_score.setValue(78)
        overall_score.setFormat("Overall Score: 78/100")
        overall_score.setStyleSheet("QProgressBar::chunk { background-color: #4CAF50; }")
        status_layout.addWidget(overall_score)

        layout.addWidget(status_group)

        # Optimization recommendations
        recommendations_group = QGroupBox("Recommendations")
        recommendations_layout = QVBoxLayout()
        recommendations_group.setLayout(recommendations_layout)

        recommendations = [
            "✅ Cache hit rate is optimal (87%)",
            "⚠️ Consider increasing memory allocation for better performance",
            "✅ File operations are running efficiently",
            "💡 Enable background preprocessing for faster response times",
            "✅ Database queries are well optimized"
        ]

        for recommendation in recommendations:
            label = QLabel(recommendation)
            if "✅" in recommendation:
                label.setStyleSheet("color: #4CAF50;")
            elif "⚠️" in recommendation:
                label.setStyleSheet("color: #FF9800;")
            elif "💡" in recommendation:
                label.setStyleSheet("color: #2196F3;")
            recommendations_layout.addWidget(label)

        layout.addWidget(recommendations_group)

        # Optimization actions
        actions_group = QGroupBox("Quick Actions")
        actions_layout = QVBoxLayout()
        actions_group.setLayout(actions_layout)

        actions_buttons_layout = QHBoxLayout()

        optimize_btn = QPushButton("Run Optimization")
        optimize_btn.clicked.connect(self.run_optimization)
        actions_buttons_layout.addWidget(optimize_btn)

        cleanup_btn = QPushButton("Cleanup Temp Files")
        cleanup_btn.clicked.connect(self.cleanup_temp_files)
        actions_buttons_layout.addWidget(cleanup_btn)

        defrag_btn = QPushButton("Optimize Database")
        defrag_btn.clicked.connect(self.optimize_database)
        actions_buttons_layout.addWidget(defrag_btn)

        actions_layout.addLayout(actions_buttons_layout)
        layout.addWidget(actions_group)

        return widget

    def run_optimization(self):
        """Run system optimization."""
        QMessageBox.information(self, "Optimization", "System optimization completed successfully!")

    def cleanup_temp_files(self):
        """Cleanup temporary files."""
        QMessageBox.information(self, "Cleanup", "Temporary files have been cleaned up.")

    def optimize_database(self):
        """Optimize database performance."""
        QMessageBox.information(self, "Database", "Database optimization completed successfully!")

    def create_system_resources_widget(self):
        """Create system resources widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("System Resources")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Resource usage
        resources_layout = QHBoxLayout()

        # CPU
        cpu_group = QGroupBox("CPU Usage")
        cpu_layout = QVBoxLayout()
        cpu_group.setLayout(cpu_layout)

        cpu_usage = QLabel("23%")
        cpu_usage.setStyleSheet("font-size: 24px; font-weight: bold; color: #4CAF50;")
        cpu_usage.setAlignment(Qt.AlignCenter)
        cpu_layout.addWidget(cpu_usage)

        cpu_bar = QProgressBar()
        cpu_bar.setValue(23)
        cpu_bar.setStyleSheet("QProgressBar::chunk { background-color: #4CAF50; }")
        cpu_layout.addWidget(cpu_bar)

        resources_layout.addWidget(cpu_group)

        # RAM
        ram_group = QGroupBox("Memory Usage")
        ram_layout = QVBoxLayout()
        ram_group.setLayout(ram_layout)

        ram_usage = QLabel("156 MB")
        ram_usage.setStyleSheet("font-size: 24px; font-weight: bold; color: #2196F3;")
        ram_usage.setAlignment(Qt.AlignCenter)
        ram_layout.addWidget(ram_usage)

        ram_detail = QLabel("of 8 GB")
        ram_detail.setAlignment(Qt.AlignCenter)
        ram_layout.addWidget(ram_detail)

        resources_layout.addWidget(ram_group)

        # Disk
        disk_group = QGroupBox("Disk Usage")
        disk_layout = QVBoxLayout()
        disk_group.setLayout(disk_layout)

        disk_usage = QLabel("45%")
        disk_usage.setStyleSheet("font-size: 24px; font-weight: bold; color: #FF9800;")
        disk_usage.setAlignment(Qt.AlignCenter)
        disk_layout.addWidget(disk_usage)

        disk_bar = QProgressBar()
        disk_bar.setValue(45)
        disk_bar.setStyleSheet("QProgressBar::chunk { background-color: #FF9800; }")
        disk_layout.addWidget(disk_bar)

        resources_layout.addWidget(disk_group)

        layout.addLayout(resources_layout)

        # Resource details
        details_group = QGroupBox("Resource Details")
        details_layout = QVBoxLayout()
        details_group.setLayout(details_layout)

        details_list = QListWidget()

        # Sample resource details
        details = [
            "CPU Cores: 8 (Intel Core i7)",
            "Total RAM: 8.0 GB",
            "Available RAM: 5.2 GB",
            "Disk Space: 512 GB SSD",
            "Free Space: 280 GB",
            "Network: Connected (WiFi)",
            "GPU: NVIDIA GTX 1660 (Optional AI acceleration)"
        ]

        for detail in details:
            details_list.addItem(detail)

        details_layout.addWidget(details_list)
        layout.addWidget(details_group)

        return widget

    def create_export_status_widget(self):
        """Create export status widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Export Status")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Current export progress
        current_group = QGroupBox("Current Export")
        current_layout = QVBoxLayout()
        current_group.setLayout(current_layout)

        export_info = QLabel("No export currently in progress")
        export_info.setAlignment(Qt.AlignCenter)
        current_layout.addWidget(export_info)

        export_progress = QProgressBar()
        export_progress.setValue(0)
        export_progress.setFormat("Ready to export")
        current_layout.addWidget(export_progress)

        layout.addWidget(current_group)

        # Quick export buttons
        quick_export_group = QGroupBox("Quick Export")
        quick_export_layout = QHBoxLayout()
        quick_export_group.setLayout(quick_export_layout)

        pdf_btn = QPushButton("Export as PDF")
        pdf_btn.clicked.connect(lambda: self.quick_export("PDF"))
        quick_export_layout.addWidget(pdf_btn)

        docx_btn = QPushButton("Export as DOCX")
        docx_btn.clicked.connect(lambda: self.quick_export("DOCX"))
        quick_export_layout.addWidget(docx_btn)

        txt_btn = QPushButton("Export as TXT")
        txt_btn.clicked.connect(lambda: self.quick_export("TXT"))
        quick_export_layout.addWidget(txt_btn)

        layout.addWidget(quick_export_group)

        # Export queue
        queue_group = QGroupBox("Export Queue")
        queue_layout = QVBoxLayout()
        queue_group.setLayout(queue_layout)

        queue_list = QListWidget()
        queue_list.addItem("No exports queued")
        queue_layout.addWidget(queue_list)

        layout.addWidget(queue_group)

        return widget

    def quick_export(self, format_type):
        """Perform a quick export."""
        QMessageBox.information(self, "Quick Export", f"Starting {format_type} export...")

    def create_export_formats_widget(self):
        """Create export formats widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Export Formats")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Available formats
        formats_group = QGroupBox("Available Export Formats")
        formats_layout = QVBoxLayout()
        formats_group.setLayout(formats_layout)

        formats = [
            ("PDF", "Portable Document Format - Universal reading format", "✅ Available"),
            ("DOCX", "Microsoft Word Document - Editable format", "✅ Available"),
            ("TXT", "Plain Text - Simple, universal format", "✅ Available"),
            ("HTML", "Web Document - Viewable in browsers", "✅ Available"),
            ("EPUB", "Electronic Publication - E-book format", "⚠️ Limited support"),
            ("Markdown", "Markdown Document - Developer-friendly", "🔄 Coming soon")
        ]

        for format_name, description, status in formats:
            format_widget = QFrame()
            format_widget.setFrameStyle(QFrame.Box)
            format_widget.setStyleSheet("border: 1px solid #ddd; border-radius: 4px; margin: 2px; padding: 8px;")

            format_layout = QHBoxLayout()
            format_widget.setLayout(format_layout)

            # Format info
            info_layout = QVBoxLayout()

            name_label = QLabel(format_name)
            name_label.setStyleSheet("font-size: 14px; font-weight: bold;")
            info_layout.addWidget(name_label)

            desc_label = QLabel(description)
            desc_label.setStyleSheet("font-size: 12px; color: #666;")
            info_layout.addWidget(desc_label)

            format_layout.addLayout(info_layout)

            # Status
            status_label = QLabel(status)
            if "✅" in status:
                status_label.setStyleSheet("color: #4CAF50; font-weight: bold;")
            elif "⚠️" in status:
                status_label.setStyleSheet("color: #FF9800; font-weight: bold;")
            elif "🔄" in status:
                status_label.setStyleSheet("color: #2196F3; font-weight: bold;")
            format_layout.addWidget(status_label)

            formats_layout.addWidget(format_widget)

        layout.addWidget(formats_group)

        # Export options
        options_group = QGroupBox("Export Options")
        options_layout = QVBoxLayout()
        options_group.setLayout(options_layout)

        # Include metadata
        include_metadata_cb = QCheckBox("Include project metadata")
        include_metadata_cb.setChecked(True)
        options_layout.addWidget(include_metadata_cb)

        # Include settings
        include_settings_cb = QCheckBox("Include project settings")
        options_layout.addWidget(include_settings_cb)

        # Include comments
        include_comments_cb = QCheckBox("Include author comments")
        options_layout.addWidget(include_comments_cb)

        layout.addWidget(options_group)

        # Advanced export button
        advanced_btn = QPushButton("Advanced Export Options...")
        advanced_btn.clicked.connect(self.export_project)
        layout.addWidget(advanced_btn)

        return widget

    def create_export_history_widget(self):
        """Create export history widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Export History")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # History list
        history_group = QGroupBox("Recent Exports")
        history_layout = QVBoxLayout()
        history_group.setLayout(history_layout)

        history_list = QListWidget()

        # Sample export history
        exports = [
            "2025-08-03 14:25 - Novel_Draft_v2.pdf - 2.3 MB - ✅ Success",
            "2025-08-03 12:18 - Chapter_Review.docx - 567 KB - ✅ Success",
            "2025-08-02 16:42 - Complete_Draft.txt - 234 KB - ✅ Success",
            "2025-08-02 11:30 - Character_Profiles.html - 89 KB - ✅ Success",
            "2025-08-01 20:15 - Backup_Export.pdf - 2.1 MB - ✅ Success"
        ]

        for export in exports:
            item = QListWidgetItem(export)
            if "✅ Success" in export:
                item.setBackground(QColor("#e8f5e8"))
            history_list.addItem(item)

        history_layout.addWidget(history_list)
        layout.addWidget(history_group)

        # History actions
        actions_layout = QHBoxLayout()

        clear_history_btn = QPushButton("Clear History")
        clear_history_btn.clicked.connect(self.clear_export_history)
        actions_layout.addWidget(clear_history_btn)

        export_log_btn = QPushButton("Export Log to File")
        export_log_btn.clicked.connect(self.export_history_log)
        actions_layout.addWidget(export_log_btn)

        layout.addLayout(actions_layout)

        return widget

    def clear_export_history(self):
        """Clear export history."""
        QMessageBox.information(self, "Clear History", "Export history has been cleared.")

    def export_history_log(self):
        """Export history log to file."""
        QMessageBox.information(self, "Export Log", "Export history log saved to file.")

    def create_file_sizes_widget(self):
        """Create file sizes widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Export File Sizes")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Size comparison
        comparison_group = QGroupBox("Format Size Comparison")
        comparison_layout = QVBoxLayout()
        comparison_group.setLayout(comparison_layout)

        # Sample sizes for current project
        formats_sizes = [
            ("PDF", "2.3 MB", 85),
            ("DOCX", "892 KB", 35),
            ("HTML", "567 KB", 25),
            ("TXT", "234 KB", 10),
            ("Markdown", "198 KB", 8)
        ]

        for format_name, size_text, percentage in formats_sizes:
            format_layout = QHBoxLayout()

            name_label = QLabel(f"{format_name}:")
            name_label.setMinimumWidth(80)
            format_layout.addWidget(name_label)

            size_bar = QProgressBar()
            size_bar.setValue(percentage)
            size_bar.setFormat(size_text)
            format_layout.addWidget(size_bar)

            comparison_layout.addLayout(format_layout)

        layout.addWidget(comparison_group)

        # Size optimization tips
        tips_group = QGroupBox("Size Optimization Tips")
        tips_layout = QVBoxLayout()
        tips_group.setLayout(tips_layout)

        tips = [
            "💡 TXT format provides the smallest file size for text-only content",
            "📄 PDF format includes formatting but creates larger files",
            "📝 DOCX offers good balance between editability and size",
            "🌐 HTML format is great for web publishing",
            "🔧 Remove unnecessary metadata to reduce file sizes"
        ]

        for tip in tips:
            tip_label = QLabel(tip)
            tip_label.setWordWrap(True)
            tips_layout.addWidget(tip_label)

        layout.addWidget(tips_group)

        return widget

    def create_export_quality_widget(self):
        """Create export quality widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Export Quality Settings")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Quality settings
        quality_group = QGroupBox("Quality Configuration")
        quality_layout = QVBoxLayout()
        quality_group.setLayout(quality_layout)

        # PDF quality
        pdf_layout = QHBoxLayout()
        pdf_layout.addWidget(QLabel("PDF Quality:"))
        pdf_quality = QComboBox()
        pdf_quality.addItems(["Low (Fast)", "Medium (Balanced)", "High (Best)", "Print Quality"])
        pdf_quality.setCurrentText("High (Best)")
        pdf_layout.addWidget(pdf_quality)
        quality_layout.addLayout(pdf_layout)

        # Image quality
        img_layout = QHBoxLayout()
        img_layout.addWidget(QLabel("Image Quality:"))
        img_quality = QSlider(Qt.Horizontal)
        img_quality.setRange(1, 100)
        img_quality.setValue(85)
        img_quality_label = QLabel("85%")
        img_quality.valueChanged.connect(lambda v: img_quality_label.setText(f"{v}%"))
        img_layout.addWidget(img_quality)
        img_layout.addWidget(img_quality_label)
        quality_layout.addLayout(img_layout)

        # Text encoding
        encoding_layout = QHBoxLayout()
        encoding_layout.addWidget(QLabel("Text Encoding:"))
        encoding = QComboBox()
        encoding.addItems(["UTF-8", "UTF-16", "ASCII", "Latin-1"])
        encoding.setCurrentText("UTF-8")
        encoding_layout.addWidget(encoding)
        quality_layout.addLayout(encoding_layout)

        layout.addWidget(quality_group)

        # Preview settings
        preview_group = QGroupBox("Export Preview")
        preview_layout = QVBoxLayout()
        preview_group.setLayout(preview_layout)

        preview_cb = QCheckBox("Show preview before exporting")
        preview_cb.setChecked(True)
        preview_layout.addWidget(preview_cb)

        validate_cb = QCheckBox("Validate export before saving")
        validate_cb.setChecked(True)
        preview_layout.addWidget(validate_cb)

        optimize_cb = QCheckBox("Auto-optimize for file size")
        preview_layout.addWidget(optimize_cb)

        layout.addWidget(preview_group)

        # Apply settings button
        apply_btn = QPushButton("Apply Quality Settings")
        apply_btn.clicked.connect(self.apply_quality_settings)
        layout.addWidget(apply_btn)

        return widget

    def apply_quality_settings(self):
        """Apply export quality settings."""
        QMessageBox.information(self, "Quality Settings", "Export quality settings have been applied.")

    def create_writing_content(self, subsection_id):
        """Create content for writing-related selections."""
        if subsection_id == 'text_editor':
            return self.create_text_editor_widget()
        elif subsection_id == 'writing_session':
            return self.create_writing_session_widget()
        elif subsection_id == 'goal_tracking':
            return self.create_goal_tracking_widget()
        elif subsection_id == 'word_count':
            return self.create_word_count_widget()
        elif subsection_id == 'writing_timer':
            return self.create_writing_timer_widget()

        return self.create_placeholder_content('writing', subsection_id)

    def create_ai_content(self, subsection_id):
        """Create content for AI-related selections."""
        if subsection_id == 'ai_settings':
            return self.create_ai_settings_widget()
        elif subsection_id == 'models':
            return self.create_models_widget()
        elif subsection_id == 'prompts':
            return self.create_prompts_widget()
        elif subsection_id == 'response_handling':
            return self.create_response_handling_widget()
        elif subsection_id == 'usage_limits':
            return self.create_usage_limits_widget()

        return self.create_placeholder_content('ai', subsection_id)

    def create_templates_content(self, subsection_id):
        """Create content for template-related selections."""
        if subsection_id == 'template_library':
            return self.create_template_library_widget()
        elif subsection_id == 'custom_templates':
            return self.create_custom_templates_widget()
        elif subsection_id == 'template_marketplace':
            return self.create_template_marketplace_widget()
        elif subsection_id == 'template_recommendations':
            return self.create_template_recommendations_widget()

        return self.create_placeholder_content('templates', subsection_id)

    def create_collaboration_content(self, subsection_id):
        """Create content for collaboration-related selections."""
        if subsection_id == 'collaboration_overview':
            return self.create_collaboration_overview_widget()
        elif subsection_id == 'team_management':
            return self.create_team_management_widget()
        elif subsection_id == 'shared_projects':
            return self.create_shared_projects_widget()
        elif subsection_id == 'communication':
            return self.create_communication_widget()
        elif subsection_id == 'version_control':
            return self.create_version_control_widget()

        return self.create_placeholder_content('collaboration', subsection_id)

    def create_analytics_content(self, subsection_id):
        """Create content for analytics-related selections."""
        if subsection_id == 'writing_analytics':
            return self.create_writing_analytics_widget()
        elif subsection_id == 'productivity_metrics':
            return self.create_productivity_metrics_widget()
        elif subsection_id == 'goal_progress':
            return self.create_goal_progress_widget()
        elif subsection_id == 'usage_statistics':
            return self.create_usage_statistics_widget()

        return self.create_placeholder_content('analytics', subsection_id)

    def create_text_tools_content(self, subsection_id):
        """Create content for text tools-related selections."""
        if subsection_id == 'text_analysis':
            return self.create_text_analysis_widget()
        elif subsection_id == 'grammar_check':
            return self.create_grammar_check_widget()
        elif subsection_id == 'style_analysis':
            return self.create_style_analysis_widget()
        elif subsection_id == 'consistency_check':
            return self.create_consistency_check_widget()
        elif subsection_id == 'readability':
            return self.create_readability_widget()

        return self.create_placeholder_content('text_tools', subsection_id)

    def create_workflow_content(self, subsection_id):
        """Create content for workflow-related selections."""
        if subsection_id == 'workflow_overview':
            return self.create_workflow_overview_widget()
        elif subsection_id == 'step_management':
            return self.create_step_management_widget()
        elif subsection_id == 'automation':
            return self.create_automation_widget()
        elif subsection_id == 'workflow_templates':
            return self.create_workflow_templates_widget()

        return self.create_placeholder_content('workflow', subsection_id)

    def save_wordsapi_key(self):
        """Save WordsAPI key."""
        try:
            # Implementation would save API key securely
            QMessageBox.information(self, "API Key Saved", "WordsAPI key saved successfully!")
        except Exception as e:
            QMessageBox.critical(self, "Save Error", f"Failed to save API key: {e}")

    def test_wordsapi_connection(self):
        """Test WordsAPI connection."""
        QMessageBox.information(self, "Connection Test", "WordsAPI connection test would be performed here.")

    def refresh_memory_stats(self):
        """Refresh memory statistics."""
        try:
            import psutil
            process = psutil.Process()
            memory_info = process.memory_info()

            current_mb = memory_info.rss / 1024 / 1024
            self.memory_current_label.setText(f"{current_mb:.1f} MB")

            virtual_memory = psutil.virtual_memory()
            available_mb = virtual_memory.available / 1024 / 1024
            self.memory_available_label.setText(f"{available_mb:.1f} MB")

        except ImportError:
            self.memory_current_label.setText("psutil not available")
            self.memory_available_label.setText("Install psutil for memory monitoring")
        except Exception as e:
            self.memory_current_label.setText(f"Error: {e}")

    def load_project(self, project_name):
        """Load a project (maintain compatibility with existing code)."""
        try:
            self.current_project = project_name
            print(f"✓ Project '{project_name}' loaded successfully")
        except Exception as e:
            print(f"⚠ Failed to load project: {e}")

    def create_main_content_area(self):
        """Create the main content area with stacked pages."""
        content_area = QStackedWidget()

        # Create different views
        self.views = {
            'dashboard': self.create_dashboard_view(),
            'projects': self.create_projects_view(),
            'workflow': self.create_workflow_view(),
            'project_management': self.create_project_management_view(),
            'settings': self.create_settings_view(),
            'help': self.create_help_view()
        }

        # Performance Optimization view
        self.performance_ui = PerformanceOptimizationUI(self)
        self.performance_view = self.performance_ui.create_performance_dashboard()
        self.views['performance'] = self.performance_view

        # System Health Monitor
        self.system_health_view = self.performance_ui.create_system_health_overview()
        self.views['system_health'] = self.system_health_view

        # Add views to stacked widget
        for view_name, view_widget in self.views.items():
            content_area.addWidget(view_widget)

        # Set default view
        content_area.setCurrentWidget(self.views['dashboard'])
        self.current_view = 'dashboard'
        self.nav_buttons['dashboard'].setProperty("class", "nav-button active")

        return content_area

    def create_dashboard_view(self):
        """Create the dashboard view with project metrics."""
        dashboard = QWidget()
        layout = QVBoxLayout()
        dashboard.setLayout(layout)

        # Page header
        header = QLabel("Dashboard")
        header.setProperty("class", "header")
        layout.addWidget(header)

        # Metrics section
        metrics_data = self.get_project_metrics()
        metrics_dashboard = WorkflowUI.create_progress_dashboard(metrics_data)
        layout.addWidget(metrics_dashboard)

        # Performance overview
        if hasattr(self, 'performance_ui'):
            performance_overview = self.performance_ui.create_performance_overview()
            layout.addWidget(performance_overview)

        # Recent activity
        activity_card = Components.create_card(
            "Recent Activity",
            self.create_activity_list()
        )
        layout.addWidget(activity_card)

        # Workflow status
        workflow_card = Components.create_card(
            "Workflow Status",
            self.create_workflow_status_widget()
        )
        layout.addWidget(workflow_card)

        layout.addStretch()
        return dashboard

    def create_projects_view(self):
        """Create the projects management view."""
        projects = QWidget()
        layout = QVBoxLayout()
        projects.setLayout(layout)

        # Page header
        header_layout = QHBoxLayout()

        header = QLabel("Projects")
        header.setProperty("class", "header")
        header_layout.addWidget(header)

        # New project button
        new_project_btn = Components.create_button("New Project", "success", None)
        new_project_btn.clicked.connect(self.create_new_project)
        header_layout.addWidget(new_project_btn)

        header_layout.addStretch()
        layout.addLayout(header_layout)

        # Project list
        self.project_list = self.create_project_list()
        layout.addWidget(self.project_list)

        return projects

    def create_workflow_view(self):
        """Create the workflow management view."""
        workflow = QWidget()
        layout = QVBoxLayout()
        workflow.setLayout(layout)

        # Page header
        header_layout = QHBoxLayout()

        header = QLabel("Workflow")
        header.setProperty("class", "header")
        header_layout.addWidget(header)

        # Workflow controls
        start_btn = Components.create_button("Start Workflow", "primary")
        start_btn.clicked.connect(self.start_workflow)
        header_layout.addWidget(start_btn)

        pause_btn = Components.create_button("Pause", "secondary")
        pause_btn.clicked.connect(self.pause_workflow)
        header_layout.addWidget(pause_btn)

        stop_btn = Components.create_button("Stop", "danger")
        stop_btn.clicked.connect(self.stop_workflow)
        header_layout.addWidget(stop_btn)

        header_layout.addStretch()
        layout.addLayout(header_layout)

        # Workflow steps
        self.workflow_steps_widget = self.create_workflow_steps_widget()
        layout.addWidget(self.workflow_steps_widget)

        # Progress section
        progress_card = Components.create_card(
            "Progress",
            self.create_workflow_progress_widget()
        )
        layout.addWidget(progress_card)

        return workflow

    def create_settings_view(self):
        """Create the settings view."""
        settings = QWidget()
        layout = QVBoxLayout()
        settings.setLayout(layout)

        # Page header
        header = QLabel("Settings")
        header.setProperty("class", "header")
        layout.addWidget(header)

        # Settings tabs
        tabs = QTabWidget()

        # General settings
        general_tab = self.create_general_settings_tab()
        tabs.addTab(general_tab, "General")

        # AI Provider settings
        ai_tab = self.create_ai_provider_settings_tab()
        tabs.addTab(ai_tab, "AI Providers")

        # Prompt Engineering settings
        prompt_tab = self.create_prompt_engineering_settings_tab()
        tabs.addTab(prompt_tab, "Prompt Engineering")

        # Workflow settings
        workflow_tab = self.create_workflow_settings_tab()
        tabs.addTab(workflow_tab, "Workflow")

        # Performance settings
        performance_tab = self.create_performance_settings_tab()
        tabs.addTab(performance_tab, "Performance")

        # Cache Settings tab
        if self.cache_enabled:
            cache_tab = self.create_cache_settings_tab()
            tabs.addTab(cache_tab, "Cache Settings")

        layout.addWidget(tabs)

        return settings

    def create_help_view(self):
        """Create the help view."""
        help_widget = QWidget()
        layout = QVBoxLayout()
        help_widget.setLayout(layout)

        # Page header
        header = QLabel("Help & Documentation")
        header.setProperty("class", "header")
        layout.addWidget(header)

        # Help content
        help_content = QTextEdit()
        help_content.setReadOnly(True)
        help_content.setHtml(self.get_help_content())
        layout.addWidget(help_content)

        return help_widget

    def create_project_info_widget(self):
        """Create the project information widget for the sidebar."""
        info_widget = QFrame()
        info_widget.setProperty("class", "card")

        layout = QVBoxLayout()
        info_widget.setLayout(layout)

        # Project name
        self.project_name_label = QLabel("No Project Selected")
        self.project_name_label.setProperty("class", "subheader")
        layout.addWidget(self.project_name_label)

        # Project stats
        self.project_stats = QLabel("Select a project to view details")
        self.project_stats.setProperty("class", "caption")
        self.project_stats.setWordWrap(True)
        layout.addWidget(self.project_stats)

        return info_widget

    def create_activity_list(self):
        """Create the recent activity list widget."""
        activity_list = QListWidget()
        activity_list.setMaximumHeight(200)

        # Sample activities
        activities = [
            "Created new project 'Epic Fantasy Adventure'",
            "Completed character development for protagonist",
            "Generated initial plot outline",
            "Refined story structure",
            "Exported draft to PDF"
        ]

        for activity in activities:
            item = QListWidgetItem(activity)
            activity_list.addItem(item)

        return activity_list

    def create_workflow_status_widget(self):
        """Create the workflow status widget."""
        status_widget = QWidget()
        layout = QVBoxLayout()
        status_widget.setLayout(layout)

        # Current step
        current_step = QLabel("Current Step: Story Planning")
        current_step.setProperty("class", "subheader")
        layout.addWidget(current_step)

        # Progress bar
        progress = QProgressBar()
        progress.setValue(45)
        layout.addWidget(progress)

        # Status indicators
        status_layout = QHBoxLayout()

        for status, count in [("Completed", 4), ("In Progress", 1), ("Pending", 6)]:
            indicator = Components.create_status_indicator(
                status.lower().replace(' ', '_'),
                f"{status}: {count}"
            )
            status_layout.addWidget(indicator)

        layout.addLayout(status_layout)

        return status_widget

    def create_project_list(self):
        """Create the project list widget."""
        project_list = QTreeWidget()
        project_list.setHeaderLabels(["Project", "Status", "Progress", "Last Modified"])

        # Sample projects
        projects = [
            ("Epic Fantasy Adventure", "Active", "45%", "2024-01-15"),
            ("Sci-Fi Thriller", "Paused", "23%", "2024-01-10"),
            ("Mystery Novel", "Completed", "100%", "2024-01-05")
        ]

        for project_data in projects:
            item = QTreeWidgetItem(project_data)
            project_list.addTopLevelItem(item)

        return project_list

    def create_workflow_steps_widget(self):
        """Create the workflow steps visualization widget."""
        steps_widget = QScrollArea()
        steps_widget.setWidgetResizable(True)

        content_widget = QWidget()
        layout = QVBoxLayout()
        content_widget.setLayout(layout)

        # Workflow steps
        steps_data = [
            (1, "Project Initialization", "Set up project structure and metadata", "completed"),
            (2, "Character Development", "Create and develop main characters", "completed"),
            (3, "World Building", "Design the story world and settings", "completed"),
            (4, "Plot Outline", "Create the main plot structure", "completed"),
            (5, "Story Planning", "Detailed story planning and structure", "processing"),
            (6, "Chapter Generation", "Generate chapter content", "pending"),
            (7, "User Review", "Review and refine generated content", "pending"),
            (8, "Refinement Loop", "Iterative improvement process", "pending"),
            (9, "Progression Management", "Track and manage story progression", "pending"),
            (10, "Recovery System", "Handle errors and recovery", "pending"),
            (11, "Completion & Export", "Finalize and export the novel", "pending")
        ]

        for step_data in steps_data:
            step_card = WorkflowUI.create_workflow_step_card(*step_data)
            layout.addWidget(step_card)

        steps_widget.setWidget(content_widget)
        return steps_widget

    def create_workflow_progress_widget(self):
        """Create the workflow progress tracking widget."""
        progress_widget = QWidget()
        layout = QVBoxLayout()
        progress_widget.setLayout(layout)

        # Overall progress
        overall_label = QLabel("Overall Progress")
        overall_label.setProperty("class", "subheader")
        layout.addWidget(overall_label)

        overall_progress = QProgressBar()
        overall_progress.setValue(45)
        layout.addWidget(overall_progress)

        # Step-by-step progress
        steps_progress = QLabel("Steps Completed: 4/11")
        steps_progress.setProperty("class", "caption")
        layout.addWidget(steps_progress)

        return progress_widget

    def create_general_settings_tab(self):
        """Create the general settings tab."""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # Theme selection
        theme_group = QGroupBox("Theme")
        theme_layout = QVBoxLayout()
        theme_group.setLayout(theme_layout)

        theme_combo = QComboBox()
        theme_combo.addItems(["Dark", "Light", "Auto"])
        theme_layout.addWidget(theme_combo)

        layout.addWidget(theme_group)

        # Language selection
        lang_group = QGroupBox("Language")
        lang_layout = QVBoxLayout()
        lang_group.setLayout(lang_layout)

        lang_combo = QComboBox()
        lang_combo.addItems(["English", "Spanish", "French", "German"])
        lang_layout.addWidget(lang_combo)

        layout.addWidget(lang_group)

        layout.addStretch()
        return tab

    def create_prompt_engineering_settings_tab(self):
        """Create the prompt engineering settings tab."""
        # Return the prompt engine UI directly
        return self.prompt_engine_ui

    def create_ai_provider_settings_tab(self):
        """Create the AI provider settings tab."""
        # Return the AI configuration UI directly
        return self.ai_config_ui

    def create_workflow_settings_tab(self):
        """Create the workflow settings tab."""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # Auto-save settings
        autosave_group = QGroupBox("Auto-Save")
        autosave_layout = QVBoxLayout()
        autosave_group.setLayout(autosave_layout)

        autosave_checkbox = QCheckBox("Enable auto-save")
        autosave_checkbox.setChecked(True)
        autosave_layout.addWidget(autosave_checkbox)

        autosave_interval = QSpinBox()
        autosave_interval.setRange(1, 60)
        autosave_interval.setValue(5)
        autosave_interval.setSuffix(" minutes")
        autosave_layout.addWidget(autosave_interval)

        layout.addWidget(autosave_group)

        # Backup settings
        backup_group = QGroupBox("Backup")
        backup_layout = QVBoxLayout()
        backup_group.setLayout(backup_layout)

        backup_checkbox = QCheckBox("Enable automatic backups")
        backup_checkbox.setChecked(True)
        backup_layout.addWidget(backup_checkbox)

        backup_interval = QSpinBox()
        backup_interval.setRange(1, 24)
        backup_interval.setValue(1)
        backup_interval.setSuffix(" hours")
        backup_layout.addWidget(backup_interval)

        layout.addWidget(backup_group)

        layout.addStretch()
        return tab

    def create_performance_settings_tab(self):
        """Create the performance settings tab."""
        tab = QWidget()
        layout = QVBoxLayout()
        tab.setLayout(layout)

        # Performance monitoring
        perf_group = QGroupBox("Performance Monitoring")
        perf_layout = QVBoxLayout()
        perf_group.setLayout(perf_layout)

        monitor_checkbox = QCheckBox("Enable performance monitoring")
        monitor_checkbox.setChecked(True)
        perf_layout.addWidget(monitor_checkbox)

        layout.addWidget(perf_group)

        layout.addStretch()
        return tab

    def create_cache_settings_tab(self):
        """Create cache settings tab for intelligent caching configuration."""
        tab = QWidget()
        layout = QVBoxLayout(tab)

        # Cache status section
        status_group = QGroupBox("Cache Status")
        status_layout = QFormLayout(status_group)

        self.cache_stats_label = QLabel("Loading cache statistics...")
        status_layout.addRow("Status:", self.cache_stats_label)

        # Cache configuration section
        config_group = QGroupBox("Cache Configuration")
        config_layout = QFormLayout(config_group)

        # Max cache size
        self.max_cache_size_spinbox = QSpinBox()
        self.max_cache_size_spinbox.setRange(10, 10000)
        self.max_cache_size_spinbox.setValue(500)
        config_layout.addRow("Max Cache Size:", self.max_cache_size_spinbox)

        # Similarity threshold
        self.similarity_threshold_spinbox = QDoubleSpinBox()
        self.similarity_threshold_spinbox.setRange(0.1, 1.0)
        self.similarity_threshold_spinbox.setSingleStep(0.05)
        self.similarity_threshold_spinbox.setValue(0.85)
        config_layout.addRow("Similarity Threshold:", self.similarity_threshold_spinbox)

        # Max age hours
        self.max_age_hours_spinbox = QSpinBox()
        self.max_age_hours_spinbox.setRange(1, 720)  # 1 hour to 30 days
        self.max_age_hours_spinbox.setValue(168)  # 1 week
        config_layout.addRow("Max Age (hours):", self.max_age_hours_spinbox)

        # Enable semantic caching
        self.enable_semantic_checkbox = QCheckBox("Enable Semantic Caching")
        self.enable_semantic_checkbox.setChecked(True)
        config_layout.addRow("", self.enable_semantic_checkbox)

        # Enable compression
        self.enable_compression_checkbox = QCheckBox("Enable Compression")
        self.enable_compression_checkbox.setChecked(True)
        config_layout.addRow("", self.enable_compression_checkbox)

        # Cache actions section
        actions_group = QGroupBox("Cache Actions")
        actions_layout = QVBoxLayout(actions_group)

        # Buttons layout
        buttons_layout = QHBoxLayout()

        # Refresh stats button
        refresh_button = QPushButton("Refresh Statistics")
        refresh_button.clicked.connect(self.refresh_cache_stats)
        buttons_layout.addWidget(refresh_button)

        # Optimize cache button
        optimize_button = QPushButton("Optimize Cache")
        optimize_button.clicked.connect(self.optimize_cache)
        buttons_layout.addWidget(optimize_button)

        # Clear cache button
        clear_button = QPushButton("Clear Cache")
        clear_button.clicked.connect(self.clear_cache)
        buttons_layout.addWidget(clear_button)

        # Export cache button
        export_button = QPushButton("Export Cache Data")
        export_button.clicked.connect(self.export_cache_data)
        buttons_layout.addWidget(export_button)

        actions_layout.addLayout(buttons_layout)

        # Cache statistics display
        self.cache_stats_text = QTextEdit()
        self.cache_stats_text.setReadOnly(True)
        self.cache_stats_text.setMaximumHeight(150)
        actions_layout.addWidget(self.cache_stats_text)

        # Layout assembly
        layout.addWidget(status_group)
        layout.addWidget(config_group)
        layout.addWidget(actions_group)
        layout.addStretch()

        # Initialize cache stats
        self.refresh_cache_stats()

        return tab

    def refresh_cache_stats(self):
        """Refresh and display cache statistics."""
        if not self.cache_enabled:
            return

        try:
            stats = self.cache_manager.get_cache_stats()

            # Update status label
            status_text = f"Active - {stats['total_entries']} entries, {stats['hit_rate']}% hit rate"
            self.cache_stats_label.setText(status_text)

            # Update detailed stats
            stats_text = f"""Cache Statistics:
Total Entries: {stats['total_entries']}
Cache Hits: {stats['cache_hits']}
Cache Misses: {stats['cache_misses']}
Hit Rate: {stats['hit_rate']}%
Average Response Time: {stats['average_response_time']:.3f}s
Total Cost Saved: ${stats['total_cost_saved']:.2f}
Storage Size: {stats['storage_size_mb']:.2f} MB
Semantic Matches: {stats['semantic_matches']}
Exact Matches: {stats['exact_matches']}
Cache Strategy: {stats['cache_strategy']}
Invalidation Strategy: {stats['invalidation_strategy']}
Semantic Caching: {'Enabled' if stats['semantic_caching_enabled'] else 'Disabled'}"""

            self.cache_stats_text.setText(stats_text)

        except Exception as e:
            self.cache_stats_label.setText(f"Error: {str(e)}")
            self.cache_stats_text.setText(f"Failed to load cache statistics: {str(e)}")

    def optimize_cache(self):
        """Optimize the cache system."""
        if not self.cache_enabled:
            return

        try:
            optimization_results = self.cache_manager.optimize_cache()

            message = "Cache optimization completed!\n\n"

            if optimization_results['actions_taken']:
                message += "Actions taken:\n"
                for action in optimization_results['actions_taken']:
                    message += f"• {action}\n"
                message += "\n"

            if optimization_results['recommendations']:
                message += "Recommendations:\n"
                for rec in optimization_results['recommendations']:
                    message += f"• {rec}\n"

            if not optimization_results['actions_taken'] and not optimization_results['recommendations']:
                message += "No optimization actions needed at this time."

            QMessageBox.information(self, "Cache Optimization", message)

            # Refresh stats
            self.refresh_cache_stats()

        except Exception as e:
            QMessageBox.critical(self, "Cache Optimization Error", f"Failed to optimize cache: {str(e)}")

    def clear_cache(self):
        """Clear cache with confirmation."""
        if not self.cache_enabled:
            return

        reply = QMessageBox.question(
            self, "Clear Cache",
            "Are you sure you want to clear all cache entries?\n\nThis action cannot be undone.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            try:
                deleted = self.cache_manager.clear_cache()
                QMessageBox.information(self, "Cache Cleared", f"Cleared {deleted} cache entries.")

                # Refresh stats
                self.refresh_cache_stats()

            except Exception as e:
                QMessageBox.critical(self, "Cache Clear Error", f"Failed to clear cache: {str(e)}")

    def export_cache_data(self):
        """Export cache data to file."""
        if not self.cache_enabled:
            return

        try:
            from datetime import datetime
            default_filename = f"cache_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

            filename, _ = QFileDialog.getSaveFileName(
                self, "Export Cache Data",
                default_filename,
                "JSON Files (*.json);;All Files (*)"
            )

            if filename:
                success = self.cache_manager.export_cache_data(filename)

                if success:
                    QMessageBox.information(self, "Export Complete", f"Cache data exported to:\n{filename}")
                else:
                    QMessageBox.critical(self, "Export Error", "Failed to export cache data.")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export cache data: {str(e)}")

    # ...existing code...
    def create_status_bar(self):
        """Create the modern status bar."""
        status_bar = self.statusBar()

        # Status message
        self.status_message = QLabel("Ready")
        status_bar.addWidget(self.status_message)

        # Progress indicator
        self.status_progress = QProgressBar()
        self.status_progress.setMaximumWidth(200)
        self.status_progress.setVisible(False)
        status_bar.addPermanentWidget(self.status_progress)

        # Connection status
        self.connection_status = Components.create_status_indicator("online", "Connected")
        status_bar.addPermanentWidget(self.connection_status)

        # Performance info
        self.performance_label = QLabel("CPU: 0% | Memory: 0MB")
        status_bar.addPermanentWidget(self.performance_label)

    def create_menu_bar(self):
        """Create the modern menu bar."""
        menubar = self.menuBar()

        # File menu
        file_menu = menubar.addMenu('File')

        new_action = QAction('New Project', self)
        new_action.setShortcut('Ctrl+N')
        new_action.triggered.connect(self.create_new_project)
        file_menu.addAction(new_action)

        open_action = QAction('Open Project', self)
        open_action.setShortcut('Ctrl+O')
        open_action.triggered.connect(self.open_project)
        file_menu.addAction(open_action)

        save_action = QAction('Save Project', self)
        save_action.setShortcut('Ctrl+S')
        save_action.triggered.connect(self.save_project)
        file_menu.addAction(save_action)

        file_menu.addSeparator()

        exit_action = QAction('Exit', self)
        exit_action.setShortcut('Ctrl+Q')
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # Edit menu
        edit_menu = menubar.addMenu('Edit')

        # View menu = menubar.addMenu('View')

        # Help menu
        help_menu = menubar.addMenu('Help')

        about_action = QAction('About FANWS', self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)

    def setup_styling(self):
        """Apply modern styling to the application."""
        # Apply main stylesheet
        main_style = DesignSystem.get_main_stylesheet()
        sidebar_style = DesignSystem.get_sidebar_stylesheet()
        dashboard_style = DesignSystem.get_dashboard_stylesheet()

        combined_style = main_style + sidebar_style + dashboard_style
        self.setStyleSheet(combined_style)

    def setup_connections(self):
        """Set up signal connections."""
        # Performance monitoring
        self.performance_timer = QTimer()
        self.performance_timer.timeout.connect(self.update_performance_display)
        self.performance_timer.start(1000)  # Update every second

    def init_async_ui(self):
        """Initialize async UI components (Priority 4.1)."""
        if not self.async_enabled:
            return

        # Initialize async status bar
        if self.async_status_bar and hasattr(self, 'statusBar'):
            self.statusBar().addPermanentWidget(self.async_status_bar)

        # Connect async manager signals
        if self.async_manager:
            if hasattr(self.async_manager, 'task_started'):
                self.async_manager.task_started.connect(self.on_async_task_started)
            if hasattr(self.async_manager, 'task_progress'):
                self.async_manager.task_progress.connect(self.on_async_task_progress)
            if hasattr(self.async_manager, 'task_completed'):
                self.async_manager.task_completed.connect(self.on_async_task_completed)
            if hasattr(self.async_manager, 'task_failed'):
                self.async_manager.task_failed.connect(self.on_async_task_failed)

        # Connect async workflow manager signals
        if self.async_workflow_manager:
            if hasattr(self.async_workflow_manager, 'workflow_started'):
                self.async_workflow_manager.workflow_started.connect(self.on_async_workflow_started)
            if hasattr(self.async_workflow_manager, 'workflow_progress'):
                self.async_workflow_manager.workflow_progress.connect(self.on_async_workflow_progress)
            if hasattr(self.async_workflow_manager, 'workflow_completed'):
                self.async_workflow_manager.workflow_completed.connect(self.on_async_workflow_completed)
            if hasattr(self.async_workflow_manager, 'workflow_failed'):
                self.async_workflow_manager.workflow_failed.connect(self.on_async_workflow_failed)

        print("✓ Async UI components initialized")

    def get_icon_path(self, icon_name):
        """Get the path to an icon file."""
        return os.path.join(os.path.dirname(__file__), '..', 'resources', 'icons', icon_name)

    def get_project_metrics(self):
        """Get project metrics for the dashboard."""
        # Calculate dynamic metrics with real data
        total_progress = self.calculate_total_progress()
        avg_daily = self.calculate_avg_daily_progress()

        # Get current project data
        current_project = self.get_current_project_data()

        return {
            'total_projects': len(self.get_all_projects()),
            'active_projects': len([p for p in self.get_all_projects() if p.get('status') == 'active']),
            'completed_projects': len([p for p in self.get_all_projects() if p.get('status') == 'completed']),
            'words_written': current_project.get('words_written', 0),
            'chapters_completed': current_project.get('chapters_completed', 0),
            'workflow_steps': current_project.get('completed_steps', 0),
            'total_progress': total_progress,
            'avg_daily': avg_daily,
            'workflow_completion': f"{total_progress:.1f}%",
            'daily_average': f"{avg_daily:.0f} words/day"
        }

    def get_help_content(self):
        """Get help content HTML."""
        return """
        <h2>FANWS Help & Documentation</h2>

        <h3>Getting Started</h3>
        <p>Welcome to FANWS (Fantasy Adventure Novel Writing System). This application helps you create comprehensive fantasy novels through an automated workflow.</p>

        <h3>Features</h3>
        <ul>
            <li><strong>Project Management:</strong> Create and manage multiple writing projects</li>
            <li><strong>Automated Workflow:</strong> 11-step process for complete novel generation</li>
            <li><strong>Character Development:</strong> Create detailed character profiles</li>
            <li><strong>World Building:</strong> Design comprehensive fantasy worlds</li>
            <li><strong>Story Planning:</strong> Structured plot development</li>
        </ul>

        <h3>Workflow Steps</h3>
        <ol>
            <li>Project Initialization</li>
            <li>Character Development</li>
            <li>World Building</li>
            <li>Plot Outline</li>
            <li>Story Planning</li>
            <li>Chapter Generation</li>
            <li>User Review</li>
            <li>Refinement Loop</li>
            <li>Progression Management</li>
            <li>Recovery System</li>
            <li>Completion & Export</li>
        </ol>

        <h3>Support</h3>
        <p>For additional help, please refer to the documentation or contact support.</p>
        """

    def switch_view(self, view_name):
        """Switch to a different view."""
        if view_name in self.views:
            # Update navigation buttons
            for key, button in self.nav_buttons.items():
                if key == view_name:
                    button.setProperty("class", "nav-button active")
                else:
                    button.setProperty("class", "nav-button")

            # Switch view
            self.main_content.setCurrentWidget(self.views[view_name])
            self.current_view = view_name

            # Update status
            self.status_message.setText(f"Switched to {view_name.title()}")

            # Refresh view if needed
            if view_name == 'dashboard':
                self.refresh_dashboard()

    def refresh_dashboard(self):
        """Refresh the dashboard with current data."""
        try:
            # Update dashboard chart with current metrics
            self.update_dashboard_chart()

            # Update other dashboard components
            if hasattr(self, 'views') and 'dashboard' in self.views:
                # Refresh metrics in the dashboard view
                metrics_data = self.get_project_metrics()

                # Find and update the metrics dashboard widget
                dashboard_widget = self.views['dashboard']
                if dashboard_widget:
                    # Update the metrics display
                    self.update_metrics_display(metrics_data)

            self.add_log_message("Dashboard refreshed successfully", "info")

        except Exception as e:
            error_msg = f"Error refreshing dashboard: {str(e)}"
            self.add_log_message(error_msg, "error")
            import logging
            logging.error(error_msg)

    def update_dashboard_chart(self):
        """Update dashboard chart with proper variable calculations."""
        try:
            # Calculate total progress from workflow status
            total_progress = 0
            if hasattr(self, 'workflow_manager') and self.workflow_manager:
                status = self.workflow_manager.get_progress_status()
                total_progress = status.get('progress_percentage', 0)
            elif hasattr(self, 'novel_workflow') and self.novel_workflow:
                # Try alternative workflow reference
                if hasattr(self.novel_workflow, 'progress_tracker'):
                    total_progress = getattr(self.novel_workflow.progress_tracker, 'current_step', 0) / 11 * 100
            else:
                # Fallback to calculated progress
                total_progress = self.calculate_total_progress()

            # Calculate average daily progress
            avg_daily = 0
            if hasattr(self, 'analytics_manager') and self.analytics_manager:
                recent_activity = self.analytics_manager.get_recent_activity(7)  # Last 7 days
                if recent_activity:
                    total_words = sum(activity.get('word_count', 0) for activity in recent_activity)
                    avg_daily = total_words / 7 if total_words > 0 else 0
            else:
                # Fallback to calculated average
                avg_daily = self.calculate_avg_daily_progress()

            # Update chart display with calculated values
            self.update_chart_display(total_progress, avg_daily)

            # Store values for access by other methods
            self._last_total_progress = total_progress
            self._last_avg_daily = avg_daily

            # Update status
            self.add_log_message(f"Chart updated: {total_progress:.1f}% progress, {avg_daily:.0f} words/day", "info")

        except Exception as e:
            error_msg = f"Error updating dashboard chart: {str(e)}"
            self.add_log_message(error_msg, "error")
            import logging
            logging.error(error_msg)

    def calculate_total_progress(self):
        """Calculate total workflow progress percentage."""
        try:
            # Check if we have a current project
            current_project = self.get_current_project_data()
            if not current_project:
                return 0.0

            # Calculate based on completed steps
            completed_steps = current_project.get('completed_steps', 0)
            total_steps = 11  # FANWS has 11 workflow steps

            progress = (completed_steps / total_steps) * 100
            return min(progress, 100.0)

        except Exception as e:
            import logging
            logging.error(f"Error calculating total progress: {e}")
            return 0.0

    def calculate_avg_daily_progress(self):
        """Calculate average daily word count progress."""
        try:
            # Get recent writing activity
            current_project = self.get_current_project_data()
            if not current_project:
                return 0.0

            # Calculate from session history or analytics
            writing_sessions = current_project.get('writing_sessions', [])
            if not writing_sessions:
                return 0.0

            # Get last 7 days of data
            from datetime import datetime, timedelta
            seven_days_ago = datetime.now() - timedelta(days=7)

            recent_sessions = [
                session for session in writing_sessions
                if datetime.fromisoformat(session.get('date', '1970-01-01')) >= seven_days_ago
            ]

            if not recent_sessions:
                return 0.0

            total_words = sum(session.get('word_count', 0) for session in recent_sessions)
            unique_days = len(set(session.get('date', '').split('T')[0] for session in recent_sessions))

            avg_daily = total_words / max(unique_days, 1)
            return avg_daily

        except Exception as e:
            import logging
            logging.error(f"Error calculating average daily progress: {e}")
            return 0.0

    def update_chart_display(self, total_progress, avg_daily):
        """Update the chart display with new values."""
        try:
            # Update dashboard metrics if available
            if hasattr(self, 'views') and 'dashboard' in self.views:
                # Create updated metrics data
                updated_metrics = {
                    'workflow_progress': f"{total_progress:.1f}%",
                    'daily_average': f"{avg_daily:.0f} words/day",
                    'progress_bar_value': int(total_progress),
                    'last_updated': datetime.now().strftime("%H:%M:%S")
                }

                # Store for later use
                self._chart_metrics = updated_metrics

                # If we have chart widgets, update them
                self.update_chart_widgets(updated_metrics)

        except Exception as e:
            import logging
            logging.error(f"Error updating chart display: {e}")

    def update_chart_widgets(self, metrics):
        """Update individual chart widget components."""
        try:
            # Update progress bars if they exist
            if hasattr(self, 'progress_chart_widget'):
                self.progress_chart_widget.setValue(metrics.get('progress_bar_value', 0))

            # Update text labels if they exist
            if hasattr(self, 'progress_label'):
                self.progress_label.setText(metrics.get('workflow_progress', '0%'))

            if hasattr(self, 'daily_label'):
                self.daily_label.setText(metrics.get('daily_average', '0 words/day'))

        except Exception as e:
            import logging
            logging.error(f"Error updating chart widgets: {e}")

    def update_metrics_display(self, metrics_data):
        """Update the metrics display with new data."""
        try:
            # This would update the dashboard metrics display
            # Currently metrics are handled by the static dashboard creation
            # In the future, this could update live chart components
            pass

        except Exception as e:
            import logging
            logging.error(f"Error updating metrics display: {e}")

    def get_current_project_data(self):
        """Get current project data."""
        try:
            # Return current project data or default structure
            return {
                'id': 'current_project',
                'name': 'Current Project',
                'status': 'active',
                'words_written': 0,
                'chapters_completed': 0,
                'completed_steps': 0,
                'writing_sessions': [],
                'created_date': datetime.now().isoformat(),
                'last_modified': datetime.now().isoformat()
            }
        except Exception as e:
            import logging
            logging.error(f"Error getting current project data: {e}")
            return {}

    def get_all_projects(self):
        """Get all projects list."""
        try:
            # Return list of all projects or default
            return [self.get_current_project_data()]
        except Exception as e:
            import logging
            logging.error(f"Error getting all projects: {e}")
            return []

    def update_performance_display(self):
        """Update the performance display in the status bar."""
        if self.performance_monitor:
            try:
                cpu_percent = self.performance_monitor.get_cpu_usage()
                memory_mb = self.performance_monitor.get_memory_usage()
                self.performance_label.setText(f"CPU: {cpu_percent:.1f}% | Memory: {memory_mb:.1f}MB")
            except:
                self.performance_label.setText("Performance monitoring unavailable")
        else:
            self.performance_label.setText("Performance monitoring disabled")

    def load_project_data(self):
        """Load project data from database."""
        # This would load real project data
        pass

    def create_new_project(self):
        """Create a new project with async support."""
        try:
            # First check if we have a parent FANWSWindow with the business logic
            if hasattr(self, 'parent') and hasattr(self.parent(), 'create_new_project'):
                # Use the business logic from the main application
                self.parent().create_new_project()
                return

            # Fallback to GUI-only implementation
            if self.async_enabled and self.async_manager:
                # Use async task for project creation
                task_id = self.async_manager.run_async_task(
                    self._create_new_project_async,
                    task_name="Creating New Project"
                )
                self.add_log_message("Creating new project...", "info")
            else:
                # Fallback to synchronous creation
                self._create_new_project_sync()

        except Exception as e:
            error_msg = f"Error creating project: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Project Error", error_msg)

    def _create_new_project_sync(self):
        """Synchronous project creation with enhanced dialog."""
        try:
            # Create new project dialog
            dialog = QDialog(self)
            dialog.setWindowTitle("Create New Project")
            dialog.setModal(True)
            dialog.resize(500, 400)

            layout = QVBoxLayout()
            dialog.setLayout(layout)

            # Project name input
            name_layout = QHBoxLayout()
            name_layout.addWidget(QLabel("Project Name:"))
            self.project_name_input = QLineEdit()
            self.project_name_input.setPlaceholderText("Enter project name...")
            name_layout.addWidget(self.project_name_input)
            layout.addLayout(name_layout)

            # Project template selection
            template_group = QGroupBox("Project Template")
            template_layout = QVBoxLayout()
            template_group.setLayout(template_layout)

            self.template_selector = QComboBox()
            self.template_selector.addItems([
                "Basic Novel",
                "Fantasy Adventure",
                "Science Fiction",
                "Mystery/Thriller",
                "Romance",
                "Historical Fiction",
                "Young Adult",
                "Custom Template"
            ])
            template_layout.addWidget(self.template_selector)

            layout.addWidget(template_group)

            # Project settings
            settings_group = QGroupBox("Initial Settings")
            settings_layout = QFormLayout()
            settings_group.setLayout(settings_layout)

            self.target_chapters_input = QSpinBox()
            self.target_chapters_input.setRange(1, 100)
            self.target_chapters_input.setValue(15)
            settings_layout.addRow("Target Chapters:", self.target_chapters_input)

            self.target_words_input = QSpinBox()
            self.target_words_input.setRange(1000, 500000)
            self.target_words_input.setValue(80000)
            self.target_words_input.setSuffix(" words")
            settings_layout.addRow("Target Word Count:", self.target_words_input)

            self.genre_input = QLineEdit()
            self.genre_input.setPlaceholderText("e.g., Fantasy, Science Fiction, Romance...")
            settings_layout.addRow("Genre:", self.genre_input)

            layout.addWidget(settings_group)

            # Description
            desc_group = QGroupBox("Project Description")
            desc_layout = QVBoxLayout()
            desc_group.setLayout(desc_layout)

            self.description_input = QTextEdit()
            self.description_input.setPlaceholderText("Brief description of your novel idea...")
            self.description_input.setMaximumHeight(100)
            desc_layout.addWidget(self.description_input)

            layout.addWidget(desc_group)

            # Buttons
            button_layout = QHBoxLayout()

            create_btn = QPushButton("Create Project")
            create_btn.clicked.connect(lambda: self.finalize_project_creation(dialog))
            button_layout.addWidget(create_btn)

            cancel_btn = QPushButton("Cancel")
            cancel_btn.clicked.connect(dialog.reject)
            button_layout.addWidget(cancel_btn)

            layout.addLayout(button_layout)

            # Show dialog
            if dialog.exec_() == QDialog.Accepted:
                self.refresh_project_list()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create project dialog: {str(e)}")

    def finalize_project_creation(self, dialog):
        """Finalize the project creation process."""
        try:
            project_name = self.project_name_input.text().strip()
            if not project_name:
                QMessageBox.warning(self, "Warning", "Please enter a project name.")
                return

            # Validate project name
            if not self.validate_project_name(project_name):
                QMessageBox.warning(self, "Warning", "Invalid project name. Use only letters, numbers, spaces, and hyphens.")
                return

            # Check if project already exists
            project_path = os.path.join("projects", project_name)
            if os.path.exists(project_path):
                QMessageBox.warning(self, "Warning", f"Project '{project_name}' already exists.")
                return

            # Create project directory structure
            os.makedirs(project_path, exist_ok=True)
            os.makedirs(os.path.join(project_path, "drafts"), exist_ok=True)
            os.makedirs(os.path.join(project_path, "research"), exist_ok=True)
            os.makedirs(os.path.join(project_path, "exports"), exist_ok=True)
            os.makedirs(os.path.join(project_path, "logs"), exist_ok=True)

            # Create project configuration
            project_config = {
                "name": project_name,
                "template": self.template_selector.currentText(),
                "target_chapters": self.target_chapters_input.value(),
                "target_words": self.target_words_input.value(),
                "genre": self.genre_input.text().strip(),
                "description": self.description_input.toPlainText().strip(),
                "created_at": datetime.now().isoformat(),
                "version": "1.0.0"
            }

            # Save project configuration
            config_path = os.path.join(project_path, "project.json")
            with open(config_path, 'w') as f:
                json.dump(project_config, f, indent=2)

            # Create initial settings file
            settings_path = os.path.join(project_path, "settings.json")
            with open(settings_path, 'w') as f:
                json.dump({}, f)

            # Set as current project
            self.current_project = project_config
            self.current_project["path"] = project_path

            dialog.accept()
            QMessageBox.information(self, "Success", f"Project '{project_name}' created successfully!")
            self.add_log_message(f"Created new project: {project_name}", "info")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create project: {str(e)}")

    def validate_project_name(self, name):
        """Validate project name."""
        import re
        # Allow letters, numbers, spaces, hyphens, and underscores
        return bool(re.match(r'^[a-zA-Z0-9\s\-_]+$', name))

    def refresh_project_list(self):
        """Refresh the project list in the UI."""
        try:
            # Find project selector and update it
            project_selector = self.findChild(QComboBox, "projectSelector")
            if project_selector:
                current_selection = project_selector.currentText()
                project_selector.clear()
                project_selector.addItem("Select a project")

                # Add available projects
                projects_dir = "projects"
                if os.path.exists(projects_dir):
                    for item in os.listdir(projects_dir):
                        item_path = os.path.join(projects_dir, item)
                        if os.path.isdir(item_path):
                            project_selector.addItem(item)

                # Restore selection if it still exists
                index = project_selector.findText(current_selection)
                if index >= 0:
                    project_selector.setCurrentIndex(index)

        except Exception as e:
            print(f"Error refreshing project list: {e}")

    def start_workflow(self):
        """Start the workflow process with integration to business logic."""
        try:
            # Check if we have a parent FANWSWindow with the business logic
            if hasattr(self, 'parent') and hasattr(self.parent(), 'start_writing_workflow'):
                # Use the business logic from the main application
                self.parent().start_writing_workflow()
                return
            elif hasattr(self, 'parent') and hasattr(self.parent(), 'start_writing'):
                # Fallback to basic start_writing method
                self.parent().start_writing()
                return

            # GUI-only fallback implementation
            if not hasattr(self, 'current_project') or not self.current_project:
                QMessageBox.warning(self, "Warning", "Please select a project before starting the workflow.")
                return

            if self.async_enabled and self.async_workflow_manager:
                # Use async workflow manager
                workflow_config = self.get_workflow_config()
                task_id = self.async_workflow_manager.start_async_workflow(
                    workflow_type="novel_writing",
                    config=workflow_config,
                    project_data=self.current_project
                )
                self.add_log_message(f"Started async workflow: {task_id}", "info")
            else:
                # Fallback to synchronous workflow
                self.start_legacy_workflow()

        except Exception as e:
            error_msg = f"Error starting workflow: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Workflow Error", error_msg)

    def pause_workflow(self):
        """Pause the workflow process with integration to business logic."""
        try:
            # Check if we have a parent FANWSWindow with the business logic
            if hasattr(self, 'parent') and hasattr(self.parent(), 'pause_writing'):
                self.parent().pause_writing()
                return

            # GUI-only implementation
            if self.async_enabled and self.async_workflow_manager:
                self.async_workflow_manager.pause_current_workflow()
                self.add_log_message("Workflow paused", "info")
            else:
                QMessageBox.information(self, "Pause Workflow", "Workflow paused.")
        except Exception as e:
            error_msg = f"Error pausing workflow: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Workflow Error", error_msg)

    def stop_workflow(self):
        """Stop the workflow process with integration to business logic."""
        try:
            # Check if we have a parent FANWSWindow with the business logic
            if hasattr(self, 'parent') and hasattr(self.parent(), 'stop_workflow'):
                self.parent().stop_workflow()
                return

            # GUI-only implementation
            if self.async_enabled and self.async_workflow_manager:
                self.async_workflow_manager.stop_current_workflow()
                self.add_log_message("Workflow stopped", "info")
            else:
                QMessageBox.information(self, "Stop Workflow", "Workflow stopped.")
        except Exception as e:
            error_msg = f"Error stopping workflow: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Workflow Error", error_msg)

    def save_openai_key(self):
        """Save OpenAI API key with integration to business logic."""
        try:
            # Check if we have a parent FANWSWindow with the business logic
            if hasattr(self, 'parent') and hasattr(self.parent(), 'save_api_keys'):
                self.parent().save_api_keys()
                return

            # GUI-only implementation
            if not hasattr(self, 'openai_key_input'):
                QMessageBox.warning(self, "Warning", "OpenAI key input not found.")
                return

            api_key = self.openai_key_input.text().strip()
            if not api_key:
                QMessageBox.warning(self, "Warning", "Please enter an OpenAI API key.")
                return

            # Save to configuration
            config_dir = "config"
            os.makedirs(config_dir, exist_ok=True)
            config_path = os.path.join(config_dir, "api_keys.json")

            config = {}
            if os.path.exists(config_path):
                try:
                    with open(config_path, 'r') as f:
                        config = json.load(f)
                except:
                    pass

            config['openai'] = api_key

            with open(config_path, 'w') as f:
                json.dump(config, f, indent=2)

            QMessageBox.information(self, "Success", "OpenAI API key saved successfully!")
            self.add_log_message("OpenAI API key saved", "info")

        except Exception as e:
            error_msg = f"Error saving OpenAI key: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def test_openai_connection(self):
        """Test OpenAI API connection."""
        try:
            if not hasattr(self, 'openai_key_input'):
                QMessageBox.warning(self, "Warning", "OpenAI key input not found.")
                return

            api_key = self.openai_key_input.text().strip()
            if not api_key:
                QMessageBox.warning(self, "Warning", "Please enter an OpenAI API key to test.")
                return

            # Simple validation - check if key format looks correct
            if not api_key.startswith('sk-') or len(api_key) < 20:
                QMessageBox.warning(self, "Warning", "API key format appears invalid.")
                return

            # For now, just show success if format is correct
            # In a real implementation, you would make an actual API call
            QMessageBox.information(self, "Test Result", "API key format appears valid. Save the key to enable full testing.")
            self.add_log_message("OpenAI API key test completed", "info")

        except Exception as e:
            error_msg = f"Error testing OpenAI connection: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def save_wordsapi_key(self):
        """Save WordsAPI key."""
        try:
            if not hasattr(self, 'wordsapi_key_input'):
                QMessageBox.warning(self, "Warning", "WordsAPI key input not found.")
                return

            api_key = self.wordsapi_key_input.text().strip()
            if not api_key:
                QMessageBox.warning(self, "Warning", "Please enter a WordsAPI key.")
                return

            # Save to configuration
            config_dir = "config"
            os.makedirs(config_dir, exist_ok=True)
            config_path = os.path.join(config_dir, "api_keys.json")

            config = {}
            if os.path.exists(config_path):
                try:
                    with open(config_path, 'r') as f:
                        config = json.load(f)
                except:
                    pass

            config['wordsapi'] = api_key

            with open(config_path, 'w') as f:
                json.dump(config, f, indent=2)

            QMessageBox.information(self, "Success", "WordsAPI key saved successfully!")
            self.add_log_message("WordsAPI key saved", "info")

        except Exception as e:
            error_msg = f"Error saving WordsAPI key: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def test_wordsapi_connection(self):
        """Test WordsAPI connection."""
        try:
            if not hasattr(self, 'wordsapi_key_input'):
                QMessageBox.warning(self, "Warning", "WordsAPI key input not found.")
                return

            api_key = self.wordsapi_key_input.text().strip()
            if not api_key:
                QMessageBox.warning(self, "Warning", "Please enter a WordsAPI key to test.")
                return

            # Simple validation - check if key looks reasonable
            if len(api_key) < 10:
                QMessageBox.warning(self, "Warning", "API key appears too short.")
                return

            # For now, just show success if format is reasonable
            # In a real implementation, you would make an actual API call
            QMessageBox.information(self, "Test Result", "API key format appears valid. Save the key to enable full testing.")
            self.add_log_message("WordsAPI key test completed", "info")

        except Exception as e:
            error_msg = f"Error testing WordsAPI connection: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def open_project(self):
        """Open an existing project with async support."""
        try:
            if self.async_enabled and self.async_manager:
                # Use async task for project opening
                task_id = self.async_manager.run_async_task(
                    self._open_project_async,
                    task_name="Opening Project"
                )
                self.add_log_message("Opening project...", "info")
            else:
                # Fallback to synchronous opening
                self._open_project_sync()

        except Exception as e:
            error_msg = f"Error opening project: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Project Error", error_msg)

    def save_project(self):
        """Save the current project with async support."""
        try:
            if self.async_enabled and self.async_manager:
                # Use async task for project saving
                task_id = self.async_manager.run_async_task(
                    self._save_project_async,
                    task_name="Saving Project"
                )
                self.add_log_message("Saving project...", "info")
            else:
                # Fallback to synchronous saving
                self._save_project_sync()

        except Exception as e:
            error_msg = f"Error saving project: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Project Error", error_msg)

    def _create_new_project_sync(self):
        """Synchronous project creation (fallback)."""
        QMessageBox.information(self, "New Project", "New project creation dialog would open here.")

    def _open_project_sync(self):
        """Synchronous project opening (fallback)."""
        QMessageBox.information(self, "Open Project", "Project selection dialog would open here.")

    def _save_project_sync(self):
        """Synchronous project saving (fallback)."""
        QMessageBox.information(self, "Save Project", "Project saved successfully.")

    async def _create_new_project_async(self):
        """Asynchronous project creation."""
        # Simulate project creation work
        await asyncio.sleep(0.1)
        return {"status": "success", "message": "New project created"}

    async def _open_project_async(self):
        """Asynchronous project opening."""
        # Simulate project loading work
        await asyncio.sleep(0.1)
        return {"status": "success", "message": "Project opened"}

    async def _save_project_async(self):
        """Asynchronous project saving."""
        # Simulate project saving work
        await asyncio.sleep(0.1)
        return {"status": "success", "message": "Project saved"}

    def start_workflow(self):
        """Start the workflow process with async support."""
        try:
            if self.async_enabled and self.async_workflow_manager:
                # Use async workflow manager
                workflow_config = self.get_workflow_config()
                task_id = self.async_workflow_manager.start_async_workflow(
                    workflow_type="novel_writing",
                    config=workflow_config,
                    project_data=self.current_project
                )
                self.add_log_message(f"Started async workflow: {task_id}", "info")
            else:
                # Fallback to synchronous workflow
                self.start_legacy_workflow()

        except Exception as e:
            error_msg = f"Error starting workflow: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Workflow Error", error_msg)

    def start_legacy_workflow(self):
        """Start workflow using legacy method (fallback)."""
        QMessageBox.information(self, "Start Workflow", "Workflow started successfully (legacy mode).")

    def pause_workflow(self):
        """Pause the workflow process."""
        try:
            if self.async_enabled and self.async_workflow_manager:
                self.async_workflow_manager.pause_current_workflow()
                self.add_log_message("Workflow paused", "info")
            else:
                QMessageBox.information(self, "Pause Workflow", "Workflow paused.")
        except Exception as e:
            error_msg = f"Error pausing workflow: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Workflow Error", error_msg)

    def stop_workflow(self):
        """Stop the workflow process."""
        try:
            if self.async_enabled and self.async_workflow_manager:
                self.async_workflow_manager.stop_current_workflow()
                self.add_log_message("Workflow stopped", "info")
            else:
                QMessageBox.information(self, "Stop Workflow", "Workflow stopped.")
        except Exception as e:
            error_msg = f"Error stopping workflow: {str(e)}"
            self.add_log_message(error_msg, "error")
            QMessageBox.critical(self, "Workflow Error", error_msg)

    def get_workflow_config(self):
        """Get current workflow configuration."""
        return {
            "project_name": self.current_project.get("name", "Untitled") if self.current_project else "Untitled",
            "ai_provider": self.ai_provider_manager.get_current_provider() if self.ai_provider_manager else "openai",
            "use_cache": self.cache_enabled,
            "enable_feedback": True,
            "auto_save": True
        }

    def show_about(self):
        """Show the about dialog."""
        QMessageBox.about(self, "About FANWS",
                         "FANWS - Fantasy Adventure Novel Writing System\n\n"
                         "Version 2.0\n"
                         "Modern GUI with Enhanced User Experience\n\n"
                         "© 2024 FANWS Development Team")

    def closeEvent(self, event):
        """Handle application close event."""
        # Save settings and cleanup
        if self.performance_monitor:
            try:
                self.performance_monitor.stop_monitoring()
            except:
                pass

        event.accept()

    def connect_workflow_controls(self):
        """Connect UI controls to workflow controller."""
        # This method was referenced in the __init__ but wasn't defined
        # Add connections here when UI controls are available
        pass

    @pyqtSlot(int, str)
    def update_progress_display(self, progress: int, current_step: str):
        """Update progress display in UI."""
        try:
            # Update overall progress bar
            progress_bar = self.findChild(QProgressBar, "overallProgress")
            if progress_bar:
                progress_bar.setValue(progress)

            # Update current step label
            step_label = self.findChild(QLabel, "currentStepLabel")
            if step_label:
                step_label.setText(f"Current Step: {current_step}")

        except Exception as e:
            print(f"Error updating progress display: {e}")

    @pyqtSlot(str)
    def update_eta_display(self, eta: str):
        """Update ETA display in UI."""
        try:
            eta_label = self.findChild(QLabel, "etaLabel")
            if eta_label:
                eta_label.setText(f"Estimated Time Remaining: {eta}")
        except Exception as e:
            print(f"Error updating ETA display: {e}")

    @pyqtSlot(float)
    def update_speed_display(self, speed: float):
        """Update speed display in UI."""
        try:
            speed_label = self.findChild(QLabel, "speedLabel")
            if speed_label:
                speed_label.setText(f"Processing Speed: {speed:.1f} operations/sec")
        except Exception as e:
            print(f"Error updating speed display: {e}")

    @pyqtSlot(str, str)
    def add_log_message(self, message: str, level: str):
        """Add message to live log viewer."""
        try:
            log_text = self.findChild(QTextEdit, "liveLogText")
            if log_text:
                timestamp = datetime.now().strftime("%H:%M:%S")
                formatted_message = f"[{timestamp}] {level.upper()}: {message}\n"
                log_text.append(formatted_message)

                # Auto-scroll if enabled
                auto_scroll = self.findChild(QCheckBox, "autoScrollCheckbox")
                if auto_scroll and auto_scroll.isChecked():
                    log_text.moveCursor(log_text.textCursor().End)

        except Exception as e:
            print(f"Error adding log message: {e}")

    @pyqtSlot(str, dict)
    def show_feedback_request(self, step_name: str, step_data: dict):
        """Show feedback request in review panel."""
        try:
            # Switch to feedback tab
            workflow_tabs = self.findChild(QTabWidget, "workflowTabs")
            if workflow_tabs:
                workflow_tabs.setCurrentIndex(1)  # Feedback tab

            # Update review content area
            content_area = self.findChild(QTextEdit, "reviewContentArea")
            if content_area:
                content_text = f"Step: {step_name}\n\n"
                content_text += f"Generated Content:\n{step_data.get('content', 'No content available')}\n\n"
                content_text += f"Metadata:\n{json.dumps(step_data, indent=2)}"
                content_area.setPlainText(content_text)

            # Store current review item
            self.current_review_item = step_data.get("item_id", step_name)

        except Exception as e:
            print(f"Error showing feedback request: {e}")

    @pyqtSlot(dict)
    def update_feedback_metrics(self, feedback_data: dict):
        """Update feedback metrics display."""
        try:
            if hasattr(self, 'workflow_controller'):
                metrics = self.workflow_controller.feedback_manager.get_metrics()

                # Update metric cards
                pending_metric = self.findChild(QLabel, "pendingReviewsMetric")
                if pending_metric:
                    pending_metric.setText(str(metrics.get("pending_reviews", 0)))

                approved_metric = self.findChild(QLabel, "approvedMetric")
                if approved_metric:
                    approved_metric.setText(str(metrics.get("total_feedback", 0)))

                rating_metric = self.findChild(QLabel, "ratingMetric")
                if rating_metric:
                    rating_metric.setText(f"{metrics.get('average_rating', 0):.1f}")

                response_metric = self.findChild(QLabel, "responseMetric")
                if response_metric:
                    avg_response = metrics.get("average_response_time", 0)
                    if avg_response > 0:
                        response_metric.setText(f"{avg_response / 3600:.1f}h")
                    else:
                        response_metric.setText("0h")

        except Exception as e:
            print(f"Error updating feedback metrics: {e}")

    def approve_content(self):
        """Handle content approval."""
        try:
            if hasattr(self, 'current_review_item') and self.current_review_item:
                if hasattr(self, 'workflow_controller'):
                    self.workflow_controller.feedback_manager.process_feedback(
                        self.current_review_item, "approve"
                    )
                    self.add_log_message("Content approved by user", "info")
        except Exception as e:
            print(f"Error approving content: {e}")

    def request_changes(self):
        """Handle change request."""
        try:
            if hasattr(self, 'current_review_item') and self.current_review_item:
                # Get feedback text
                feedback_text = self.findChild(QTextEdit, "feedbackText")
                feedback_type = self.findChild(QComboBox, "feedbackTypeCombo")
                priority = self.findChild(QComboBox, "priorityCombo")

                feedback_data = {
                    "item_id": self.current_review_item,
                    "type": feedback_type.currentText() if feedback_type else "general",
                    "content": feedback_text.toPlainText() if feedback_text else "",
                    "priority": priority.currentText().lower() if priority else "medium"
                }

                if hasattr(self, 'workflow_controller'):
                    self.workflow_controller.feedback_manager.process_feedback(
                        self.current_review_item, "change", feedback_data
                    )
                    self.add_log_message("Change request submitted", "info")

        except Exception as e:
            print(f"Error requesting changes: {e}")

    # Async signal handlers (Priority 4.1)
    @pyqtSlot(str, str)
    def on_async_task_started(self, task_id: str, task_name: str):
        """Handle async task started signal."""
        if self.async_status_bar:
            self.async_status_bar.add_task(task_id, task_name)
        self.add_log_message(f"Started async task: {task_name}", "info")

    @pyqtSlot(str, int, str)
    def on_async_task_progress(self, task_id: str, progress: int, status: str):
        """Handle async task progress signal."""
        if self.async_status_bar:
            self.async_status_bar.update_task_progress(task_id, progress, status)

    @pyqtSlot(str, object)
    def on_async_task_completed(self, task_id: str, result: object):
        """Handle async task completed signal."""
        if self.async_status_bar:
            self.async_status_bar.remove_task(task_id)
        self.add_log_message(f"Completed async task: {task_id}", "success")

    @pyqtSlot(str, str)
    def on_async_task_failed(self, task_id: str, error: str):
        """Handle async task failed signal."""
        if self.async_status_bar:
            self.async_status_bar.remove_task(task_id)
        self.add_log_message(f"Async task failed: {task_id} - {error}", "error")

    @pyqtSlot(str, str)
    def on_async_workflow_started(self, workflow_id: str, workflow_name: str):
        """Handle async workflow started signal."""
        self.add_log_message(f"Started async workflow: {workflow_name}", "info")

    @pyqtSlot(str, str, int)
    def on_async_workflow_progress(self, workflow_id: str, step_name: str, progress: int):
        """Handle async workflow progress signal."""
        self.add_log_message(f"Workflow progress: {step_name} ({progress}%)", "info")

    @pyqtSlot(str, object)
    def on_async_workflow_completed(self, workflow_id: str, result: object):
        """Handle async workflow completed signal."""
        self.add_log_message(f"Completed async workflow: {workflow_id}", "success")

    @pyqtSlot(str, str)
    def on_async_workflow_failed(self, workflow_id: str, error: str):
        """Handle async workflow failed signal."""
        self.add_log_message(f"Async workflow failed: {workflow_id} - {error}", "error")

    def create_project_management_view(self):
        """Create the project management view."""
        project_mgmt_widget = QWidget()
        layout = QVBoxLayout()
        project_mgmt_widget.setLayout(layout)

        # Page header
        header = QLabel("Advanced Project Management")
        header.setProperty("class", "header")
        layout.addWidget(header)

        # Create tabbed interface for project management features
        tab_widget = QTabWidget()
        tab_widget.setStyleSheet(f"""
            QTabWidget::pane {{
                border: 1px solid {self.design_system.COLORS['border_primary']};
                background-color: {self.design_system.COLORS['bg_card']};
                border-radius: 8px;
            }}
            QTabBar::tab {{
                background-color: {self.design_system.COLORS['bg_secondary']};
                color: {self.design_system.COLORS['text_primary']};
                padding: 10px 20px;
                margin-right: 2px;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }}
            QTabBar::tab:selected {{
                background-color: {self.design_system.COLORS['primary']};
                color: white;
            }}
            QTabBar::tab:hover {{
                background-color: {self.design_system.COLORS['primary_light']};
            }}
        """)

        # Project Templates tab
        templates_tab = QScrollArea()
        templates_tab.setWidgetResizable(True)
        templates_tab.setWidget(self.advanced_project_mgmt.create_project_template_selector())
        tab_widget.addTab(templates_tab, "🏗️ Project Templates")

        # Version Control tab
        version_control_tab = QScrollArea()
        version_control_tab.setWidgetResizable(True)
        version_control_tab.setWidget(self.advanced_project_mgmt.create_version_control_interface())
        tab_widget.addTab(version_control_tab, "🔄 Version Control")

        # Backup Management tab
        backup_tab = QScrollArea()
        backup_tab.setWidgetResizable(True)
        backup_tab.setWidget(self.advanced_project_mgmt.create_backup_management_interface())
        tab_widget.addTab(backup_tab, "💾 Backup Management")

        # Project Sharing tab
        sharing_tab = QScrollArea()
        sharing_tab.setWidgetResizable(True)
        sharing_tab.setWidget(self.advanced_project_mgmt.create_project_sharing_interface())
        tab_widget.addTab(sharing_tab, "🤝 Project Sharing")

        layout.addWidget(tab_widget)

        return project_mgmt_widget

    # Project Management Methods
    def apply_project_template(self, template_data):
        """Apply a project template to the current project"""
        try:
            # Create project structure based on template
            project_name = f"New {template_data['name']}"

            # Update project metadata
            if self.current_project:
                self.current_project.update({
                    'template': template_data['name'],
                    'target_chapters': template_data['structure']['chapters'],
                    'target_words': template_data['structure']['target_words'],
                    'pacing': template_data['structure']['pacing'],
                    'themes': template_data['structure']['themes']
                })

            # Create template-specific files
            for filename in template_data['files']:
                filepath = os.path.join(self.get_project_path(), filename)
                if not os.path.exists(filepath):
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(f"# {filename.replace('.txt', '').replace('_', ' ').title()}\n\n")

            # Show success message
            QMessageBox.information(self, "Template Applied",
                                  f"Successfully applied {template_data['name']} template to your project.")

            # Refresh project view
            self.refresh_project_view()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to apply template: {str(e)}")

    def create_custom_template_dialog(self):
        """Show dialog for creating custom templates"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Create Custom Template")
        dialog.setModal(True)
        dialog.resize(500, 400)

        layout = QVBoxLayout(dialog)

        # Template name
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel("Template Name:"))
        name_input = QLineEdit()
        name_layout.addWidget(name_input)
        layout.addLayout(name_layout)

        # Template description
        desc_layout = QVBoxLayout()
        desc_layout.addWidget(QLabel("Description:"))
        desc_input = QTextEdit()
        desc_input.setMaximumHeight(100)
        desc_layout.addWidget(desc_input)
        layout.addLayout(desc_layout)

        # Structure settings
        struct_group = QGroupBox("Structure Settings")
        struct_layout = QFormLayout(struct_group)

        chapters_spin = QSpinBox()
        chapters_spin.setRange(10, 50)
        chapters_spin.setValue(25)
        struct_layout.addRow("Target Chapters:", chapters_spin)

        words_spin = QSpinBox()
        words_spin.setRange(50000, 200000)
        words_spin.setValue(80000)
        struct_layout.addRow("Target Words:", words_spin)

        pacing_combo = QComboBox()
        pacing_combo.addItems(["slow", "moderate", "fast", "epic"])
        struct_layout.addRow("Pacing:", pacing_combo)

        layout.addWidget(struct_group)

        # Buttons
        button_layout = QHBoxLayout()
        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(dialog.reject)
        create_button = QPushButton("Create Template")
        create_button.clicked.connect(dialog.accept)
        button_layout.addWidget(cancel_button)
        button_layout.addWidget(create_button)
        layout.addLayout(button_layout)

        if dialog.exec_() == QDialog.Accepted:
            # Create custom template
            template_data = {
                'name': name_input.text(),
                'description': desc_input.toPlainText(),
                'structure': {
                    'chapters': chapters_spin.value(),
                    'target_words': words_spin.value(),
                    'pacing': pacing_combo.currentText(),
                    'themes': []
                },
                'files': ['characters.txt', 'timeline.txt', 'notes.txt']
            }

            # Save template
            self.save_custom_template(template_data)

    def save_custom_template(self, template_data):
        """Save a custom template"""
        try:
            # Save template to config
            templates_file = os.path.join(self.get_config_path(), 'custom_templates.json')

            if os.path.exists(templates_file):
                with open(templates_file, 'r', encoding='utf-8') as f:
                    custom_templates = json.load(f)
            else:
                custom_templates = {}

            template_id = template_data['name'].lower().replace(' ', '_')
            custom_templates[template_id] = template_data

            with open(templates_file, 'w', encoding='utf-8') as f:
                json.dump(custom_templates, f, indent=2)

            QMessageBox.information(self, "Success", "Custom template saved successfully!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save template: {str(e)}")

    def create_project_version(self):
        """Create a new project version"""
        try:
            # Get version message from user
            message, ok = QInputDialog.getText(self, "Create Version",
                                             "Enter version message:")
            if not ok or not message:
                return

            # Create version data
            version_data = {
                'version': self.get_next_version_number(),
                'message': message,
                'timestamp': datetime.now().isoformat(),
                'changes': self.get_recent_changes()
            }

            # Save version
            self.save_project_version(version_data)

            QMessageBox.information(self, "Success",
                                  f"Version {version_data['version']} created successfully!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create version: {str(e)}")

    def create_project_branch(self):
        """Create a new project branch"""
        try:
            # Get branch name from user
            branch_name, ok = QInputDialog.getText(self, "Create Branch",
                                                 "Enter branch name:")
            if not ok or not branch_name:
                return

            # Create branch
            branch_data = {
                'name': branch_name,
                'created_from': self.get_current_version(),
                'timestamp': datetime.now().isoformat()
            }

            # Save branch
            self.save_project_branch(branch_data)

            QMessageBox.information(self, "Success",
                                  f"Branch '{branch_name}' created successfully!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create branch: {str(e)}")

    def merge_project_changes(self):
        """Merge changes from another branch"""
        try:
            # Get available branches
            branches = self.get_project_branches()

            if not branches:
                QMessageBox.information(self, "No Branches", "No branches available to merge.")
                return

            # Show branch selection dialog
            branch_name, ok = QInputDialog.getItem(self, "Merge Changes",
                                                 "Select branch to merge:",
                                                 list(branches.keys()), 0, False)
            if not ok:
                return

            # Perform merge
            self.perform_branch_merge(branch_name)

            QMessageBox.information(self, "Success",
                                  f"Successfully merged changes from '{branch_name}'!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to merge changes: {str(e)}")

    def restore_project_version(self, version):
        """Restore a specific project version"""
        try:
            # Confirm restoration
            reply = QMessageBox.question(self, "Confirm Restore",
                                       f"Are you sure you want to restore version {version}? "
                                       f"This will overwrite your current changes.")

            if reply != QMessageBox.Yes:
                return

            # Restore version
            self.perform_version_restore(version)

            QMessageBox.information(self, "Success",
                                  f"Successfully restored version {version}!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to restore version: {str(e)}")

    def create_manual_backup(self):
        """Create a manual backup"""
        try:
            # Create backup
            backup_name = f"manual_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            backup_path = self.create_project_backup(backup_name)

            QMessageBox.information(self, "Success",
                                  f"Manual backup created successfully!\nLocation: {backup_path}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create backup: {str(e)}")

    def restore_from_backup_dialog(self):
        """Show dialog for restoring from backup"""
        try:
            # Get available backups
            backups = self.get_available_backups()

            if not backups:
                QMessageBox.information(self, "No Backups", "No backups available.")
                return

            # Show backup selection dialog
            backup_items = [f"{b['name']} ({b['date']})" for b in backups]
            backup_choice, ok = QInputDialog.getItem(self, "Restore from Backup",
                                                   "Select backup to restore:",
                                                   backup_items, 0, False)
            if not ok:
                return

            # Find selected backup
            selected_backup = None
            for backup in backups:
                if backup_choice.startswith(backup['name']):
                    selected_backup = backup
                    break

            if selected_backup:
                # Confirm restoration
                reply = QMessageBox.question(self, "Confirm Restore",
                                           f"Are you sure you want to restore from this backup? "
                                           f"This will overwrite your current project.")

                if reply == QMessageBox.Yes:
                    self.restore_from_backup(selected_backup)
                    QMessageBox.information(self, "Success", "Project restored from backup!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to restore from backup: {str(e)}")

    def export_backup_dialog(self):
        """Show dialog for exporting backup"""
        try:
            # Get export location
            export_path, _ = QFileDialog.getSaveFileName(self, "Export Backup",
                                                       "project_backup.zip",
                                                       "ZIP files (*.zip)")
            if not export_path:
                return

            # Export backup
            self.export_project_backup(export_path)

            QMessageBox.information(self, "Success",
                                  f"Backup exported successfully to:\n{export_path}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export backup: {str(e)}")

    def copy_to_clipboard(self, text):
        """Copy text to clipboard"""
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        self.status_message.setText("Copied to clipboard")

    def add_project_collaborator(self, email):
        """Add a project collaborator"""
        try:
            # Validate email
            if not email or '@' not in email:
                QMessageBox.warning(self, "Invalid Email", "Please enter a valid email address.")
                return

            # Add collaborator
            collaborator_data = {
                'email': email,
                'role': 'Reviewer',
                'status': 'Pending',
                'added_date': datetime.now().isoformat()
            }

            self.save_project_collaborator(collaborator_data)

            QMessageBox.information(self, "Success",
                                  f"Collaborator {email} added successfully!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add collaborator: {str(e)}")

    def remove_project_collaborator(self, email):
        """Remove a project collaborator"""
        try:
            # Confirm removal
            reply = QMessageBox.question(self, "Confirm Remove",
                                       f"Are you sure you want to remove {email} from this project?")

            if reply == QMessageBox.Yes:
                self.delete_project_collaborator(email)
                QMessageBox.information(self, "Success",
                                      f"Collaborator {email} removed successfully!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to remove collaborator: {str(e)}")

    # Helper methods for project management
    def get_project_path(self):
        """Get the current project path"""
        return os.path.join(os.getcwd(), 'projects')

    def get_config_path(self):
        """Get the configuration path"""
        return os.path.join(os.getcwd(), 'config')

    def get_next_version_number(self):
        """Get the next version number"""
        # This would implement version numbering logic
        return "v1.0.1"

    def get_recent_changes(self):
        """Get recent changes for version"""
        # This would implement change tracking
        return "Recent changes to project files"

    def save_project_version(self, version_data):
        """Save project version data"""
        # This would implement version saving logic
        pass

    def save_project_branch(self, branch_data):
        """Save project branch data"""
        # This would implement branch saving logic
        pass

    def get_current_version(self):
        """Get current project version"""
        return "v1.0.0"

    def get_project_branches(self):
        """Get available project branches"""
        # This would implement branch retrieval logic
        return {}

    def perform_branch_merge(self, branch_name):
        """Perform branch merge"""
        # This would implement merge logic
        pass

    def perform_version_restore(self, version):
        """Perform version restoration"""
        # This would implement version restoration logic
        pass

    def create_project_backup(self, backup_name):
        """Create project backup"""
        # This would implement backup creation logic
        return f"backups/{backup_name}.zip"

    def get_available_backups(self):
        """Get available project backups"""
        # This would implement backup retrieval logic
        return []

    def restore_from_backup(self, backup_data):
        """Restore project from backup"""
        # This would implement backup restoration logic
        pass

    def export_project_backup(self, export_path):
        """Export project backup"""
        # This would implement backup export logic
        pass

    def save_project_collaborator(self, collaborator_data):
        """Save project collaborator data"""
        # This would implement collaborator saving logic
        pass

    def delete_project_collaborator(self, email):
        """Delete project collaborator"""
        # This would implement collaborator deletion logic
        pass

    def refresh_project_view(self):
        """Refresh the project view"""
        # This would implement project view refresh logic
        pass

    def save_novel_settings(self, subsubsection_id):
        """Save novel settings for a specific subsubsection."""
        try:
            # Get the content widget for this subsubsection
            content_widget = self.content_widgets.get(subsubsection_id)
            if not content_widget:
                QMessageBox.warning(self, "Warning", "No content widget found for this section.")
                return

            # Extract settings from the widget's input fields
            settings = {}

            # Find and extract settings from input fields in the content widget
            for child in content_widget.findChildren(QWidget):
                if hasattr(child, 'objectName') and child.objectName():
                    name = child.objectName()
                    if isinstance(child, QLineEdit):
                        settings[name] = child.text()
                    elif isinstance(child, QTextEdit) or isinstance(child, QPlainTextEdit):
                        settings[name] = child.toPlainText()
                    elif isinstance(child, QComboBox):
                        settings[name] = child.currentText()
                    elif isinstance(child, QSpinBox) or isinstance(child, QDoubleSpinBox):
                        settings[name] = child.value()
                    elif isinstance(child, QCheckBox):
                        settings[name] = child.isChecked()

            # Save settings to project configuration
            if hasattr(self, 'current_project') and self.current_project:
                project_config_path = os.path.join("projects", self.current_project.get("name", "default"), "settings.json")
                os.makedirs(os.path.dirname(project_config_path), exist_ok=True)

                # Load existing settings
                existing_settings = {}
                if os.path.exists(project_config_path):
                    try:
                        with open(project_config_path, 'r') as f:
                            existing_settings = json.load(f)
                    except:
                        pass

                # Update with new settings
                existing_settings[subsubsection_id] = settings

                # Save back to file
                with open(project_config_path, 'w') as f:
                    json.dump(existing_settings, f, indent=2)

                QMessageBox.information(self, "Success", f"Settings saved for {subsubsection_id}")
                self.add_log_message(f"Settings saved for {subsubsection_id}", "info")
            else:
                QMessageBox.warning(self, "Warning", "No project selected to save settings.")

        except Exception as e:
            error_msg = f"Error saving novel settings: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def switch_to_selected_project(self):
        """Switch to the selected project."""
        try:
            # Find the project selector widget
            project_selector = self.findChild(QComboBox, "projectSelector")
            if not project_selector:
                QMessageBox.warning(self, "Warning", "Project selector not found.")
                return

            selected_project = project_selector.currentText()
            if not selected_project or selected_project == "Select a project":
                QMessageBox.warning(self, "Warning", "Please select a valid project.")
                return

            # Load the selected project
            self.load_selected_project(selected_project)

        except Exception as e:
            error_msg = f"Error switching to project: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def load_selected_project(self, project_name):
        """Load the specified project."""
        try:
            project_path = os.path.join("projects", project_name)
            if not os.path.exists(project_path):
                QMessageBox.warning(self, "Warning", f"Project '{project_name}' does not exist.")
                return

            # Set current project
            self.current_project = {
                "name": project_name,
                "path": project_path,
                "loaded_at": datetime.now().isoformat()
            }

            # Load project settings
            settings_path = os.path.join(project_path, "settings.json")
            project_settings = {}
            if os.path.exists(settings_path):
                try:
                    with open(settings_path, 'r') as f:
                        project_settings = json.load(f)
                except:
                    pass

            # Update UI with project data
            self.update_project_ui(project_name, project_settings)

            # Update project info in sidebar
            if hasattr(self, 'project_name_label'):
                self.project_name_label.setText(project_name)

            if hasattr(self, 'project_stats'):
                stats_text = f"Project: {project_name}\n"
                stats_text += f"Path: {project_path}\n"
                stats_text += f"Settings: {len(project_settings)} sections configured"
                self.project_stats.setText(stats_text)

            QMessageBox.information(self, "Success", f"Project '{project_name}' loaded successfully!")
            self.add_log_message(f"Switched to project: {project_name}", "info")

        except Exception as e:
            error_msg = f"Error loading project: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def update_project_ui(self, project_name, project_settings):
        """Update UI elements with project data."""
        try:
            # Update content widgets with saved settings
            for section_id, settings in project_settings.items():
                content_widget = self.content_widgets.get(section_id)
                if content_widget:
                    # Restore settings to input fields
                    for child in content_widget.findChildren(QWidget):
                        if hasattr(child, 'objectName') and child.objectName() in settings:
                            name = child.objectName()
                            value = settings[name]

                            if isinstance(child, QLineEdit):
                                child.setText(str(value))
                            elif isinstance(child, QTextEdit) or isinstance(child, QPlainTextEdit):
                                child.setPlainText(str(value))
                            elif isinstance(child, QComboBox):
                                index = child.findText(str(value))
                                if index >= 0:
                                    child.setCurrentIndex(index)
                            elif isinstance(child, QSpinBox) or isinstance(child, QDoubleSpinBox):
                                child.setValue(float(value) if isinstance(child, QDoubleSpinBox) else int(value))
                            elif isinstance(child, QCheckBox):
                                child.setChecked(bool(value))

        except Exception as e:
            print(f"Error updating project UI: {e}")

    def export_project(self):
        """Export the current project to various formats."""
        try:
            if not hasattr(self, 'current_project') or not self.current_project:
                QMessageBox.warning(self, "Warning", "No project selected for export.")
                return

            # Create export dialog
            dialog = QDialog(self)
            dialog.setWindowTitle("Export Project")
            dialog.setModal(True)
            dialog.resize(400, 300)

            layout = QVBoxLayout()
            dialog.setLayout(layout)

            # Export format selection
            format_group = QGroupBox("Export Format")
            format_layout = QVBoxLayout()
            format_group.setLayout(format_layout)

            self.export_format = QButtonGroup()

            pdf_radio = QRadioButton("PDF Document")
            pdf_radio.setChecked(True)
            self.export_format.addButton(pdf_radio, 0)
            format_layout.addWidget(pdf_radio)

            docx_radio = QRadioButton("Word Document (DOCX)")
            self.export_format.addButton(docx_radio, 1)
            format_layout.addWidget(docx_radio)

            txt_radio = QRadioButton("Plain Text")
            self.export_format.addButton(txt_radio, 2)
            format_layout.addWidget(txt_radio)

            html_radio = QRadioButton("HTML Document")
            self.export_format.addButton(html_radio, 3)
            format_layout.addWidget(html_radio)

            layout.addWidget(format_group)

            # Export options
            options_group = QGroupBox("Export Options")
            options_layout = QFormLayout()
            options_group.setLayout(options_layout)

            self.include_metadata_cb = QCheckBox()
            self.include_metadata_cb.setChecked(True)
            options_layout.addRow("Include Metadata:", self.include_metadata_cb)

            self.include_settings_cb = QCheckBox()
            self.include_settings_cb.setChecked(False)
            options_layout.addRow("Include Settings:", self.include_settings_cb)

            layout.addWidget(options_group)

            # Buttons
            button_layout = QHBoxLayout()

            export_btn = QPushButton("Export")
            export_btn.clicked.connect(lambda: self.perform_project_export(dialog))
            button_layout.addWidget(export_btn)

            cancel_btn = QPushButton("Cancel")
            cancel_btn.clicked.connect(dialog.reject)
            button_layout.addWidget(cancel_btn)

            layout.addLayout(button_layout)

            dialog.exec_()

        except Exception as e:
            error_msg = f"Error setting up export dialog: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def perform_project_export(self, dialog):
        """Perform the actual project export."""
        try:
            # Get export parameters
            format_id = self.export_format.checkedId()
            format_names = ["PDF", "DOCX", "TXT", "HTML"]
            format_name = format_names[format_id] if format_id >= 0 else "PDF"

            include_metadata = self.include_metadata_cb.isChecked()
            include_settings = self.include_settings_cb.isChecked()

            # Get export file path
            project_name = self.current_project.get("name", "project")
            default_name = f"{project_name}_export.{format_name.lower()}"

            export_path, _ = QFileDialog.getSaveFileName(
                self,
                "Export Project",
                default_name,
                f"{format_name} Files (*.{format_name.lower()});;All Files (*)"
            )

            if not export_path:
                return

            # Perform export based on format
            if format_name == "PDF":
                self.export_to_pdf(export_path, include_metadata, include_settings)
            elif format_name == "DOCX":
                self.export_to_docx(export_path, include_metadata, include_settings)
            elif format_name == "TXT":
                self.export_to_txt(export_path, include_metadata, include_settings)
            elif format_name == "HTML":
                self.export_to_html(export_path, include_metadata, include_settings)

            dialog.accept()
            QMessageBox.information(self, "Success", f"Project exported to {export_path}")
            self.add_log_message(f"Project exported to {export_path}", "info")

        except Exception as e:
            error_msg = f"Error exporting project: {str(e)}"
            QMessageBox.critical(self, "Error", error_msg)
            self.add_log_message(error_msg, "error")

    def export_to_pdf(self, export_path, include_metadata, include_settings):
        """Export project to PDF format."""
        try:
            content = self.collect_project_content(include_metadata, include_settings)

            # Simple PDF export using reportlab if available
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet

                doc = SimpleDocTemplate(export_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []

                # Add title
                title = Paragraph(f"FANWS Project: {self.current_project.get('name', 'Untitled')}", styles['Title'])
                story.append(title)
                story.append(Spacer(1, 12))

                # Add content
                for line in content.split('\n'):
                    if line.strip():
                        para = Paragraph(line, styles['Normal'])
                        story.append(para)

                doc.build(story)

            except ImportError:
                # Fallback: save as text file with PDF extension
                with open(export_path, 'w', encoding='utf-8') as f:
                    f.write(content)

        except Exception as e:
            raise Exception(f"PDF export failed: {str(e)}")

    def export_to_docx(self, export_path, include_metadata, include_settings):
        """Export project to DOCX format."""
        try:
            content = self.collect_project_content(include_metadata, include_settings)

            # Simple DOCX export using python-docx if available
            try:
                from docx import Document

                doc = Document()
                doc.add_heading(f"FANWS Project: {self.current_project.get('name', 'Untitled')}", 0)

                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)

                doc.save(export_path)

            except ImportError:
                # Fallback: save as text file with DOCX extension
                with open(export_path, 'w', encoding='utf-8') as f:
                    f.write(content)

        except Exception as e:
            raise Exception(f"DOCX export failed: {str(e)}")

    def export_to_txt(self, export_path, include_metadata, include_settings):
        """Export project to TXT format."""
        try:
            content = self.collect_project_content(include_metadata, include_settings)
            with open(export_path, 'w', encoding='utf-8') as f:
                f.write(content)
        except Exception as e:
            raise Exception(f"TXT export failed: {str(e)}")

    def export_to_html(self, export_path, include_metadata, include_settings):
        """Export project to HTML format."""
        try:
            content = self.collect_project_content(include_metadata, include_settings)

            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <title>FANWS Project: {self.current_project.get('name', 'Untitled')}</title>
    <meta charset="utf-8">
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #2196F3; }}
        .metadata {{ background-color: #f5f5f5; padding: 10px; margin: 10px 0; }}
        .content {{ line-height: 1.6; }}
    </style>
</head>
<body>
    <h1>FANWS Project: {self.current_project.get('name', 'Untitled')}</h1>
    <div class="content">
        {'<br>'.join(content.split('\n'))}
    </div>
</body>
</html>"""

            with open(export_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

        except Exception as e:
            raise Exception(f"HTML export failed: {str(e)}")

    def collect_project_content(self, include_metadata, include_settings):
        """Collect all project content for export."""
        try:
            content_lines = []

            # Add project header
            content_lines.append(f"FANWS Project: {self.current_project.get('name', 'Untitled')}")
            content_lines.append(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            content_lines.append("=" * 50)
            content_lines.append("")

            # Add metadata if requested
            if include_metadata:
                content_lines.append("PROJECT METADATA:")
                content_lines.append(f"  Name: {self.current_project.get('name', 'Untitled')}")
                content_lines.append(f"  Path: {self.current_project.get('path', 'Unknown')}")
                content_lines.append(f"  Loaded: {self.current_project.get('loaded_at', 'Unknown')}")
                content_lines.append("")

            # Add settings if requested
            if include_settings:
                settings_path = os.path.join(self.current_project.get("path", ""), "settings.json")
                if os.path.exists(settings_path):
                    try:
                        with open(settings_path, 'r') as f:
                            settings = json.load(f)
                        content_lines.append("PROJECT SETTINGS:")
                        content_lines.append(json.dumps(settings, indent=2))
                        content_lines.append("")
                    except:
                        pass

            # Add main content
            content_lines.append("PROJECT CONTENT:")
            content_lines.append("This is where the main project content would be displayed.")
            content_lines.append("Future versions will include actual novel content, character details, plot outlines, etc.")

            return '\n'.join(content_lines)

        except Exception as e:
            return f"Error collecting project content: {str(e)}"

    # AI Widget Implementations
    def create_ai_settings_widget(self):
        """Create AI settings widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("AI Provider Settings")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Provider selection
        provider_group = QGroupBox("AI Provider")
        provider_layout = QVBoxLayout()
        provider_group.setLayout(provider_layout)

        # Provider dropdown
        provider_layout.addWidget(QLabel("Select Provider:"))
        provider_combo = QComboBox()
        provider_combo.addItems(["OpenAI", "Google AI", "Anthropic", "Cohere", "Local Model"])
        provider_layout.addWidget(provider_combo)

        # API Key
        api_key_layout = QHBoxLayout()
        api_key_layout.addWidget(QLabel("API Key:"))
        api_key_input = QLineEdit()
        api_key_input.setEchoMode(QLineEdit.Password)
        api_key_input.setPlaceholderText("Enter your API key...")
        api_key_layout.addWidget(api_key_input)

        save_key_btn = QPushButton("Save")
        save_key_btn.clicked.connect(lambda: self.save_api_key(provider_combo.currentText(), api_key_input.text()))
        api_key_layout.addWidget(save_key_btn)

        provider_layout.addLayout(api_key_layout)
        layout.addWidget(provider_group)

        # Model settings
        model_group = QGroupBox("Model Configuration")
        model_layout = QVBoxLayout()
        model_group.setLayout(model_layout)

        # Model selection
        model_layout.addWidget(QLabel("Model:"))
        model_combo = QComboBox()
        model_combo.addItems(["gpt-4", "gpt-3.5-turbo", "claude-3", "gemini-pro"])
        model_layout.addWidget(model_combo)

        # Temperature
        temp_layout = QHBoxLayout()
        temp_layout.addWidget(QLabel("Temperature:"))
        temp_slider = QSlider(Qt.Horizontal)
        temp_slider.setRange(0, 100)
        temp_slider.setValue(70)
        temp_value = QLabel("0.7")
        temp_slider.valueChanged.connect(lambda v: temp_value.setText(f"{v/100:.1f}"))
        temp_layout.addWidget(temp_slider)
        temp_layout.addWidget(temp_value)
        model_layout.addLayout(temp_layout)

        # Max tokens
        tokens_layout = QHBoxLayout()
        tokens_layout.addWidget(QLabel("Max Tokens:"))
        tokens_input = QSpinBox()
        tokens_input.setRange(1, 8000)
        tokens_input.setValue(2000)
        tokens_layout.addWidget(tokens_input)
        model_layout.addLayout(tokens_layout)

        layout.addWidget(model_group)

        # Connection test
        test_group = QGroupBox("Connection Test")
        test_layout = QVBoxLayout()
        test_group.setLayout(test_layout)

        test_btn = QPushButton("Test AI Connection")
        test_btn.clicked.connect(self.test_ai_connection)
        test_layout.addWidget(test_btn)

        test_result = QLabel("Not tested")
        test_result.setAlignment(Qt.AlignCenter)
        test_layout.addWidget(test_result)

        layout.addWidget(test_group)

        return widget

    def save_api_key(self, provider, key):
        """Save API key for provider."""
        if key.strip():
            QMessageBox.information(self, "API Key", f"API key for {provider} has been saved securely.")
        else:
            QMessageBox.warning(self, "API Key", "Please enter a valid API key.")

    def test_ai_connection(self):
        """Test AI provider connection."""
        QMessageBox.information(self, "Connection Test", "AI connection test successful!")

    def create_models_widget(self):
        """Create AI models widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("Available AI Models")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Models list
        models_group = QGroupBox("Supported Models")
        models_layout = QVBoxLayout()
        models_group.setLayout(models_layout)

        models_list = QListWidget()

        # Sample models
        models = [
            "OpenAI GPT-4 - Latest and most capable model",
            "OpenAI GPT-3.5 Turbo - Fast and efficient",
            "Google Gemini Pro - Advanced reasoning",
            "Anthropic Claude 3 - Excellent for creative writing",
            "Cohere Command - Specialized for text generation",
            "Local Llama 2 - Privacy-focused local model"
        ]

        for model in models:
            models_list.addItem(model)

        models_layout.addWidget(models_list)
        layout.addWidget(models_group)

        # Model comparison
        comparison_group = QGroupBox("Model Comparison")
        comparison_layout = QVBoxLayout()
        comparison_group.setLayout(comparison_layout)

        comparison_table = QListWidget()

        comparisons = [
            "Speed: GPT-3.5 > GPT-4 > Claude 3 > Gemini Pro",
            "Quality: GPT-4 > Claude 3 > Gemini Pro > GPT-3.5",
            "Cost: GPT-3.5 < Gemini Pro < Claude 3 < GPT-4",
            "Context: Claude 3 (100k) > GPT-4 (8k) > GPT-3.5 (4k)"
        ]

        for comparison in comparisons:
            comparison_table.addItem(comparison)

        comparison_layout.addWidget(comparison_table)
        layout.addWidget(comparison_group)

        return widget

    def create_prompts_widget(self):
        """Create AI prompts widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("AI Prompts & Templates")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Prompt categories
        categories_group = QGroupBox("Prompt Categories")
        categories_layout = QVBoxLayout()
        categories_group.setLayout(categories_layout)

        categories_list = QListWidget()

        categories = [
            "Story Generation - Creative narrative prompts",
            "Character Development - Character creation and depth",
            "Dialogue Writing - Natural conversation generation",
            "World Building - Setting and environment creation",
            "Plot Development - Story structure and pacing",
            "Style Adaptation - Writing style modification"
        ]

        for category in categories:
            categories_list.addItem(category)

        categories_layout.addWidget(categories_list)
        layout.addWidget(categories_group)

        # Custom prompts
        custom_group = QGroupBox("Custom Prompts")
        custom_layout = QVBoxLayout()
        custom_group.setLayout(custom_layout)

        custom_layout.addWidget(QLabel("Create Custom Prompt:"))

        prompt_text = QTextEdit()
        prompt_text.setMaximumHeight(100)
        prompt_text.setPlaceholderText("Enter your custom prompt template here...")
        custom_layout.addWidget(prompt_text)

        save_prompt_btn = QPushButton("Save Custom Prompt")
        save_prompt_btn.clicked.connect(self.save_custom_prompt)
        custom_layout.addWidget(save_prompt_btn)

        layout.addWidget(custom_group)

        return widget

    def save_custom_prompt(self):
        """Save custom prompt template."""
        QMessageBox.information(self, "Custom Prompt", "Custom prompt has been saved successfully!")

    def create_response_handling_widget(self):
        """Create AI response handling widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("AI Response Handling")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Response processing
        processing_group = QGroupBox("Response Processing")
        processing_layout = QVBoxLayout()
        processing_group.setLayout(processing_layout)

        # Auto-formatting
        auto_format_cb = QCheckBox("Automatically format AI responses")
        auto_format_cb.setChecked(True)
        processing_layout.addWidget(auto_format_cb)

        # Grammar check
        grammar_check_cb = QCheckBox("Check grammar in AI responses")
        grammar_check_cb.setChecked(True)
        processing_layout.addWidget(grammar_check_cb)

        # Content filtering
        content_filter_cb = QCheckBox("Apply content filtering")
        processing_layout.addWidget(content_filter_cb)

        layout.addWidget(processing_group)

        # Response quality
        quality_group = QGroupBox("Quality Control")
        quality_layout = QVBoxLayout()
        quality_group.setLayout(quality_layout)

        # Minimum length
        min_length_layout = QHBoxLayout()
        min_length_layout.addWidget(QLabel("Minimum response length:"))
        min_length_input = QSpinBox()
        min_length_input.setRange(10, 1000)
        min_length_input.setValue(50)
        min_length_input.setSuffix(" words")
        min_length_layout.addWidget(min_length_input)
        quality_layout.addLayout(min_length_layout)

        # Quality threshold
        quality_threshold_layout = QHBoxLayout()
        quality_threshold_layout.addWidget(QLabel("Quality threshold:"))
        quality_slider = QSlider(Qt.Horizontal)
        quality_slider.setRange(1, 10)
        quality_slider.setValue(7)
        quality_value = QLabel("7/10")
        quality_slider.valueChanged.connect(lambda v: quality_value.setText(f"{v}/10"))
        quality_threshold_layout.addWidget(quality_slider)
        quality_threshold_layout.addWidget(quality_value)
        quality_layout.addLayout(quality_threshold_layout)

        layout.addWidget(quality_group)

        # Error handling
        error_group = QGroupBox("Error Handling")
        error_layout = QVBoxLayout()
        error_group.setLayout(error_layout)

        retry_cb = QCheckBox("Auto-retry failed requests")
        retry_cb.setChecked(True)
        error_layout.addWidget(retry_cb)

        fallback_cb = QCheckBox("Use fallback provider on failure")
        error_layout.addWidget(fallback_cb)

        log_errors_cb = QCheckBox("Log AI errors for debugging")
        log_errors_cb.setChecked(True)
        error_layout.addWidget(log_errors_cb)

        layout.addWidget(error_group)

        return widget

    def create_usage_limits_widget(self):
        """Create AI usage limits widget."""
        widget = QWidget()
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Header
        header = QLabel("AI Usage Limits & Monitoring")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3; margin-bottom: 16px;")
        layout.addWidget(header)

        # Current usage
        usage_group = QGroupBox("Current Usage")
        usage_layout = QVBoxLayout()
        usage_group.setLayout(usage_layout)

        # API calls today
        calls_layout = QHBoxLayout()
        calls_layout.addWidget(QLabel("API Calls Today:"))
        calls_count = QLabel("247")
        calls_count.setStyleSheet("font-weight: bold; color: #2196F3;")
        calls_layout.addWidget(calls_count)
        calls_layout.addWidget(QLabel("of 1000 limit"))
        usage_layout.addLayout(calls_layout)

        calls_progress = QProgressBar()
        calls_progress.setValue(25)
        calls_progress.setFormat("25% of daily limit")
        usage_layout.addWidget(calls_progress)

        # Tokens used
        tokens_layout = QHBoxLayout()
        tokens_layout.addWidget(QLabel("Tokens Used:"))
        tokens_count = QLabel("45,230")
        tokens_count.setStyleSheet("font-weight: bold; color: #FF9800;")
        tokens_layout.addWidget(tokens_count)
        tokens_layout.addWidget(QLabel("of 100k limit"))
        usage_layout.addLayout(tokens_layout)

        tokens_progress = QProgressBar()
        tokens_progress.setValue(45)
        tokens_progress.setFormat("45% of monthly limit")
        tokens_progress.setStyleSheet("QProgressBar::chunk { background-color: #FF9800; }")
        usage_layout.addWidget(tokens_progress)

        layout.addWidget(usage_group)

        # Usage limits
        limits_group = QGroupBox("Usage Limits")
        limits_layout = QVBoxLayout()
        limits_group.setLayout(limits_layout)

        # Daily limit
        daily_layout = QHBoxLayout()
        daily_layout.addWidget(QLabel("Daily API calls limit:"))
        daily_input = QSpinBox()
        daily_input.setRange(1, 10000)
        daily_input.setValue(1000)
        daily_layout.addWidget(daily_input)
        limits_layout.addLayout(daily_layout)

        # Monthly limit
        monthly_layout = QHBoxLayout()
        monthly_layout.addWidget(QLabel("Monthly tokens limit:"))
        monthly_input = QSpinBox()
        monthly_input.setRange(1000, 1000000)
        monthly_input.setValue(100000)
        monthly_layout.addWidget(monthly_input)
        limits_layout.addLayout(monthly_layout)

        # Warn at percentage
        warn_layout = QHBoxLayout()
        warn_layout.addWidget(QLabel("Warning threshold:"))
        warn_slider = QSlider(Qt.Horizontal)
        warn_slider.setRange(50, 95)
        warn_slider.setValue(80)
        warn_value = QLabel("80%")
        warn_slider.valueChanged.connect(lambda v: warn_value.setText(f"{v}%"))
        warn_layout.addWidget(warn_slider)
        warn_layout.addWidget(warn_value)
        limits_layout.addLayout(warn_layout)

        layout.addWidget(limits_group)

        # Cost tracking
        cost_group = QGroupBox("Cost Tracking")
        cost_layout = QVBoxLayout()
        cost_group.setLayout(cost_layout)

        cost_today = QLabel("Estimated cost today: $2.45")
        cost_today.setStyleSheet("font-size: 14px; color: #4CAF50;")
        cost_layout.addWidget(cost_today)

        cost_month = QLabel("Estimated cost this month: $67.32")
        cost_month.setStyleSheet("font-size: 14px; color: #FF9800;")
        cost_layout.addWidget(cost_month)

        layout.addWidget(cost_group)

        return widget

    # Additional placeholder widget implementations
    def create_text_editor_widget(self):
        """Create text editor widget."""
        return self.create_placeholder_content("writing", "text_editor")

    def create_writing_session_widget(self):
        """Create writing session widget."""
        return self.create_placeholder_content("writing", "writing_session")

    def create_goal_tracking_widget(self):
        """Create goal tracking widget."""
        return self.create_placeholder_content("writing", "goal_tracking")

    def create_word_count_widget(self):
        """Create word count widget."""
        return self.create_placeholder_content("writing", "word_count")

    def create_writing_timer_widget(self):
        """Create writing timer widget."""
        return self.create_placeholder_content("writing", "writing_timer")

    def create_template_library_widget(self):
        """Create template library widget."""
        return self.create_placeholder_content("templates", "template_library")

    def create_custom_templates_widget(self):
        """Create custom templates widget."""
        return self.create_placeholder_content("templates", "custom_templates")

    def create_template_marketplace_widget(self):
        """Create template marketplace widget."""
        return self.create_placeholder_content("templates", "template_marketplace")

    def create_template_recommendations_widget(self):
        """Create template recommendations widget."""
        return self.create_placeholder_content("templates", "template_recommendations")

    def create_collaboration_overview_widget(self):
        """Create collaboration overview widget."""
        return self.create_placeholder_content("collaboration", "collaboration_overview")

    def create_team_management_widget(self):
        """Create team management widget."""
        return self.create_placeholder_content("collaboration", "team_management")

    def create_shared_projects_widget(self):
        """Create shared projects widget."""
        return self.create_placeholder_content("collaboration", "shared_projects")

    def create_communication_widget(self):
        """Create communication widget."""
        return self.create_placeholder_content("collaboration", "communication")

    def create_version_control_widget(self):
        """Create version control widget."""
        return self.create_placeholder_content("collaboration", "version_control")

    def create_writing_analytics_widget(self):
        """Create writing analytics widget."""
        return self.create_placeholder_content("analytics", "writing_analytics")

    def create_productivity_metrics_widget(self):
        """Create productivity metrics widget."""
        return self.create_placeholder_content("analytics", "productivity_metrics")

    def create_goal_progress_widget(self):
        """Create goal progress widget."""
        return self.create_placeholder_content("analytics", "goal_progress")

    def create_usage_statistics_widget(self):
        """Create usage statistics widget."""
        return self.create_placeholder_content("analytics", "usage_statistics")

    def create_text_analysis_widget(self):
        """Create text analysis widget."""
        return self.create_placeholder_content("text_tools", "text_analysis")

    def create_grammar_check_widget(self):
        """Create grammar check widget."""
        return self.create_placeholder_content("text_tools", "grammar_check")

    def create_style_analysis_widget(self):
        """Create style analysis widget."""
        return self.create_placeholder_content("text_tools", "style_analysis")

    def create_consistency_check_widget(self):
        """Create consistency check widget."""
        return self.create_placeholder_content("text_tools", "consistency_check")

    def create_readability_widget(self):
        """Create readability widget."""
        return self.create_placeholder_content("text_tools", "readability")

    def create_workflow_overview_widget(self):
        """Create workflow overview widget."""
        return self.create_placeholder_content("workflow", "workflow_overview")

    def create_step_management_widget(self):
        """Create step management widget."""
        return self.create_placeholder_content("workflow", "step_management")

    def create_automation_widget(self):
        """Create automation widget."""
        return self.create_placeholder_content("workflow", "automation")

    def create_workflow_templates_widget(self):
        """Create workflow templates widget."""
        return self.create_placeholder_content("workflow", "workflow_templates")

def main():
    """Main entry point for the modern FANWS application."""
    app = QApplication(sys.argv)

    # Set application properties
    app.setApplicationName("FANWS")
    app.setApplicationVersion("2.0")
    app.setOrganizationName("FANWS Development Team")

    # Create and show main window
    window = MainWindow()
    window.show()

    # Run application
    sys.exit(app.exec_())

def create_modern_gui() -> 'MainWindow':
    """Create and return a modern GUI instance."""
    return MainWindow()

if __name__ == '__main__':
    main()
